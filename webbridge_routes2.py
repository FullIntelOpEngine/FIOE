# webbridge_routes2.py — Third-segment routes for webbridge.py.
# Contains: static HTML, porting, user-service-config, external provider APIs
#           (ContactOut, Apollo, RocketReach, LinkDAPI, ScrapingDog, BrightData),
#           admin VIP config, criteria load/save, and assessment report endpoints.
# This module is imported at the bottom of webbridge_routes.py after all shared
# state (first-half routes, webbridge_cv) is fully defined.
# Circular import is safe because webbridge and webbridge_routes are already in
# sys.modules by the time this file is executed.

import os
import sys
import re
import json
import threading
import time
import uuid
import io
import hashlib
import logging
import heapq
import difflib
import secrets
from csv import DictWriter
from datetime import datetime
from functools import lru_cache, wraps
import requests
from flask import request, send_from_directory, jsonify, abort, Response, stream_with_context
from werkzeug.middleware.dispatcher import DispatcherMiddleware
from werkzeug.security import generate_password_hash, check_password_hash

# ---------------------------------------------------------------------------
# __main__ / module-name fix — same pattern as webbridge_cv.py and webbridge_routes.py
if 'webbridge' not in sys.modules:
    _main = sys.modules.get('__main__')
    if _main is not None and os.path.basename(os.path.normpath(getattr(_main, '__file__', ''))) == 'webbridge.py':
        sys.modules['webbridge'] = _main
# ---------------------------------------------------------------------------

from webbridge import (
    app, logger, genai,
    BASE_DIR, OUTPUT_DIR, SEARCH_XLS_DIR, REPORT_TEMPLATES_DIR,
    BUCKET_COMPANIES, BUCKET_JOB_TITLES,
    SECTORS_INDEX,
    CSE_PAGE_SIZE, CSE_PAGE_DELAY,
    GOOGLE_CSE_API_KEY, GOOGLE_CSE_CX, SEARCH_RESULTS_TARGET,
    GEMINI_API_KEY, GEMINI_SUGGEST_MODEL,
    SINGAPORE_CONTEXT, SEARCH_RULES,
    CV_TRANSLATION_MAX_CHARS, LANG_DETECTION_SAMPLE_LENGTH, CV_ANALYSIS_MAX_CHARS,
    MAX_COMMENT_LENGTH, COMMENT_TRUNCATE_LENGTH,
    ASSESSMENT_EXCELLENT_THRESHOLD, ASSESSMENT_GOOD_THRESHOLD, ASSESSMENT_MODERATE_THRESHOLD,
    CITY_TO_COUNTRY_DATA,
    _CV_ANALYZE_SEMAPHORE, _SINGLE_FILE_MAX,
    _rate, _check_user_rate, _check_gp_rate_limit, _csrf_required, _require_admin, _require_session,
    _user_has_custom_providers,
    _is_pdf_bytes,
    _extract_json_object, _extract_confirmed_skills,
    translate_text_pipeline,
    _infer_region_from_country,
    _find_best_sector_match_for_text, _map_keyword_to_sector_label,
    _compute_search_target,
    _should_overwrite_existing, _ensure_rating_metadata_columns, _ensure_search_indexes,
    _persist_jskillset, _fetch_jskillset, _fetch_jskillset_from_process,
    _sync_login_jskillset_to_process, _sync_criteria_jskillset_to_process,
    _increment_cse_query_count, _increment_gemini_query_count, _load_rate_limits, _save_rate_limits,
    _make_flask_limit,
    _pg_connect, _ensure_admin_columns,
    dedupe,
    _normalize_seniority_single, _map_gemini_seniority_to_dropdown,
    _gemini_talent_pool_suggestion,
    _token_set, _build_sectors_token_index, _is_pharma_company, _sectors_allow_pharma,
    _nllb_available, nllb_translate,
    log_identity, log_infrastructure, log_financial, log_security, log_error, log_approval,
    read_all_logs,
    _APP_LOGGER_AVAILABLE,
    _load_search_provider_config,
    _load_llm_provider_config,
    _load_email_verif_config,
    _load_get_profiles_config,
)

# Names defined in webbridge_routes (first half) that this module uses.
# Imported after webbridge_routes is fully initialised (guaranteed by the
# import chain: webbridge → webbridge_routes → webbridge_routes2).
from webbridge_routes import (
    unified_llm_call_text,
    CRITERIA_OUTPUT_DIR,
    _get_criteria_filepath,
    _role_tag_session_column_ensured,
)


@app.get("/")
def index():
    html_file=os.path.join(BASE_DIR, "AutoSourcing.html")
    if os.path.isfile(html_file): return send_from_directory(BASE_DIR, "AutoSourcing.html")
    return "AutoSourcing WebBridge is running! (AutoSourcing.html not found)", 200

@app.get("/AutoSourcing.html")
def autosourcing_explicit(): return send_from_directory(BASE_DIR, "AutoSourcing.html")

@app.get("/sales_rep_register.html")
def sales_rep_register_html():
    return send_from_directory(BASE_DIR, "sales_rep_register.html")

@app.get('/favicon.ico')
def favicon():
    path=os.path.join(BASE_DIR, 'favicon.ico')
    if not os.path.isfile(path): abort(404)
    return send_from_directory(BASE_DIR, 'favicon.ico', mimetype='image/vnd.microsoft.icon')

# --- START: New Endpoint to serve data_sorter.json ---
@app.get("/data_sorter.json")
def get_data_sorter_json():
    """
    Serve data_sorter.json if present in static folder.
    This allows frontend or other services to access reference lists (JobFamilyRoles, GeoCountries)
    even when data_sorter.py is not active or directly reachable.
    """
    try:
        # Check standard static location relative to BASE_DIR
        static_folder = os.path.join(BASE_DIR, "static")
        filename = "data_sorter.json"
        file_path = os.path.join(static_folder, filename)
        
        if os.path.isfile(file_path):
            return send_from_directory(static_folder, filename, mimetype='application/json')
        else:
            # Fallback check in base dir just in case
            if os.path.isfile(os.path.join(BASE_DIR, filename)):
                return send_from_directory(BASE_DIR, filename, mimetype='application/json')
            
            return jsonify({"error": "data_sorter.json not found"}), 404
    except Exception as e:
        logger.warning(f"Failed to serve data_sorter.json: {e}")
        return jsonify({"error": str(e)}), 500
# --- END: New Endpoint ---

# --- START: Integration of data_sorter.py ---
try:
    import data_sorter
    if hasattr(data_sorter, 'app'):
        # Ensure session compatibility if needed by sharing secret key
        try:
            data_sorter.app.secret_key = app.secret_key
        except Exception:
            pass

        # Mount data_sorter app at /data_sorter prefix
        app.wsgi_app = DispatcherMiddleware(app.wsgi_app, {
            '/data_sorter': data_sorter.app.wsgi_app
        })
        logger.info("Integrated data_sorter app mounted at /data_sorter")
    else:
        logger.warning("data_sorter module found but has no 'app' attribute.")
except ImportError:
    logger.warning("data_sorter.py not found. Skipping integration.")
except Exception as e:
    logger.warning(f"Failed to integrate data_sorter: {e}")
# --- END: Integration ---

def _startup_backfill_role_tag_session():
    """
    One-time startup backfill: for every login row where role_tag is set but
    role_tag_session is NULL, generate a timestamp (NOW()) and transfer it to
    all matching sourcing rows (WHERE username matches AND role_tag matches).

    This handles rows that existed before the role_tag_session column was
    introduced via ALTER TABLE … ADD COLUMN IF NOT EXISTS (which sets NULL for
    pre-existing rows).  Called once when the server process starts.
    """
    try:
        import psycopg2
        pg_host = os.getenv("PGHOST", "localhost")
        pg_port = int(os.getenv("PGPORT", "5432"))
        pg_user = os.getenv("PGUSER", "postgres")
        pg_password = os.getenv("PGPASSWORD", "")
        pg_db = os.getenv("PGDATABASE", "candidate_db")
        conn = _pg_connect()
        cur = conn.cursor()
        try:
            # Ensure columns exist before touching them
            cur.execute("ALTER TABLE login ADD COLUMN IF NOT EXISTS session TIMESTAMPTZ")
            cur.execute("ALTER TABLE sourcing ADD COLUMN IF NOT EXISTS session TIMESTAMPTZ")
            cur.execute("ALTER TABLE sourcing ALTER COLUMN session DROP DEFAULT")
            # Find all login rows with role_tag set but role_tag_session NULL
            cur.execute(
                "SELECT username, role_tag FROM login"
                " WHERE role_tag IS NOT NULL AND role_tag <> '' AND session IS NULL"
            )
            rows = cur.fetchall()
            count = 0
            for username, role_tag in rows:
                if not username:
                    continue
                # Generate a timestamp for this row and write it to login
                cur.execute(
                    "UPDATE login SET session = NOW()"
                    " WHERE username = %s AND role_tag = %s AND session IS NULL"
                    " RETURNING session",
                    (username, role_tag)
                )
                ts_row = cur.fetchone()
                if ts_row and ts_row[0] is not None:
                    # Transfer the same timestamp to sourcing for matching rows
                    cur.execute(
                        "UPDATE sourcing SET session = %s"
                        " WHERE username = %s AND role_tag = %s",
                        (ts_row[0], username, role_tag)
                    )
                    count += 1
            conn.commit()
            if count:
                logger.info(f"[Startup] Backfilled role_tag_session for {count} user(s) missing a session timestamp.")
            else:
                logger.info("[Startup] role_tag_session backfill: no rows needed backfilling (all sessions already set or no role_tag entries found).")
            # Update the flag in webbridge_routes (first half) where route
            # handlers read it from.
            import webbridge_routes as _wr_main
            _wr_main._role_tag_session_column_ensured = True
        finally:
            cur.close()
            conn.close()
    except Exception as e:
        logger.error(f"[Startup] role_tag_session backfill failed: {e}")


_startup_backfill_role_tag_session()

# ── API Porting routes ─────────────────────────────────────────────────────────
import re as _re

_PORTING_INPUT_DIR = os.path.normpath(
    os.getenv("PORTING_INPUT_DIR", os.path.join(os.path.dirname(os.path.abspath(__file__)), "porting_input"))
)
_PORTING_MAPPINGS_DIR = os.path.normpath(
    os.getenv("PORTING_MAPPINGS_DIR", os.path.join(os.path.dirname(os.path.abspath(__file__)), "porting_mappings"))
)
_PROCESS_TABLE_FIELDS = [
    'id','name','company','jobtitle','country','linkedinurl','username','userid',
    'product','sector','jobfamily','geographic','seniority','skillset',
    'sourcingstatus','email','mobile','office','role_tag','experience','cv',
    'education','exp','rating','pic','tenure','comment','vskillset',
    'compensation','lskillset','jskillset',
]

def _porting_safe_name(s):
    return _re.sub(r'[^a-zA-Z0-9_\-]', '_', str(s))

def _porting_get_key() -> bytes:
    """Return a stable 32-byte encryption key.

    Priority:
    1. PORTING_SECRET env var (set by the operator for production use).
    2. Persisted key file  <porting_input>/porting.key  (auto-created on first run).
    """
    secret = os.getenv("PORTING_SECRET", "").strip()
    if secret:
        return (secret + "!" * 32)[:32].encode()[:32]
    # Auto-generate / reuse a persistent random key so restarts stay compatible.
    key_path = os.path.join(_PORTING_INPUT_DIR, "porting.key")
    os.makedirs(_PORTING_INPUT_DIR, exist_ok=True)
    if os.path.exists(key_path):
        with open(key_path, "rb") as fh:
            raw = fh.read()
        if len(raw) >= 32:
            return raw[:32]
        logger.warning("[porting] porting.key is shorter than 32 bytes — regenerating.")
    raw = os.urandom(32)
    import tempfile
    fd, tmp = tempfile.mkstemp(dir=_PORTING_INPUT_DIR)
    try:
        os.write(fd, raw)
        os.close(fd)
        try:
            os.chmod(tmp, 0o600)
        except OSError:
            pass
        os.replace(tmp, key_path)
    except Exception:
        try:
            os.close(fd)
        except OSError:
            pass
        try:
            os.unlink(tmp)
        except OSError:
            pass
        raise
    return raw


def _porting_encrypt(data: bytes) -> bytes:
    """AES-256-GCM encrypt.  Returns nonce(12) + ciphertext + tag(16).
    Auto-installs the 'cryptography' package if it is not already present."""
    try:
        from cryptography.hazmat.primitives.ciphers.aead import AESGCM
    except ImportError:
        import subprocess
        import sys as _sys
        logger.info("[porting] 'cryptography' not found — installing…")
        result = subprocess.run(
            [_sys.executable, "-m", "pip", "install", "cryptography"],
            capture_output=True, text=True,
        )
        if result.returncode != 0:
            logger.error("[porting] pip install cryptography failed: %s", result.stderr)
            raise RuntimeError(
                "The 'cryptography' package is required for encryption. "
                "Install it with: pip install cryptography"
            ) from None
        from cryptography.hazmat.primitives.ciphers.aead import AESGCM
    key = _porting_get_key()
    nonce = os.urandom(12)
    ct = AESGCM(key).encrypt(nonce, data, None)
    return nonce + ct

def _porting_login_required():
    """Return (username, None) or (None, error_response)."""
    username = (request.cookies.get("username") or "").strip()
    if not username:
        return None, (jsonify({"error": "Authentication required"}), 401)
    return username, None

@app.post("/api/porting/upload")
def porting_upload():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        body = request.get_json(silent=True) or {}
        upload_type = body.get("type", "")
        content = body.get("content", "")
        filename = body.get("filename", "")
        if not upload_type or not content:
            return jsonify({"error": "Missing type or content"}), 400
        if upload_type not in ("file", "text"):
            return jsonify({"error": 'type must be "file" or "text"'}), 400
        import base64
        if upload_type == "file":
            raw = base64.b64decode(content)
        else:
            raw = content.encode("utf-8")
        if len(raw) > 1024 * 1024:
            return jsonify({"error": "Content too large (max 1 MB)"}), 413
        safe_fname = os.path.basename(str(filename)).replace(" ", "_") if filename else (
            "upload.env" if upload_type == "file" else "api_keys.txt"
        )
        safe_fname = _re.sub(r'[^a-zA-Z0-9_\-\.]', '_', safe_fname)
        safe_fname = f"{_porting_safe_name(username)}_{int(__import__('time').time()*1000)}_{safe_fname}"
        os.makedirs(_PORTING_INPUT_DIR, exist_ok=True)
        encrypted = _porting_encrypt(raw)
        dest = os.path.join(_PORTING_INPUT_DIR, safe_fname + ".enc")
        with open(dest, "wb") as fh:
            fh.write(encrypted)
        return jsonify({"ok": True, "stored": safe_fname + ".enc"})
    except Exception as exc:
        logger.exception("[porting/upload]")
        return jsonify({"error": "Upload failed", "detail": str(exc)}), 500

@app.post("/api/porting/map")
def porting_map():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        body = request.get_json(silent=True) or {}
        names = body.get("names", [])
        if not isinstance(names, list) or not names:
            return jsonify({"error": "names must be a non-empty array"}), 400
        fields_str = ", ".join(_PROCESS_TABLE_FIELDS)
        names_str = ", ".join(f'"{str(n)}"' for n in names)
        prompt = (
            f'You are a database field mapping assistant.\n'
            f'Available target fields (PostgreSQL "process" table): {fields_str}\n\n'
            f'Map each of the following external API field names to the SINGLE best-matching target field.\n'
            f'If there is no reasonable match, use null.\n'
            f'Return ONLY a JSON object (no markdown, no explanation) where each key is the input name and '
            f'each value is the matching target field name or null.\n\n'
            f'Input names: {names_str}'
        )
        raw = (unified_llm_call_text(prompt) or "").strip()
        if not raw:
            return jsonify({"error": "No LLM provider configured."}), 500
        _increment_gemini_query_count(username)
        raw = _re.sub(r'^```(?:json)?\s*', '', raw, flags=_re.IGNORECASE)
        raw = _re.sub(r'\s*```$', '', raw).strip()
        try:
            mapping = json.loads(raw)
        except Exception:
            return jsonify({"error": "LLM returned invalid JSON", "raw": raw}), 500
        cleaned = {k: (v if v and v in _PROCESS_TABLE_FIELDS else None) for k, v in mapping.items()}
        return jsonify({"ok": True, "mapping": cleaned})
    except Exception as exc:
        logger.exception("[porting/map]")
        return jsonify({"error": "Mapping failed", "detail": str(exc)}), 500

@app.post("/api/porting/confirm")
def porting_confirm():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        body = request.get_json(silent=True) or {}
        mapping = body.get("mapping")
        if not mapping or not isinstance(mapping, dict):
            return jsonify({"error": "mapping is required"}), 400
        for k, v in mapping.items():
            if v is not None and v not in _PROCESS_TABLE_FIELDS:
                return jsonify({"error": f"Invalid target field: {v}"}), 400
        os.makedirs(_PORTING_MAPPINGS_DIR, exist_ok=True)
        path_out = os.path.join(_PORTING_MAPPINGS_DIR, _porting_safe_name(username) + ".json")
        with open(path_out, "w", encoding="utf-8") as fh:
            json.dump({"username": username, "mapping": mapping}, fh, indent=2)
        return jsonify({"ok": True})
    except Exception as exc:
        logger.exception("[porting/confirm]")
        return jsonify({"error": "Confirm failed", "detail": str(exc)}), 500

@app.get("/api/porting/mapping")
def porting_get_mapping():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        path_in = os.path.join(_PORTING_MAPPINGS_DIR, _porting_safe_name(username) + ".json")
        if not os.path.isfile(path_in):
            return jsonify({"mapping": None})
        with open(path_in, encoding="utf-8") as fh:
            data = json.load(fh)
        return jsonify({"mapping": data.get("mapping")})
    except Exception as exc:
        logger.exception("[porting/mapping]")
        return jsonify({"error": "Could not load mapping", "detail": str(exc)}), 500

@app.post("/api/porting/export")
def porting_export():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        path_map = os.path.join(_PORTING_MAPPINGS_DIR, _porting_safe_name(username) + ".json")
        if not os.path.isfile(path_map):
            return jsonify({"error": "No confirmed mapping found. Please complete the mapping step first."}), 400
        with open(path_map, encoding="utf-8") as fh:
            mapping = json.load(fh).get("mapping", {})
        cols = [c for c in _PROCESS_TABLE_FIELDS if c not in ("cv", "pic")]
        conn = _pg_connect()
        try:
            cur = conn.cursor()
            col_sql = ", ".join(f'"{c}"' for c in cols)
            cur.execute(f'SELECT {col_sql} FROM "process" WHERE username = %s', (username,))
            rows = cur.fetchall()
            cur.close()
        finally:
            conn.close()
        if not rows:
            return jsonify({"error": "No data found for this user in the process table."}), 404
        reverse_map = {proc: ext for ext, proc in mapping.items() if proc}
        exported = [
            {reverse_map.get(col, col): (row[i] if row[i] is not None else None) for i, col in enumerate(cols)}
            for row in rows
        ]
        json_str = json.dumps(exported, indent=2, default=str)
        body_req = request.get_json(silent=True) or {}
        target_url = body_req.get("targetUrl", "")
        if target_url:
            try:
                import urllib.parse as _up
                import urllib.request as _ur
                import ipaddress as _ipaddr
                import socket as _sock
                parsed = _up.urlparse(target_url)
                if parsed.scheme not in ("http", "https"):
                    raise ValueError("targetUrl must use http or https scheme")
                _host = parsed.hostname or ""
                if not _host:
                    raise ValueError("targetUrl must include a hostname")
                # Block requests to loopback / link-local / private ranges (SSRF guard)
                try:
                    _resolved = _sock.getaddrinfo(_host, None, proto=_sock.IPPROTO_TCP)
                    for _af, _st, _pr, _cn, _sa in _resolved:
                        _ip = _ipaddr.ip_address(_sa[0])
                        if _ip.is_loopback or _ip.is_private or _ip.is_link_local or _ip.is_reserved:
                            raise ValueError(f"targetUrl resolves to a disallowed address: {_sa[0]}")
                except _sock.gaierror:
                    raise ValueError("targetUrl hostname could not be resolved")
                # Reconstruct URL from parsed (and validated) components so the
                # request is made to a known-safe value, not the raw user string.
                _safe_url = _up.urlunparse((
                    parsed.scheme, parsed.netloc,
                    parsed.path, parsed.params, parsed.query, ""
                ))
                req_obj = _ur.Request(
                    _safe_url,
                    data=json_str.encode("utf-8"),
                    headers={"Content-Type": "application/json"},
                    method="POST",
                )
                with _ur.urlopen(req_obj, timeout=15):
                    pass
            except Exception as push_err:
                logger.warning(f"[porting/export] push to {target_url} failed: {push_err}")
        from flask import make_response
        resp = make_response(json_str)
        resp.headers["Content-Type"] = "application/json"
        resp.headers["Content-Disposition"] = f'attachment; filename="porting_export_{_porting_safe_name(username)}.json"'
        log_approval(action="export_pdf_triggered", username=username,
                     detail=f"Data export triggered; {len(exported)} row(s)")
        return resp
    except Exception as exc:
        logger.exception("[porting/export]")
        log_error(source="porting_export", message=str(exc), severity="error",
                  username=username, endpoint="/api/porting/export")
        return jsonify({"error": "Export failed", "detail": str(exc)}), 500


# ── BYOK (Bring Your Own Keys) routes ─────────────────────────────────────────
_BYOK_REQUIRED_KEYS = [
    'GEMINI_API_KEY', 'GOOGLE_CSE_API_KEY', 'GOOGLE_API_KEY',
    'GOOGLE_CSE_CX', 'GOOGLE_CLIENT_ID', 'GOOGLE_CLIENT_SECRET',
]

def _byok_path(username: str) -> str:
    byok_dir = os.path.join(_PORTING_INPUT_DIR, 'byok')
    os.makedirs(byok_dir, exist_ok=True)
    return os.path.join(byok_dir, _porting_safe_name(username) + '.enc')


@app.post("/api/porting/byok/activate")
def byok_activate():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        body = request.get_json(silent=True) or {}
        keys = {}
        missing = []
        for k in _BYOK_REQUIRED_KEYS:
            val = str(body.get(k, '')).strip()
            if not val:
                missing.append(k)
            else:
                keys[k] = val
        if missing:
            return jsonify({"error": f"Missing required keys: {', '.join(missing)}"}), 400
        raw = json.dumps({'username': username, 'keys': keys}).encode('utf-8')
        encrypted = _porting_encrypt(raw)
        dest = _byok_path(username)
        with open(dest, 'wb') as fh:
            fh.write(encrypted)
        log_infrastructure("byok_activated", username=username,
                           detail="BYOK keys activated", status="success")
        return jsonify({"ok": True, "byok_active": True})
    except Exception as exc:
        logger.exception("[porting/byok/activate]")
        log_error(source="byok_activate", message=str(exc), severity="error",
                  username=username, endpoint="/api/porting/byok/activate")
        return jsonify({"error": "BYOK activation failed", "detail": str(exc)}), 500


@app.get("/api/porting/byok/status")
def byok_status():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        active = os.path.isfile(_byok_path(username))
        return jsonify({"byok_active": active})
    except Exception as exc:
        logger.exception("[porting/byok/status]")
        return jsonify({"error": "Could not check BYOK status", "detail": str(exc)}), 500


@app.get("/api/porting/credentials/status")
def porting_credentials_status():
    """Return whether the user has any uploaded credential files on file."""
    username, err = _porting_login_required()
    if err:
        return err
    try:
        safe_prefix = _porting_safe_name(username) + "_"
        has_creds = any(
            f.startswith(safe_prefix) and f.endswith(".enc")
            for f in os.listdir(_PORTING_INPUT_DIR)
        ) if os.path.isdir(_PORTING_INPUT_DIR) else False
        return jsonify({"credentials_on_file": has_creds})
    except Exception as exc:
        logger.exception("[porting/credentials/status]")
        return jsonify({"error": "Could not check credential status", "detail": str(exc)}), 500


@app.delete("/api/porting/byok/deactivate")
def byok_deactivate():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        dest = _byok_path(username)
        if os.path.isfile(dest):
            os.remove(dest)
        log_infrastructure(
            "byok_deactivated",
            username=username,
            detail="BYOK keys file removed",
            status="success",
            key_type="ALL",
            deactivation_reason="manual",
        )
        return jsonify({"ok": True, "byok_active": False})
    except Exception as exc:
        logger.exception("[porting/byok/deactivate]")
        return jsonify({"error": "Could not deactivate BYOK", "detail": str(exc)}), 500


@app.post("/api/porting/byok/validate")
def byok_validate():
    """Validate BYOK keys by probing live Google Cloud APIs + checking credential formats.
    Steps:
      1. Gemini API  — list models (validates GEMINI_API_KEY + billing)
      2. Custom Search API — single query (validates GOOGLE_CSE_API_KEY + GOOGLE_CSE_CX)
      3. GOOGLE_API_KEY format check
      4. OAuth client credential format check
    Returns a structured results array without storing anything."""
    username, err = _porting_login_required()
    if err:
        return err
    try:
        import re
        import urllib.request as _ureq
        import urllib.parse as _uparse

        body = request.get_json(silent=True) or {}
        keys = {}
        missing = []
        for k in _BYOK_REQUIRED_KEYS:
            raw = body.get(k)
            if not isinstance(raw, (str, int, float)):
                missing.append(k); continue
            val = str(raw).strip()
            if not val or len(val) > 512:
                missing.append(k)
            else:
                keys[k] = val
        if missing:
            return jsonify({"error": f"Missing required keys: {', '.join(missing)}"}), 400

        def _probe(url, timeout=8):
            """GET url; returns (http_status_or_None, body_text)."""
            try:
                with _ureq.urlopen(url, timeout=timeout) as resp:
                    return resp.status, resp.read().decode('utf-8', errors='replace')
            except Exception as exc:
                if hasattr(exc, 'code'):
                    try:
                        return exc.code, exc.read().decode('utf-8', errors='replace')
                    except Exception:
                        return exc.code, ''
                return None, str(exc)

        def _err_msg(body_text, fallback):
            try:
                return json.loads(body_text).get('error', {}).get('message', fallback)
            except Exception:
                return fallback

        results = []

        # ── Step 1: Gemini API (GEMINI_API_KEY + billing) ────────────────────────
        gemini_url = (
            "https://generativelanguage.googleapis.com/v1beta/models?key="
            + _uparse.quote(keys['GEMINI_API_KEY'], safe='')
        )
        status, body_text = _probe(gemini_url)
        if status == 200:
            results.append({'step': 'gemini', 'label': 'Gemini API', 'status': 'ok',
                            'detail': 'API key is valid and billing is active.'})
        elif status == 403:
            results.append({'step': 'gemini', 'label': 'Gemini API', 'status': 'error',
                            'detail': _err_msg(body_text, 'Gemini API is not enabled or billing is inactive on this project.'),
                            'consoleUrl': 'https://console.cloud.google.com/apis/library/generativelanguage.googleapis.com'})
        elif status == 400:
            results.append({'step': 'gemini', 'label': 'Gemini API', 'status': 'error',
                            'detail': _err_msg(body_text, 'Invalid GEMINI_API_KEY.'),
                            'consoleUrl': 'https://console.cloud.google.com/apis/credentials'})
        else:
            detail = f'Unexpected HTTP {status}' if status else f'Could not reach Google APIs: {body_text}'
            results.append({'step': 'gemini', 'label': 'Gemini API', 'status': 'warn', 'detail': detail})

        # ── Step 2: Custom Search API (GOOGLE_CSE_API_KEY + GOOGLE_CSE_CX) ───────
        cse_url = (
            "https://customsearch.googleapis.com/customsearch/v1?key="
            + _uparse.quote(keys['GOOGLE_CSE_API_KEY'], safe='')
            + "&cx=" + _uparse.quote(keys['GOOGLE_CSE_CX'], safe='')
            + "&q=test&num=1"
        )
        status, body_text = _probe(cse_url)
        if status == 200:
            results.append({'step': 'cse', 'label': 'Custom Search API', 'status': 'ok',
                            'detail': 'CSE API key and Search Engine ID are valid.'})
        elif status == 403:
            results.append({'step': 'cse', 'label': 'Custom Search API', 'status': 'error',
                            'detail': _err_msg(body_text, 'Custom Search API is not enabled or billing is required.'),
                            'consoleUrl': 'https://console.cloud.google.com/apis/library/customsearch.googleapis.com'})
        elif status == 400:
            results.append({'step': 'cse', 'label': 'Custom Search API', 'status': 'error',
                            'detail': _err_msg(body_text, 'Invalid GOOGLE_CSE_API_KEY or GOOGLE_CSE_CX Search Engine ID.'),
                            'consoleUrl': 'https://console.cloud.google.com/apis/credentials'})
        else:
            detail = f'Unexpected HTTP {status}' if status else f'Could not reach Custom Search API: {body_text}'
            results.append({'step': 'cse', 'label': 'Custom Search API', 'status': 'warn', 'detail': detail})

        # ── Step 3: GOOGLE_API_KEY format ─────────────────────────────────────────
        google_api_key_ok = bool(re.fullmatch(r'AIza[0-9A-Za-z\-_]{35}', keys['GOOGLE_API_KEY']))
        results.append({
            'step': 'google_api_key', 'label': 'GOOGLE_API_KEY Format',
            'status': 'ok' if google_api_key_ok else 'warn',
            'detail': ('Key format is valid (AIza… 39-character format).' if google_api_key_ok
                       else 'Key format looks unusual — expected a 39-character key starting with "AIza".'),
            **({'consoleUrl': 'https://console.cloud.google.com/apis/credentials'} if not google_api_key_ok else {}),
        })

        # ── Step 4: OAuth client credentials ──────────────────────────────────────
        client_id_ok = bool(re.fullmatch(r'\d+-[a-zA-Z0-9]+\.apps\.googleusercontent\.com', keys['GOOGLE_CLIENT_ID']))
        client_secret_ok = bool(re.match(r'^(GOCSPX-[A-Za-z0-9_\-]{28,}|[A-Za-z0-9_\-]{24,})$', keys['GOOGLE_CLIENT_SECRET']))
        if not client_id_ok:
            results.append({'step': 'oauth', 'label': 'OAuth Client Credentials', 'status': 'error',
                            'detail': 'GOOGLE_CLIENT_ID must have the format <numbers>-<id>.apps.googleusercontent.com',
                            'consoleUrl': 'https://console.cloud.google.com/apis/credentials'})
        elif not client_secret_ok:
            results.append({'step': 'oauth', 'label': 'OAuth Client Credentials', 'status': 'warn',
                            'detail': 'GOOGLE_CLIENT_SECRET format looks unusual (expected "GOCSPX-…"). Verify it was copied from Google Cloud Console → Credentials → OAuth 2.0 Client.',
                            'consoleUrl': 'https://console.cloud.google.com/apis/credentials'})
        else:
            results.append({'step': 'oauth', 'label': 'OAuth Client Credentials', 'status': 'ok',
                            'detail': 'Client ID and Client Secret formats are valid.'})

        all_ok = all(r['status'] in ('ok', 'warn') for r in results)
        overall_status = "success" if all_ok else "fail"
        failed_steps = [r['label'] for r in results if r['status'] == 'error']
        log_infrastructure("byok_validation", username=username,
                           detail="; ".join(failed_steps) if failed_steps else "All checks passed",
                           status=overall_status)
        return jsonify({'ok': all_ok, 'results': results})
    except Exception as exc:
        logger.exception("[porting/byok/validate]")
        log_error(source="byok_validate", message=str(exc), severity="error",
                  username=username, endpoint="/api/porting/byok/validate")
        return jsonify({"error": "Validation failed", "detail": str(exc)}), 500


# ── Per-User Service Config (Option A) ───────────────────────────────────────
# Encrypted per-user storage for Search Engine / LLM / Email Verification keys.
# Keys are AES-256-GCM encrypted using the same _porting_get_key() helper.
# Each user's config is stored at <PORTING_INPUT_DIR>/user-services/<username>.enc
# Format matches server.js: IV(16) + tag(16) + ciphertext so server.js can decrypt it.

def _svc_config_path(username: str) -> str:
    svc_dir = os.path.join(_PORTING_INPUT_DIR, 'user-services')
    os.makedirs(svc_dir, exist_ok=True)
    return os.path.join(svc_dir, _porting_safe_name(username) + '.enc')


def _svc_config_json_path(username: str) -> str:
    """Plaintext JSON fallback path (when PORTING_SECRET is not set) — matches server.js."""
    svc_dir = os.path.join(_PORTING_INPUT_DIR, 'user-services')
    os.makedirs(svc_dir, exist_ok=True)
    return os.path.join(svc_dir, _porting_safe_name(username) + '.json')


def _load_user_gp_cfg(req_username: str) -> dict:
    """Return the get_profile section from *req_username*'s per-user service config, or {}.

    Returns {} when the user has no config, cannot be loaded, or has
    ``get_profile.provider == 'platform'`` (meaning "use global platform config").
    Used by the GP endpoints to allow VIP / user-defined keys to override the
    global platform API keys.
    """
    if not req_username:
        return {}
    safe = _porting_safe_name(req_username)
    svc_dir   = os.path.realpath(os.path.join(_PORTING_INPUT_DIR, 'user-services'))
    enc_path  = _svc_config_path(safe)
    json_path = _svc_config_json_path(safe)
    # Confinement: ensure resolved paths stay within the expected directory.
    if (not os.path.realpath(enc_path).startswith(svc_dir + os.sep) and
            os.path.realpath(enc_path) != svc_dir):
        return {}
    if (not os.path.realpath(json_path).startswith(svc_dir + os.sep) and
            os.path.realpath(json_path) != svc_dir):
        return {}
    u_cfg = None
    if os.path.isfile(enc_path):
        try:
            with open(enc_path, 'rb') as _fh:
                u_cfg = json.loads(_svc_config_decrypt(_fh.read()).decode('utf-8'))
        except Exception:
            pass
    if u_cfg is None and os.path.isfile(json_path):
        try:
            with open(json_path, 'r', encoding='utf-8') as _fh:
                u_cfg = json.load(_fh)
        except Exception:
            pass
    if not isinstance(u_cfg, dict):
        return {}
    gp = u_cfg.get('get_profile') or {}
    provider = (gp.get('provider') or 'platform').strip()
    if provider == 'platform':
        return {}
    return gp


def _svc_config_encrypt(data: bytes) -> bytes:
    """AES-256-GCM encrypt in Node.js-compatible format: IV(16) + tag(16) + ciphertext."""
    from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
    key = _porting_get_key()
    iv = os.urandom(16)  # 16-byte IV to match server.js
    cipher = Cipher(algorithms.AES(key), modes.GCM(iv))
    encryptor = cipher.encryptor()
    ct = encryptor.update(data) + encryptor.finalize()
    tag = encryptor.tag
    return iv + tag + ct  # IV(16) + tag(16) + ciphertext


def _svc_config_decrypt(data: bytes) -> bytes:
    """AES-256-GCM decrypt in Node.js-compatible format: IV(16) + tag(16) + ciphertext."""
    if len(data) < 33:  # 16-byte IV + 16-byte tag + at least 1 byte
        raise ValueError("Encrypted data too short")
    from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
    key = _porting_get_key()
    iv = data[:16]
    tag = data[16:32]
    ct = data[32:]
    cipher = Cipher(algorithms.AES(key), modes.GCM(iv, tag))
    decryptor = cipher.decryptor()
    return decryptor.update(ct) + decryptor.finalize()


def _porting_decrypt(data: bytes) -> bytes:
    """AES-256-GCM decrypt.  Expects nonce(12) + ciphertext + tag(16)."""
    if len(data) < 28:  # 12-byte nonce + 16-byte tag minimum
        raise ValueError("Encrypted data too short")
    from cryptography.hazmat.primitives.ciphers.aead import AESGCM
    key = _porting_get_key()
    nonce = data[:12]
    ct = data[12:]
    return AESGCM(key).decrypt(nonce, ct, None)


@app.get("/api/user-service-config/status")
def user_svc_config_status():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        stored = None
        # Try encrypted .enc first, then plaintext .json (matches server.js readUserServiceConfig)
        enc_path = _svc_config_path(username)
        json_path = _svc_config_json_path(username)
        if os.path.isfile(enc_path):
            try:
                with open(enc_path, 'rb') as fh:
                    raw = fh.read()
                decrypted = _svc_config_decrypt(raw)
                stored = json.loads(decrypted.decode('utf-8'))
            except Exception:
                logger.warning("[user-service-config/status] .enc decrypt failed for %s — trying .json", username, exc_info=True)
        if stored is None and os.path.isfile(json_path):
            try:
                with open(json_path, 'r', encoding='utf-8') as fh:
                    stored = json.load(fh)
            except Exception:
                logger.warning("[user-service-config/status] .json parse failed for %s", username, exc_info=True)
        if stored is None:
            return jsonify({"active": False})
        providers = {
            'search': stored.get('search', {}).get('provider', 'google_cse'),
            'llm': stored.get('llm', {}).get('provider', 'gemini'),
            'email_verif': stored.get('email_verif', {}).get('provider', 'default'),
            'contact_gen': stored.get('contact_gen', {}).get('provider', 'gemini'),
        }
        return jsonify({"active": True, "providers": providers})
    except Exception as exc:
        logger.exception("[user-service-config/status]")
        return jsonify({"error": "Could not retrieve service config status"}), 500


@app.get("/api/user-service-config/search-keys")
def user_svc_config_search_keys():
    """Return decrypted search credentials for the authenticated user.
    Used by AutoSourcing.html to inject per-user search keys into the /start_job payload."""
    username, err = _porting_login_required()
    if err:
        return err
    try:
        cfg = None
        # Try .enc then .json (matches server.js readUserServiceConfig)
        enc_path = _svc_config_path(username)
        json_path = _svc_config_json_path(username)
        if os.path.isfile(enc_path):
            try:
                with open(enc_path, 'rb') as fh:
                    raw = fh.read()
                decrypted = _svc_config_decrypt(raw)
                cfg = json.loads(decrypted.decode('utf-8'))
            except Exception:
                logger.warning("[user-service-config/search-keys] .enc decrypt/parse failed for %s", username, exc_info=True)
        if cfg is None and os.path.isfile(json_path):
            try:
                with open(json_path, 'r', encoding='utf-8') as fh:
                    cfg = json.load(fh)
            except Exception:
                logger.warning("[user-service-config/search-keys] .json parse failed for %s", username, exc_info=True)
        if cfg is None:
            return jsonify({"provider": "google_cse"})
        search = cfg.get('search', {})
        result = {"provider": search.get('provider', 'google_cse')}
        if search.get('provider') == 'serper' and search.get('SERPER_API_KEY'):
            result['SERPER_API_KEY'] = search['SERPER_API_KEY']
        if search.get('provider') == 'dataforseo' and search.get('DATAFORSEO_LOGIN'):
            result['DATAFORSEO_LOGIN'] = search['DATAFORSEO_LOGIN']
        if search.get('provider') == 'dataforseo' and search.get('DATAFORSEO_PASSWORD'):
            result['DATAFORSEO_PASSWORD'] = search['DATAFORSEO_PASSWORD']
        if search.get('provider') == 'linkedin' and search.get('LINKEDIN_API_KEY'):
            result['LINKEDIN_API_KEY'] = search['LINKEDIN_API_KEY']
        return jsonify(result)
    except Exception as exc:
        logger.exception("[user-service-config/search-keys]")
        return jsonify({"provider": "google_cse"})


@app.post("/api/user-service-config/validate")
def user_svc_config_validate():
    """Validate provided keys by calling each service's API. Does NOT store anything."""
    username, err = _porting_login_required()
    if err:
        return err
    try:
        body = request.get_json(silent=True) or {}
        results = _run_svc_config_validation(body)
        has_error = any(r['status'] == 'error' for r in results)
        return jsonify({'ok': not has_error, 'results': results})
    except Exception:
        logger.exception("[user-service-config/validate]")
        return jsonify({"error": "Validation failed"}), 500


def _run_svc_config_validation(body: dict) -> list:
    """Run API-key validation for all service sections in *body*.
    Returns a list of result dicts: {label, status, detail}.
    Does NOT store anything and has no Flask request context dependency."""
    import urllib.request as _ureq2
    import urllib.parse as _uparse2
    import base64

    search = body.get('search') or {}
    llm = body.get('llm') or {}
    email_verif = body.get('email_verif') or {}
    contact_gen = body.get('contact_gen') or {}

    def _probe_get(url, headers=None, timeout=8):
        req = _ureq2.Request(url, headers=headers or {})
        try:
            with _ureq2.urlopen(req, timeout=timeout) as resp:
                return resp.status, resp.read().decode('utf-8', errors='replace')
        except Exception as exc2:
            if hasattr(exc2, 'code'):
                try:
                    return exc2.code, exc2.read().decode('utf-8', errors='replace')
                except Exception:
                    return exc2.code, ''
            return None, ''

    def _probe_post(url, data, headers=None, timeout=8):
        req = _ureq2.Request(url, data=data, headers=headers or {}, method='POST')
        try:
            with _ureq2.urlopen(req, timeout=timeout) as resp:
                return resp.status, resp.read().decode('utf-8', errors='replace')
        except Exception as exc2:
            if hasattr(exc2, 'code'):
                try:
                    return exc2.code, exc2.read().decode('utf-8', errors='replace')
                except Exception:
                    return exc2.code, ''
            return None, ''

    results = []

    # ── Search Engine ──────────────────────────────────────────────────────
    sp = (search.get('provider') or '').strip()
    if sp == 'google_cse' or not sp:
        results.append({'label': 'Search Engine', 'status': 'ok',
                        'detail': 'Using platform Google CSE — no custom key required.'})
    elif sp == 'serper':
        key = (search.get('SERPER_API_KEY') or '').strip()
        if not key:
            results.append({'label': 'Serper.dev', 'status': 'error',
                            'detail': 'SERPER_API_KEY is required.'})
        else:
            payload = json.dumps({'q': 'test', 'num': 1}).encode('utf-8')
            status, _ = _probe_post('https://google.serper.dev/search', payload,
                                    headers={'X-API-KEY': key, 'Content-Type': 'application/json'})
            if status == 200:
                results.append({'label': 'Serper.dev', 'status': 'ok', 'detail': 'API key is valid.'})
            elif status in (401, 403):
                results.append({'label': 'Serper.dev', 'status': 'error',
                                'detail': f'Authentication failed (HTTP {status}). Check your SERPER_API_KEY.'})
            else:
                results.append({'label': 'Serper.dev', 'status': 'warn',
                                'detail': f'Unexpected HTTP {status} — key may be valid but quota or plan issue possible.'
                                if status else 'Could not reach Serper API.'})
    elif sp == 'dataforseo':
        login = (search.get('DATAFORSEO_LOGIN') or '').strip()
        pwd   = (search.get('DATAFORSEO_PASSWORD') or '').strip()
        if not login or not pwd:
            results.append({'label': 'DataforSEO', 'status': 'error',
                            'detail': 'DATAFORSEO_LOGIN and DATAFORSEO_PASSWORD are both required.'})
        else:
            auth = base64.b64encode(f'{login}:{pwd}'.encode('utf-8')).decode('ascii')
            status, _ = _probe_get(
                'https://api.dataforseo.com/v3/appendix/user_data',
                headers={'Authorization': f'Basic {auth}'}
            )
            if status == 200:
                results.append({'label': 'DataforSEO', 'status': 'ok', 'detail': 'Credentials are valid.'})
            elif status in (401, 403):
                results.append({'label': 'DataforSEO', 'status': 'error',
                                'detail': f'Authentication failed (HTTP {status}). Check login/password.'})
            else:
                results.append({'label': 'DataforSEO', 'status': 'warn',
                                'detail': f'Unexpected HTTP {status}.' if status else 'Could not reach DataforSEO API.'})
    elif sp == 'linkedin':
        key = (search.get('LINKEDIN_API_KEY') or '').strip()
        if not key:
            results.append({'label': 'LinkedIn', 'status': 'error',
                            'detail': 'LINKEDIN_API_KEY is required.'})
        else:
            payload = json.dumps({'q': 'test', 'num': 1}).encode('utf-8')
            status, _ = _probe_post('https://api.linkedapi.io/v1/search', payload,
                                    headers={'X-API-KEY': key, 'Content-Type': 'application/json'})
            if status == 200:
                results.append({'label': 'LinkedIn', 'status': 'ok', 'detail': 'API key is valid.'})
            elif status in (401, 403):
                results.append({'label': 'LinkedIn', 'status': 'error',
                                'detail': f'Authentication failed (HTTP {status}). Check your LINKEDIN_API_KEY.'})
            else:
                results.append({'label': 'LinkedIn', 'status': 'warn',
                                'detail': f'Unexpected HTTP {status} — key may be valid but quota or plan issue possible.'
                                if status else 'Could not reach LinkedIn API.'})

    # ── LLM ───────────────────────────────────────────────────────────────
    lp = (llm.get('provider') or '').strip()
    if lp == 'gemini' or not lp:
        results.append({'label': 'LLM', 'status': 'ok',
                        'detail': 'Using platform Gemini — no custom key required.'})
    elif lp == 'openai':
        key = (llm.get('OPENAI_API_KEY') or '').strip()
        if not key:
            results.append({'label': 'OpenAI', 'status': 'error', 'detail': 'OPENAI_API_KEY is required.'})
        else:
            status, _ = _probe_get('https://api.openai.com/v1/models',
                                   headers={'Authorization': f'Bearer {key}'})
            if status == 200:
                results.append({'label': 'OpenAI', 'status': 'ok', 'detail': 'API key is valid.'})
            elif status == 401:
                results.append({'label': 'OpenAI', 'status': 'error',
                                'detail': 'Authentication failed. Check your OPENAI_API_KEY.'})
            else:
                results.append({'label': 'OpenAI', 'status': 'warn',
                                'detail': f'Unexpected HTTP {status}.' if status else 'Could not reach OpenAI API.'})
    elif lp == 'anthropic':
        key = (llm.get('ANTHROPIC_API_KEY') or '').strip()
        if not key:
            results.append({'label': 'Anthropic', 'status': 'error', 'detail': 'ANTHROPIC_API_KEY is required.'})
        else:
            payload = json.dumps({
                'model': 'claude-3-haiku-20240307', 'max_tokens': 1,
                'messages': [{'role': 'user', 'content': 'hi'}]
            }).encode('utf-8')
            status, _ = _probe_post('https://api.anthropic.com/v1/messages', payload, headers={
                'x-api-key': key, 'anthropic-version': '2023-06-01',
                'Content-Type': 'application/json',
            })
            if status == 401:
                results.append({'label': 'Anthropic', 'status': 'error',
                                'detail': 'Authentication failed. Check your ANTHROPIC_API_KEY.'})
            elif status:
                results.append({'label': 'Anthropic', 'status': 'ok',
                                'detail': f'API key accepted (HTTP {status}).'})
            else:
                results.append({'label': 'Anthropic', 'status': 'warn',
                                'detail': 'Could not reach Anthropic API.'})

    # ── Email Verification ────────────────────────────────────────────────
    ep = (email_verif.get('provider') or '').strip()
    if ep == 'default' or not ep:
        results.append({'label': 'Email Verification', 'status': 'ok',
                        'detail': 'Using platform default verification — no custom key required.'})
    elif ep == 'neverbounce':
        key = (email_verif.get('NEVERBOUNCE_API_KEY') or '').strip()
        if not key:
            results.append({'label': 'NeverBounce', 'status': 'error',
                            'detail': 'NEVERBOUNCE_API_KEY is required.'})
        else:
            status, _ = _probe_get(
                f'https://api.neverbounce.com/v4/account/info?key={_uparse2.quote(key, safe="")}')
            if status == 200:
                results.append({'label': 'NeverBounce', 'status': 'ok', 'detail': 'API key is valid.'})
            elif status in (401, 403):
                results.append({'label': 'NeverBounce', 'status': 'error',
                                'detail': f'Authentication failed (HTTP {status}).'})
            else:
                results.append({'label': 'NeverBounce', 'status': 'warn',
                                'detail': f'Unexpected HTTP {status}.' if status else 'Could not reach NeverBounce API.'})
    elif ep == 'zerobounce':
        key = (email_verif.get('ZEROBOUNCE_API_KEY') or '').strip()
        if not key:
            results.append({'label': 'ZeroBounce', 'status': 'error',
                            'detail': 'ZEROBOUNCE_API_KEY is required.'})
        else:
            status, body_text = _probe_get(
                f'https://api.zerobounce.net/v2/getcredits?api_key={_uparse2.quote(key, safe="")}')
            if status == 200:
                try:
                    credits = json.loads(body_text).get('Credits')
                except Exception:
                    credits = None
                try:
                    credits_num = int(credits) if credits is not None else None
                except (TypeError, ValueError):
                    credits_num = None
                if credits_num is not None and credits_num > 0:
                    results.append({'label': 'ZeroBounce', 'status': 'ok',
                                    'detail': f'API key valid. Credits remaining: {credits_num}.'})
                elif credits_num == 0:
                    results.append({'label': 'ZeroBounce', 'status': 'warn',
                                    'detail': 'API key valid but account has 0 credits.'})
                else:
                    results.append({'label': 'ZeroBounce', 'status': 'ok', 'detail': 'API key accepted.'})
            elif status in (400, 401):
                results.append({'label': 'ZeroBounce', 'status': 'error',
                                'detail': f'Authentication failed (HTTP {status}).'})
            else:
                results.append({'label': 'ZeroBounce', 'status': 'warn',
                                'detail': f'Unexpected HTTP {status}.' if status else 'Could not reach ZeroBounce API.'})
    elif ep == 'bouncer':
        key = (email_verif.get('BOUNCER_API_KEY') or '').strip()
        if not key:
            results.append({'label': 'Bouncer', 'status': 'error',
                            'detail': 'BOUNCER_API_KEY is required.'})
        else:
            # Use the verify endpoint for validation — the /account endpoint returns
            # 403 on some Bouncer plans even with a valid key.  A test verification
            # call returns 401 only when the key itself is invalid; any other code
            # (200 = success, 402 = no credits, 429 = rate-limited) means the key works.
            # timeout=5 in the query string is the Bouncer per-request MX lookup limit.
            status, _ = _probe_get(
                'https://api.usebouncer.com/v1.1/email/verify?email=test%40usebouncer.com&timeout=5',
                headers={'x-api-key': key})
            if status == 401:
                results.append({'label': 'Bouncer', 'status': 'error',
                                'detail': 'Authentication failed — invalid API key (HTTP 401).'})
            elif status == 402:
                results.append({'label': 'Bouncer', 'status': 'warn',
                                'detail': 'API key is valid but the account has no credits (HTTP 402).'})
            elif status == 429:
                results.append({'label': 'Bouncer', 'status': 'warn',
                                'detail': 'API key is valid but the request was rate-limited (HTTP 429).'})
            elif status == 200:
                results.append({'label': 'Bouncer', 'status': 'ok', 'detail': 'API key is valid.'})
            elif status is not None:
                # Any non-401 HTTP response means the key was accepted by Bouncer.
                results.append({'label': 'Bouncer', 'status': 'ok',
                                'detail': f'API key accepted (HTTP {status}).'})
            else:
                results.append({'label': 'Bouncer', 'status': 'warn',
                                'detail': 'Could not reach Bouncer API — please try again.'})

    # ── Contact Generation ────────────────────────────────────────────────
    cp = (contact_gen.get('provider') or '').strip()
    if cp == 'gemini' or not cp:
        results.append({'label': 'Contact Generation', 'status': 'ok',
                        'detail': 'Using platform Gemini — no custom key required.'})
    elif cp == 'contactout':
        key = (contact_gen.get('CONTACTOUT_API_KEY') or '').strip()
        if not key:
            results.append({'label': 'ContactOut', 'status': 'error',
                            'detail': 'CONTACTOUT_API_KEY is required.'})
        else:
            status, _ = _probe_get(
                'https://api.contactout.com/v1/people/linkedin?profile=https://www.linkedin.com/in/test&email_type=none&include_phone=false',
                headers={'Content-Type': 'application/json', 'Accept': 'application/json', 'token': key})
            if status == 401:
                results.append({'label': 'ContactOut', 'status': 'error',
                                'detail': 'Authentication failed (HTTP 401). Check your CONTACTOUT_API_KEY.'})
            elif status == 403:
                # 403 from ContactOut means account suspended or quota exceeded, NOT invalid key
                results.append({'label': 'ContactOut', 'status': 'warn',
                                'detail': 'ContactOut returned HTTP 403 — key may be valid but your account may be suspended or quota exceeded.'})
            elif status in (200, 404, 422):
                results.append({'label': 'ContactOut', 'status': 'ok', 'detail': 'API key accepted.'})
            else:
                results.append({'label': 'ContactOut', 'status': 'warn',
                                'detail': f'Unexpected HTTP {status} — key may be valid but check your account or try again.'
                                if status else 'Could not reach ContactOut API.'})
    elif cp == 'apollo':
        key = (contact_gen.get('APOLLO_API_KEY') or '').strip()
        if not key:
            results.append({'label': 'Apollo', 'status': 'error',
                            'detail': 'APOLLO_API_KEY is required.'})
        else:
            status, _ = _probe_get('https://api.apollo.io/v1/auth/health',
                                   headers={'x-api-key': key, 'Content-Type': 'application/json'})
            if status == 200:
                results.append({'label': 'Apollo', 'status': 'ok', 'detail': 'API key is valid.'})
            elif status == 401:
                results.append({'label': 'Apollo', 'status': 'error',
                                'detail': 'Authentication failed (HTTP 401). Check your APOLLO_API_KEY.'})
            elif status == 403:
                results.append({'label': 'Apollo', 'status': 'warn',
                                'detail': 'Apollo returned HTTP 403 — key may be valid but your account may lack access. Check your plan.'})
            elif status is not None and status >= 500:
                results.append({'label': 'Apollo', 'status': 'warn',
                                'detail': f'Apollo API returned HTTP {status} — server may be temporarily unavailable. Try again.'})
            else:
                results.append({'label': 'Apollo', 'status': 'warn',
                                'detail': f'Unexpected HTTP {status} — key may be valid but check your plan.'
                                if status else 'Could not reach Apollo API.'})
    elif cp == 'rocketreach':
        key = (contact_gen.get('ROCKETREACH_API_KEY') or '').strip()
        if not key:
            results.append({'label': 'RocketReach', 'status': 'error',
                            'detail': 'ROCKETREACH_API_KEY is required.'})
        else:
            status, _ = _probe_get('https://api.rocketreach.co/api/v2/checkStatus',
                                   headers={'Api-Key': key})
            if status == 200:
                results.append({'label': 'RocketReach', 'status': 'ok', 'detail': 'API key is valid.'})
            elif status == 401:
                results.append({'label': 'RocketReach', 'status': 'error',
                                'detail': 'Authentication failed (HTTP 401). Check your ROCKETREACH_API_KEY.'})
            elif status == 403:
                results.append({'label': 'RocketReach', 'status': 'warn',
                                'detail': 'RocketReach returned HTTP 403 — key may be valid but your account may be suspended or quota exceeded.'})
            elif status is not None and status >= 500:
                results.append({'label': 'RocketReach', 'status': 'warn',
                                'detail': f'RocketReach API returned HTTP {status} — server may be temporarily unavailable. Try again.'})
            else:
                results.append({'label': 'RocketReach', 'status': 'warn',
                                'detail': f'Unexpected HTTP {status} — key may be valid but check your plan.'
                                if status else 'Could not reach RocketReach API.'})

    return results



@app.get("/api/contactout/download-profile")
@_require_session
def contactout_download_profile():
    """Fetch the full ContactOut profile for a given LinkedIn URL.

    Query parameters:
      linkedin_url – the LinkedIn profile URL (required)

    Calls the ContactOut LinkedIn lookup endpoint
    (GET /v1/people/linkedin) with ``reveal_info=true`` so that contact
    details are included in the response, and returns the raw JSON document.
    """
    linkedin_url = (request.args.get("linkedin_url") or "").strip()
    if not linkedin_url:
        return jsonify({"error": "linkedin_url is required"}), 400

    ev_cfg = _load_email_verif_config()
    _co_admin = ev_cfg.get("contactout", {})
    co_key = (_co_admin.get("api_key") or "").strip() if _co_admin.get("enabled") == "enabled" else ""
    if not co_key:
        return jsonify({"error": "ContactOut API key is not configured or not enabled"}), 503

    try:
        r = requests.get(
            "https://api.contactout.com/v1/people/linkedin",
            params={
                "profile": linkedin_url,
                "email_type": "personal",
                "include_phone": "true",
                "reveal_info": "true",
            },
            headers={
                "token": co_key,
                "Accept": "application/json",
            },
            timeout=30,
        )
        if r.status_code == 401:
            return jsonify({"error": "ContactOut authentication failed (HTTP 401)"}), 401
        if r.status_code == 403:
            return jsonify({"error": "ContactOut returned HTTP 403 — quota may be exceeded"}), 403
        r.raise_for_status()
        profile_data = r.json()
        return jsonify(profile_data)
    except requests.exceptions.HTTPError as http_err:
        logger.warning(f"[ContactOut] download-profile HTTP error: {http_err}")
        return jsonify({"error": "ContactOut API request failed"}), 502
    except Exception as exc:
        logger.warning(f"[ContactOut] download-profile error: {exc}")
        return jsonify({"error": "Failed to fetch ContactOut profile"}), 500


@app.get("/api/apollo/download-profile")
@_require_session
def apollo_download_profile():
    """Fetch the Apollo contact profile for a given person ID or LinkedIn URL.

    Query parameters (at least one required):
      person_id    – the Apollo contact/person ID (preferred)
      linkedin_url – the LinkedIn profile URL (fallback lookup)

    Calls the Apollo mixed_people/search endpoint.  The response prominently
    exposes ``email``, ``mobile_phone``, and ``office_phone``; the complete
    contact record is included under ``_details`` for reference.
    """
    person_id = (request.args.get("person_id") or "").strip()
    linkedin_url = (request.args.get("linkedin_url") or "").strip()
    if not person_id and not linkedin_url:
        return jsonify({"error": "person_id or linkedin_url is required"}), 400

    # Resolve the Apollo API key: check per-user service config first, then admin config.
    # request._session_user is set by @_require_session after DB-validated session check.
    ap_key = ""
    _req_user = getattr(request, "_session_user", None) or (request.cookies.get("username") or "").strip()
    if _req_user:
        # Sanitise username to a safe filesystem identifier (strips path-traversal chars)
        _safe_user = _porting_safe_name(_req_user)
        try:
            _u_enc = _svc_config_path(_safe_user)
            _u_json = _svc_config_json_path(_safe_user)
            _u_cfg = None
            if os.path.isfile(_u_enc):
                try:
                    with open(_u_enc, "rb") as _fh:
                        _u_cfg = json.loads(_svc_config_decrypt(_fh.read()).decode("utf-8"))
                except Exception:
                    pass
            if _u_cfg is None and os.path.isfile(_u_json):
                try:
                    with open(_u_json, "r", encoding="utf-8") as _fh:
                        _u_cfg = json.load(_fh)
                except Exception:
                    pass
            if _u_cfg:
                ap_key = (_u_cfg.get("contact_gen", {}).get("APOLLO_API_KEY") or "").strip()
        except Exception:
            pass
    if not ap_key:
        ev_cfg = _load_email_verif_config()
        _ap_admin = ev_cfg.get("apollo", {})
        ap_key = (_ap_admin.get("api_key") or "").strip() if _ap_admin.get("enabled") == "enabled" else ""
    if not ap_key:
        return jsonify({"error": "Apollo API key is not configured or not enabled"}), 503

    def _extract_contact_fields(contact):
        """Extract email, mobile_phone, office_phone from an Apollo person/contact dict."""
        email = (contact.get("email") or "").strip()
        mobile_phone = ""
        office_phone = ""
        _mobile_types = {"mobile", "cell", "home", "personal"}
        _office_types = {"work", "office", "direct", "direct_phone", "work_hq"}
        for ph in (contact.get("phone_numbers") or []):
            ph_type = (ph.get("type") or "").lower().replace(" ", "_")
            ph_num = (ph.get("sanitized_number") or ph.get("raw_number") or "").strip()
            if not ph_num:
                continue
            if ph_type in _mobile_types and not mobile_phone:
                mobile_phone = ph_num
            elif ph_type in _office_types and not office_phone:
                office_phone = ph_num
        if not office_phone:
            _acct = contact.get("account") or {}
            office_phone = (
                _acct.get("sanitized_phone") or _acct.get("phone") or ""
            ).strip()
        return email, mobile_phone, office_phone

    def _apollo_people_match_fallback(ap_key, linkedin_url):
        """Fallback enrichment via POST /v1/people/match. Returns person dict or None."""
        logger.debug(f"[Apollo] people/match fallback for linkedin_url={linkedin_url!r}")
        try:
            r2 = requests.post(
                "https://api.apollo.io/v1/people/match",
                headers={
                    "Cache-Control": "no-cache",
                    "Content-Type": "application/json",
                    "Accept": "application/json",
                    "x-api-key": ap_key,
                },
                json={
                    "linkedin_url": linkedin_url,
                    "reveal_personal_emails": True,
                    "reveal_phone_number": True,
                },
                timeout=30,
            )
            if r2.status_code == 401:
                logger.warning("[Apollo] people/match returned HTTP 401")
                return None
            r2.raise_for_status()
            return r2.json().get("person")
        except Exception as exc2:
            logger.warning(f"[Apollo] people/match fallback error: {exc2}")
            return None

    try:
        if person_id:
            # Search global database by person ID array
            payload = {"ids": [person_id], "per_page": 1, "page": 1}
            logger.debug(f"[Apollo] mixed_people/search by person_id={person_id!r}")
        else:
            # Search global database using the LinkedIn URL
            payload = {"person_linkedin_urls": [linkedin_url], "per_page": 1, "page": 1}
            logger.debug(f"[Apollo] mixed_people/search by linkedin_url={linkedin_url!r}")

        contact = None
        used_fallback = False

        try:
            r = requests.post(
                "https://api.apollo.io/api/v1/mixed_people/search",
                headers={
                    "Cache-Control": "no-cache",
                    "Content-Type": "application/json",
                    "Accept": "application/json",
                    "x-api-key": ap_key,
                },
                json=payload,
                timeout=30,
            )
            if r.status_code == 401:
                return jsonify({"error": "Apollo authentication failed (HTTP 401)"}), 401

            # Check for free-plan restriction message before raising for status
            _mixed_body = {}
            try:
                _mixed_body = r.json()
            except Exception:
                pass
            _mixed_msg = (_mixed_body.get("message") or _mixed_body.get("error") or "").lower()
            _is_plan_restricted = "not accessible" in _mixed_msg and "free plan" in _mixed_msg

            if r.status_code == 200 and not _is_plan_restricted:
                r.raise_for_status()
                people = _mixed_body.get("people") or []
                if people:
                    contact = people[0]
                    logger.debug(f"[Apollo] mixed_people/search found person id={contact.get('id')!r}")
                else:
                    logger.info("[Apollo] mixed_people/search returned empty people — trying people/match fallback")
            else:
                logger.info(
                    f"[Apollo] mixed_people/search failed (status={r.status_code}, "
                    f"plan_restricted={_is_plan_restricted}) — trying people/match fallback"
                )
        except requests.exceptions.HTTPError as primary_err:
            logger.info(f"[Apollo] mixed_people/search HTTP error ({primary_err}) — trying people/match fallback")
        except Exception as primary_exc:
            logger.info(f"[Apollo] mixed_people/search error ({primary_exc}) — trying people/match fallback")

        # Fallback to people/match when primary search did not return a contact
        if contact is None and linkedin_url:
            contact = _apollo_people_match_fallback(ap_key, linkedin_url)
            if contact:
                used_fallback = True

        if not contact:
            return jsonify({"error": "No matching contact found in Apollo"}), 404

        # --- Extract contact fields ---
        email, mobile_phone, office_phone = _extract_contact_fields(contact)

        # --- Build the response ---
        result = {
            "email": email,
            "mobile_phone": mobile_phone,
            "office_phone": office_phone,
            "_details": contact,
        }
        logger.info(
            f"[Apollo] {'people/match' if used_fallback else 'mixed_people/search'} "
            f"returned person id={contact.get('id')!r} "
            f"email={'***' if email else '(none)'} "
            f"mobile={'***' if mobile_phone else '(none)'} "
            f"office={'***' if office_phone else '(none)'}"
        )
        return jsonify(result)
    except Exception as exc:
        logger.warning(f"[Apollo] download-profile error: {exc}")
        return jsonify({"error": "Failed to fetch Apollo contact"}), 500


@app.get("/api/rocketreach/download-profile")
@_require_session
def rocketreach_download_profile():
    """Fetch the full RocketReach profile for a given LinkedIn URL.

    Query parameters:
      linkedin_url – the LinkedIn profile URL (required)

    Calls the RocketReach lookupProfile endpoint to retrieve the full profile JSON.
    """
    linkedin_url = (request.args.get("linkedin_url") or "").strip()
    if not linkedin_url:
        return jsonify({"error": "linkedin_url is required"}), 400

    # Resolve the RocketReach API key: check per-user service config first, then admin config.
    rr_key = ""
    _req_user = getattr(request, "_session_user", None) or (request.cookies.get("username") or "").strip()
    if _req_user:
        _safe_user = _porting_safe_name(_req_user)
        try:
            _u_enc = _svc_config_path(_safe_user)
            _u_json = _svc_config_json_path(_safe_user)
            _u_cfg = None
            if os.path.isfile(_u_enc):
                try:
                    with open(_u_enc, "rb") as _fh:
                        _u_cfg = json.loads(_svc_config_decrypt(_fh.read()).decode("utf-8"))
                except Exception:
                    pass
            if _u_cfg is None and os.path.isfile(_u_json):
                try:
                    with open(_u_json, "r", encoding="utf-8") as _fh:
                        _u_cfg = json.load(_fh)
                except Exception:
                    pass
            if _u_cfg:
                rr_key = (_u_cfg.get("contact_gen", {}).get("ROCKETREACH_API_KEY") or "").strip()
        except Exception:
            pass
    if not rr_key:
        ev_cfg = _load_email_verif_config()
        _rr_admin = ev_cfg.get("rocketreach", {})
        rr_key = (_rr_admin.get("api_key") or "").strip() if _rr_admin.get("enabled") == "enabled" else ""
    if not rr_key:
        return jsonify({"error": "RocketReach API key is not configured or not enabled"}), 503

    try:
        r = requests.get(
            "https://api.rocketreach.co/api/v2/lookupProfile",
            params={"linkedin_url": linkedin_url},
            headers={
                "Api-Key": rr_key,
                "Accept": "application/json",
            },
            timeout=30,
        )
        if r.status_code == 401:
            return jsonify({"error": "RocketReach authentication failed (HTTP 401)"}), 401
        if r.status_code == 403:
            return jsonify({"error": "RocketReach returned HTTP 403 — quota may be exceeded"}), 403
        r.raise_for_status()
        profile_data = r.json()
        return jsonify(profile_data)
    except requests.exceptions.HTTPError as http_err:
        logger.warning(f"[RocketReach] download-profile HTTP error: {http_err}")
        return jsonify({"error": "RocketReach API request failed"}), 502
    except Exception as exc:
        logger.warning(f"[RocketReach] download-profile error: {exc}")
        return jsonify({"error": "Failed to fetch RocketReach profile"}), 500


# ---------------------------------------------------------------------------
# linkdapi: full LinkedIn profile retrieval
# ---------------------------------------------------------------------------

_LINKDAPI_HOST = "linkdapi.com"
_LINKDAPI_HEADER = "X-linkdapi-apikey"
_LINKDAPI_MAX_RETRIES = 3

# ── Scrapingdog API constants ─────────────────────────────────────────────────
_SCRAPINGDOG_API_BASE = "api.scrapingdog.com"
_SCRAPINGDOG_MAX_RETRIES = 3


# ── BrightData API constants ───────────────────────────────────────────────────
_BRIGHTDATA_API_BASE = "api.brightdata.com"
_BRIGHTDATA_REQUEST_URL = f"https://{_BRIGHTDATA_API_BASE}/request"


def _brightdata_fetch_profile(linkedin_url: str, api_key: str, zone: str, timeout: int = 60):
    """Fetch a LinkedIn profile via BrightData SERP API.

    Calls POST https://api.brightdata.com/request with zone, url, and format
    parameters as required by the BrightData SERP API specification.

    Parameters
    ----------
    linkedin_url : str
        Full canonical LinkedIn profile URL (e.g. https://www.linkedin.com/in/username).
    api_key : str
        BrightData Bearer API key (from account settings).
    zone : str
        BrightData zone identifier (managed at https://brightdata.com/cp/zones).
    timeout : int
        Request timeout in seconds.  60 s is used because the SERP API is a
        single synchronous call — no polling loop is involved.

    Returns ``(body_str, http_status)`` where:
      - success: body_str is a JSON string (list or dict), status 200
      - auth error: status 401 / 403
      - other error: status >= 400 or 0 for network failure
    """
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "zone": zone,
        "url": linkedin_url,
        "format": "json",
    }
    try:
        resp = requests.post(
            _BRIGHTDATA_REQUEST_URL,
            json=payload,
            headers=headers,
            timeout=timeout,
        )
        logger.info("[brightdata] POST /request → HTTP %d (%d bytes)", resp.status_code, len(resp.text))
        if resp.status_code >= 400:
            logger.warning("[brightdata] HTTP %d; body: %.200s", resp.status_code, resp.text[:200])
        return resp.text, resp.status_code
    except requests.exceptions.Timeout:
        logger.warning("[brightdata] request timed out")
        return "", 0
    except Exception as exc:
        logger.error("[brightdata] request failed: %s", exc)
        return "", 0


def _scrapingdog_fetch_profile(linkedin_id: str, api_key: str, timeout: int = 60):
    """Fetch a LinkedIn profile from Scrapingdog and return ``(body_str, http_status)``.

    Calls GET https://api.scrapingdog.com/profile?api_key=...&id=...&type=profile&premium=true

    ``linkedin_id`` should be the full canonical LinkedIn URL
    (e.g. ``https://www.linkedin.com/in/username``).  Passing only the
    username slug may cause Scrapingdog to return HTTP 400 for some regional
    profiles (cn., jp., …).
    """
    url = f"https://{_SCRAPINGDOG_API_BASE}/profile"
    params = {
        "api_key": api_key,
        "id": linkedin_id,
        "type": "profile",
        "premium": "true",
        "webhook": "false",
        "fresh": "false",
    }

    last_exc = None
    for attempt in range(1, _SCRAPINGDOG_MAX_RETRIES + 1):
        try:
            resp = requests.get(url, params=params, timeout=timeout,
                                headers={"Accept": "application/json"})
            body = resp.text
            logger.info("[scrapingdog] attempt %d/%d → HTTP %d (%d bytes)",
                        attempt, _SCRAPINGDOG_MAX_RETRIES, resp.status_code,
                        len(body))
            return body, resp.status_code
        except requests.exceptions.Timeout as exc:
            last_exc = exc
            logger.warning("[scrapingdog] fetch attempt %d/%d timed out: %s",
                           attempt, _SCRAPINGDOG_MAX_RETRIES, exc)
        except Exception as exc:
            last_exc = exc
            logger.warning("[scrapingdog] fetch attempt %d/%d failed: %s",
                           attempt, _SCRAPINGDOG_MAX_RETRIES, exc)

    logger.error("[scrapingdog] all %d attempts failed; last error: %s",
                 _SCRAPINGDOG_MAX_RETRIES, last_exc)
    return "", 0


def _linkdapi_fetch(username: str, api_key: str, timeout: int = 30):
    """Fetch a profile from linkdapi.com and return ``(body_str, http_status)``.

    **Strategy** (dual-layer, cross-platform, with retry):

    1. *Primary* – Python ``http.client.HTTPSConnection`` with a custom
       ``connect()`` that creates a TLS 1.2-only context, suppresses SNI
       (``server_hostname=None``), and enables ``OP_LEGACY_SERVER_CONNECT``.
       This uses Python's bundled OpenSSL, *not* the OS TLS stack (so it
       works identically on Windows, Linux, and macOS).

    2. *Fallback* – ``curl`` subprocess with ``-sSk --tlsv1.2 --tls-max 1.2``
       (pins exactly TLS 1.2) plus ``--ssl-no-revoke`` for Windows Schannel.

    Both layers are retried up to ``_LINKDAPI_MAX_RETRIES`` times for
    transient TLS / network errors before giving up.

    The linkdapi.com server may send an ``unrecognized_name`` TLS warning
    alert.  TLS 1.3 (RFC 8446 §6) treats *all* warning alerts as fatal,
    killing the handshake.  TLS 1.2 (RFC 6066 §3) considers it non-fatal.
    Forcing TLS 1.2 + suppressing SNI avoids the alert entirely.
    """
    import http.client as _hc   # noqa: PLC0415
    import ssl as _ssl          # noqa: PLC0415
    import urllib.parse as _up  # noqa: PLC0415

    qs = _up.urlencode({"username": username})
    path = f"/api/v1/profile/full?{qs}"
    headers = {_LINKDAPI_HEADER: api_key, "Accept": "application/json"}

    # Sanitise api_key for safe shell-arg use (reject embedded control chars).
    _safe_key = re.sub(r"[^\x20-\x7E]", "", api_key)

    last_exc = None

    for attempt in range(1, _LINKDAPI_MAX_RETRIES + 1):
        # ---- primary: Python http.client (uses Python/OpenSSL, not OS TLS) ----
        conn = None
        try:
            ctx = _ssl.SSLContext(_ssl.PROTOCOL_TLS_CLIENT)
            ctx.check_hostname = False
            ctx.verify_mode = _ssl.CERT_NONE
            ctx.maximum_version = _ssl.TLSVersion.TLSv1_2
            ctx.minimum_version = _ssl.TLSVersion.TLSv1_2
            if hasattr(_ssl, "OP_LEGACY_SERVER_CONNECT"):
                ctx.options |= _ssl.OP_LEGACY_SERVER_CONNECT
            if hasattr(_ssl, "OP_IGNORE_UNEXPECTED_EOF"):
                ctx.options |= _ssl.OP_IGNORE_UNEXPECTED_EOF

            import socket as _sock  # noqa: PLC0415

            conn = _hc.HTTPSConnection(_LINKDAPI_HOST, timeout=timeout, context=ctx)
            # Override connect() to suppress SNI (server_hostname=None).

            def _no_sni_connect():
                conn.sock = _sock.create_connection(
                    (conn.host, conn.port or 443), conn.timeout
                )
                conn.sock = ctx.wrap_socket(conn.sock, server_hostname=None)

            conn.connect = _no_sni_connect
            conn.request("GET", path, headers=headers)
            resp = conn.getresponse()
            body = resp.read().decode("utf-8", errors="replace")
            status = resp.status
            return body, status
        except Exception as py_exc:
            last_exc = py_exc
            logger.debug(
                "[linkdapi] Python http.client attempt %d/%d failed (%s), trying curl…",
                attempt, _LINKDAPI_MAX_RETRIES, py_exc,
            )
        finally:
            if conn is not None:
                try:
                    conn.close()
                except Exception:
                    pass

        # ---- fallback: curl subprocess ----
        try:
            import subprocess as _sp  # noqa: PLC0415

            url = f"https://{_LINKDAPI_HOST}{path}"
            sep = "\n__LINKDAPI_HTTP_STATUS__"
            cmd = [
                "curl", "-sSk",
                "--tlsv1.2", "--tls-max", "1.2",
                "--ssl-no-revoke",
                "--max-time", str(timeout),
                "-H", f"{_LINKDAPI_HEADER}: {_safe_key}",
                "-H", "Accept: application/json",
                "-w", sep + "%{http_code}",
                url,
            ]
            result = _sp.run(cmd, capture_output=True, timeout=timeout + 5)
            raw = result.stdout.decode("utf-8", errors="replace")
            if sep in raw:
                body_str, st = raw.rsplit(sep, 1)
            else:
                body_str, st = raw, "0"
            try:
                sc = int(st.strip())
            except ValueError:
                sc = 0

            # Non-transient HTTP errors — do not retry
            if sc in (401, 403, 404):
                return body_str, sc

            if result.returncode != 0 and sc == 0:
                stderr = result.stderr.decode("utf-8", errors="replace").strip()
                logger.warning(
                    "[linkdapi] curl error attempt %d/%d (rc=%d): %s",
                    attempt, _LINKDAPI_MAX_RETRIES, result.returncode, stderr,
                )
                last_exc = RuntimeError(stderr)
                continue  # retry transient TLS/network errors
            return body_str, sc
        except _sp.TimeoutExpired:
            logger.warning("[linkdapi] get-profile timed out (curl, attempt %d/%d)", attempt, _LINKDAPI_MAX_RETRIES)
            last_exc = TimeoutError("curl timed out")
        except FileNotFoundError:
            logger.error("[linkdapi] curl binary not found")
            return "curl not found on server", 0
        except Exception as curl_exc:
            logger.warning("[linkdapi] curl fallback error (attempt %d/%d): %s", attempt, _LINKDAPI_MAX_RETRIES, curl_exc)
            last_exc = curl_exc

    # All retries exhausted
    logger.error("[linkdapi] all %d attempts failed; last error: %s", _LINKDAPI_MAX_RETRIES, last_exc)
    return "TLS/connection error reaching linkdapi", 0


LINKDAPI_PROFILE_OUTPUT_DIR = os.getenv(
    "LINKDAPI_PROFILE_OUTPUT_DIR",
    r"F:\Recruiting Tools\Autosourcing\output\profiles",
)


@app.get("/api/linkdapi/get-profile")
@_require_session
@_check_gp_rate_limit()
def linkdapi_get_profile():
    """Fetch the full LinkedIn profile via linkdapi for a given LinkedIn URL.

    Query parameters:
      linkedin_url – the LinkedIn profile URL (required); username is extracted
                     from the URL path (e.g. /in/ryanroslansky → ryanroslansky)

    Calls GET https://linkdapi.com/api/v1/profile/full?username=<username>
    using the admin-configured LINKDAPI_API_KEY, then saves the JSON to
    LINKDAPI_PROFILE_OUTPUT_DIR with the authenticated username appended to
    the filename.
    """
    linkedin_url = (request.args.get("linkedin_url") or "").strip()
    if not linkedin_url:
        return jsonify({"error": "linkedin_url is required"}), 400

    # Extract LinkedIn username from URL (re already imported at module level)
    _m = re.search(r"/in/([A-Za-z0-9_%-]+)", linkedin_url)
    if not _m:
        return jsonify({"error": "Could not extract LinkedIn username from URL"}), 400
    username = _m.group(1)

    # Per-user GP key (from VIP admin assignment or user's own api_porting config) takes
    # priority over the global platform key.  User-defined keys always win because both
    # admin VIP writes and user self-service writes target the same per-user config file,
    # and the user's own save (from api_porting.html) overwrites any admin-set value.
    _req_user = getattr(request, '_session_user', None) or (request.cookies.get('username') or '').strip()
    _user_gp = _load_user_gp_cfg(_req_user)
    if _user_gp.get('provider') == 'linkdapi':
        api_key = (_user_gp.get('GP_LINKDAPI_API_KEY') or '').strip()
        if not api_key:
            return jsonify({"error": "Linkdapi API key is not configured for your account"}), 503
    else:
        gp_cfg = _load_get_profiles_config()
        linkdapi = gp_cfg.get("linkdapi", {})
        if linkdapi.get("enabled") != "enabled":
            return jsonify({"error": "linkdapi is not enabled"}), 503
        api_key = (linkdapi.get("api_key") or "").strip()
        if not api_key:
            return jsonify({"error": "LINKDAPI_API_KEY is not configured"}), 503

    body_str, status_code = _linkdapi_fetch(username, api_key)

    if status_code == 401:
        return (
            jsonify({"error": "linkdapi authentication failed (HTTP 401). Check your API key."}),
            401,
        )
    if status_code == 403:
        return (
            jsonify({"error": "linkdapi returned HTTP 403 — quota may be exceeded or key restricted"}),
            403,
        )
    if status_code == 404:
        return (
            jsonify({"error": f"Profile not found for username '{username}' (HTTP 404)"}),
            404,
        )
    if status_code >= 400:
        logger.warning("[linkdapi] upstream HTTP %d", status_code)
        return (
            jsonify({"error": f"linkdapi returned an error (HTTP {status_code})"}),
            status_code,
        )
    if status_code == 0:
        return jsonify({"error": "Failed to reach linkdapi service"}), 502

    try:
        profile_data = json.loads(body_str)
    except (json.JSONDecodeError, ValueError):
        return jsonify({"error": "Invalid JSON from linkdapi"}), 502
    def _safe_slug(value: str, fallback: str) -> str:
        slug = re.sub(r'[^A-Za-z0-9_-]+', '_', value or "")
        slug = re.sub(r'_+', '_', slug).strip('_')
        return slug or fallback

    active_username = getattr(request, "_session_user", "") or ""
    safe_active_username = _safe_slug(active_username, "unknown")
    safe_profile_username = _safe_slug(username, "linkdapi_profile")
    out_filename = f"{safe_profile_username}_{safe_active_username}.json"
    out_dir = os.path.abspath(LINKDAPI_PROFILE_OUTPUT_DIR)
    out_path = os.path.abspath(os.path.join(out_dir, out_filename))
    try:
        real_out_dir = os.path.realpath(out_dir)
        real_out_path = os.path.realpath(out_path)
        if os.path.commonpath([real_out_dir, real_out_path]) != real_out_dir:
            return jsonify({"error": "Invalid output path"}), 400
    except ValueError:
        return jsonify({"error": "Invalid output path"}), 400

    try:
        os.makedirs(out_dir, exist_ok=True)
        with open(out_path, "w", encoding="utf-8") as fh:
            json.dump(profile_data, fh, ensure_ascii=False, indent=2)
    except Exception as exc:
        logger.exception("[linkdapi] failed to save profile JSON: %s", exc)
        return jsonify({"error": "Failed to save profile JSON output file"}), 500

    return jsonify({
        "profile": profile_data,
        "saved_filename": out_filename,
        "saved_for_user": safe_active_username,
    })


@app.get("/api/linkdapi/read-profile")
@_require_session
def linkdapi_read_profile():
    """Read a previously-saved GP profile JSON from the output directory.

    Query parameters:
      linkedin_url – the LinkedIn profile URL (required); username is extracted
                     from the URL path (e.g. /in/ryanroslansky → ryanroslansky)

    The file is expected at
    ``LINKDAPI_PROFILE_OUTPUT_DIR/<linkedin_slug>_<session_username>.json``.
    Returns the parsed profile JSON so the UI can use it for assessment.
    """
    linkedin_url = (request.args.get("linkedin_url") or "").strip()
    if not linkedin_url:
        return jsonify({"error": "linkedin_url is required"}), 400

    _m = re.search(r"/in/([A-Za-z0-9_%-]+)", linkedin_url)
    if not _m:
        return jsonify({"error": "Could not extract LinkedIn username from URL"}), 400
    username = _m.group(1)

    def _safe_slug(value: str, fallback: str) -> str:
        slug = re.sub(r'[^A-Za-z0-9_-]+', '_', value or "")
        slug = re.sub(r'_+', '_', slug).strip('_')
        return slug or fallback

    active_username = getattr(request, "_session_user", "") or ""
    safe_active_username = _safe_slug(active_username, "unknown")
    safe_profile_username = _safe_slug(username, "linkdapi_profile")
    out_filename = f"{safe_profile_username}_{safe_active_username}.json"

    # Filename whitelist: after _safe_slug only [A-Za-z0-9_-] remain,
    # so the composed name is always safe (no path separators, no dots
    # except the ".json" suffix).
    if os.sep in out_filename or (os.altsep and os.altsep in out_filename):
        return jsonify({"error": "Invalid filename"}), 400

    # Build path from the fixed output directory + the sanitised filename.
    out_dir = os.path.abspath(LINKDAPI_PROFILE_OUTPUT_DIR)

    # List directory and match: avoids constructing a path from
    # user-influenced values so the file-open cannot be steered by input.
    try:
        existing = set(os.listdir(out_dir))
    except FileNotFoundError:
        return jsonify({"error": "GP profiles directory does not exist"}), 404
    except OSError as exc:
        logger.exception("[linkdapi] cannot list profiles dir: %s", exc)
        return jsonify({"error": "Cannot read profiles directory"}), 500

    if out_filename not in existing:
        return jsonify({"error": "GP profile not found. Click the GP button first to fetch the profile."}), 404

    # Use the entry from the directory listing (OS-sourced, not user-derived)
    # to construct the path — this breaks the taint chain from user input.
    matched_entry = next(f for f in existing if f == out_filename)
    safe_path = os.path.join(out_dir, matched_entry)

    try:
        with open(safe_path, "r", encoding="utf-8") as fh:
            profile_data = json.load(fh)
    except (json.JSONDecodeError, ValueError):
        return jsonify({"error": "Saved GP profile JSON is corrupted"}), 500
    except Exception as exc:
        logger.exception("[linkdapi] failed to read profile JSON: %s", exc)
        return jsonify({"error": "Failed to read GP profile JSON"}), 500

    return jsonify({
        "profile": profile_data,
        "filename": out_filename,
        "username": safe_active_username,
    })


def _linkdapi_json_to_pdf_bytes(profile: dict) -> bytes:
    """Convert a linkdapi profile JSON dict to a FIOE-branded A4 PDF.

    Uses the configured LLM (Gemini by default via ``unified_llm_call_text``)
    to extract and normalise profile data, then renders with the FIOE colour
    scheme (azure dragon / cool blue / robin's egg).  Falls back to direct
    field extraction when the LLM call fails or returns unparseable output.
    """

    # ── Attempt LLM-assisted structuring ──────────────────────────────────
    formatted = None
    try:
        profile_json_str = json.dumps(profile, ensure_ascii=False, indent=2)
        prompt = (
            "You are preparing a LinkedIn profile for a professional PDF document.\n"
            "Given the profile JSON below, extract and organise the information "
            "into this exact JSON structure (return ONLY valid JSON — no markdown "
            "code blocks, no extra text):\n\n"
            "{\n"
            '  "name": "Full Name",\n'
            '  "headline": "Job Title at Company",\n'
            '  "location": "City, Country",\n'
            '  "email": "email@example.com",\n'
            '  "linkedin_url": "https://linkedin.com/in/...",\n'
            '  "summary": "Professional summary text (2-4 sentences)",\n'
            '  "experience": [\n'
            '    {"title": "Job Title", "company": "Company Name",\n'
            '     "dates": "Jan 2020 - Present",\n'
            '     "description": "Key responsibilities in 1-3 sentences"}\n'
            '  ],\n'
            '  "education": [\n'
            '    {"school": "University Name",\n'
            '     "degree": "Bachelor of Science in Computer Science",\n'
            '     "dates": "2016 - 2020"}\n'
            '  ],\n'
            '  "skills": ["Skill 1", "Skill 2"]\n'
            "}\n\n"
            "Rules:\n"
            "- Format experience dates as 'Mon YYYY - Mon YYYY' or 'Mon YYYY - Present'\n"
            "- Keep descriptions concise (max 3 sentences per role)\n"
            "- Transliterate or translate any non-English text (e.g. Japanese, Chinese, Korean) to English\n"
            "- Use only ASCII or Latin characters in all fields — no CJK, Arabic, Cyrillic or other scripts\n"
            "- Omit keys whose values are empty or unknown\n"
            "- Do NOT insert line breaks (\\n) inside any text value; keep every field as a single continuous string\n"
            "- Return ONLY valid JSON\n\n"
            f"Profile JSON:\n{profile_json_str}"
        )
        raw = unified_llm_call_text(prompt, temperature=0.1, max_output_tokens=8000)
        if raw:
            cleaned = raw.strip()
            if cleaned.startswith("```"):
                cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
                cleaned = re.sub(r"\s*```$", "", cleaned).strip()
            formatted = json.loads(cleaned)
    except Exception as exc:
        logger.warning("[linkdapi PDF] LLM formatting failed: %s", exc)
        formatted = None

    # ── Fallback: direct extraction from raw linkdapi JSON fields ──────────
    def _fmt_date(d):
        if not d:
            return ""
        if isinstance(d, str):
            return d
        _MONTHS = ["", "Jan", "Feb", "Mar", "Apr", "May", "Jun",
                   "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
        m = d.get("month", "")
        y = d.get("year", "")
        try:
            m_str = _MONTHS[int(m)] if m and 1 <= int(m) <= 12 else str(m)
        except (ValueError, TypeError):
            m_str = str(m) if m else ""
        return f"{m_str} {y}".strip() if m_str else str(y)

    def _fmt_year(d):
        if not d:
            return ""
        if isinstance(d, (int, str)):
            return str(d)
        return str(d.get("year", ""))

    if not (formatted and isinstance(formatted, dict)):
        first = (profile.get("firstName") or profile.get("first_name") or "").strip()
        last = (profile.get("lastName") or profile.get("last_name") or "").strip()
        positions_raw = profile.get("positions") or profile.get("experience") or []
        schools_raw = profile.get("schools") or profile.get("education") or []
        skills_raw = profile.get("skills") or []

        exp_list = []
        for pos in positions_raw:
            is_current = pos.get("isCurrent") or pos.get("is_current") or False
            start_str = _fmt_date(pos.get("startDate") or pos.get("start_date") or {})
            end_str = "Present" if is_current else _fmt_date(
                pos.get("endDate") or pos.get("end_date") or {})
            dates = f"{start_str} - {end_str}" if (start_str or end_str) else ""
            exp_list.append({
                "title":       (pos.get("title") or "").strip(),
                "company":     (pos.get("companyName") or pos.get("company")
                                or pos.get("company_name") or "").strip(),
                "dates":       dates,
                "description": (pos.get("description") or "").strip(),
            })

        edu_list = []
        for edu in schools_raw:
            degree = (edu.get("degree") or edu.get("degreeName") or "").strip()
            field = (edu.get("fieldOfStudy") or edu.get("field_of_study") or "").strip()
            if degree and field:
                degree = f"{degree} in {field}"
            elif field:
                degree = field
            sy = _fmt_year(edu.get("startDate") or edu.get("start_date") or {})
            ey = _fmt_year(edu.get("endDate") or edu.get("end_date") or {})
            dates = f"{sy} - {ey}" if (sy and ey) else (sy or ey)
            edu_list.append({
                "school": (edu.get("schoolName") or edu.get("school")
                           or edu.get("name") or "").strip(),
                "degree": degree,
                "dates":  dates,
            })

        skill_names = [
            (s.get("name") or s.get("skill") or str(s)) if isinstance(s, dict) else str(s)
            for s in skills_raw
        ]

        formatted = {
            "name":         (profile.get("fullName") or profile.get("full_name")
                             or f"{first} {last}").strip(),
            "headline":     (profile.get("headline") or "").strip(),
            "location":     (profile.get("location")
                             or (profile.get("geo") or {}).get("full") or "").strip(),
            "email":        (profile.get("email") or "").strip(),
            "linkedin_url": (profile.get("profileUrl") or profile.get("profile_url")
                             or profile.get("url") or "").strip(),
            "summary":      (profile.get("summary") or profile.get("about") or "").strip(),
            "experience":   exp_list,
            "education":    edu_list,
            "skills":       skill_names,
        }

    # ── Render with FIOE branding ──────────────────────────────────────────
    return _render_fioe_profile_pdf(formatted)


def _render_fioe_profile_pdf(data: dict) -> bytes:
    """Render a FIOE-branded A4 PDF from a structured profile dict using
    reportlab Platypus for proper text flow (no right-side cut-off).

    FIOE colour scheme:
      Azure Dragon  #073679  — primary header background, section titles
      Cool Blue     #4c82b8  — dates, sub-labels
      Robin's Egg   #6deaf9  — decorative rules, contact strip

    Falls back to ``_lines_to_pdf_bytes`` when reportlab is unavailable.
    """
    # ── Shared text sanitiser ──────────────────────────────────────────────
    def _s(t):
        """Normalise text for Latin-1 PDF rendering (canvas-safe).

        1. Strips XML-invalid control characters (which would crash Platypus).
        2. NFKC-normalises (handles fullwidth/halfwidth variants).
        3. Replaces entire *runs* of CJK / non-Latin characters with ``[...]``
           so the output clearly shows where untransliterated text exists
           rather than producing a string of ``?`` characters.
        4. Applies targeted replacements for common Unicode punctuation and
           maps ALL Unicode Pd (dash) / Pc (connector) characters to ASCII ``-``
           and non-Latin-1 Zs (space-separator) characters to ASCII space so
           date separators produced by LLMs never appear as ``?``.
        5. Encodes to Latin-1, replacing any remaining stragglers with ``?``.
        """
        import unicodedata as _ud
        import re as _re
        s = str(t or "")
        # 1. Strip control characters that are invalid in XML (U+0000-U+0008,
        #    U+000B, U+000C, U+000E-U+001F) — these crash Platypus's XML parser.
        #    Also strip DEL (U+007F) and the C1 control range (U+0080-U+009F)
        #    which includes NEL (U+0085) — some Platypus parsers misinterpret
        #    NEL as a newline and raise an XML error mid-paragraph.
        #    Normalise CR/TAB to a plain space; preserve LF (\n) so Gemini
        #    pre-wrapped line breaks survive into _spwrap() for PDF rendering.
        s = _re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f\x80-\x9f]', '', s)
        s = _re.sub(r'[\r\t]', ' ', s)  # preserve \n for Gemini line break formatting
        # 2. NFKC normalisation
        s = _ud.normalize('NFKC', s)
        # 3. Replace runs of non-Latin characters (CJK, Arabic, Cyrillic, etc.)
        s = _re.sub(
            r'[\u0400-\u04ff'           # Cyrillic
            r'\u0500-\u05ff'           # Cyrillic Supplement, Hebrew, Armenian
            r'\u0600-\u06ff'           # Arabic
            r'\u0700-\u08ff'           # Syriac, Thaana, NKo, Samaritan, Mandaic
            r'\u0900-\u0dff'           # Devanagari, Bengali, Gurmukhi, Gujarati,
                                       #   Oriya, Tamil, Telugu, Kannada, Malayalam
            r'\u0e00-\u0e7f'           # Thai
            r'\u0e80-\u0eff'           # Lao
            r'\u1000-\u10ff'           # Myanmar, Georgian
            r'\u1100-\u11ff'           # Hangul Jamo
            r'\u1200-\u137f'           # Ethiopic
            r'\u1700-\u17ff'           # Tagalog, Hanunoo, Buhid, Tagbanwa, Khmer
            r'\u1800-\u18af'           # Mongolian
            r'\u1e00-\u1eff'           # Latin Extended Additional (pre-composed)
            r'\u2e80-\u2eff'           # CJK Radicals Supplement
            r'\u2f00-\u2fdf'           # Kangxi Radicals
            r'\u3000-\u303f'           # CJK symbols / punctuation
            r'\u3040-\u30ff'           # Hiragana + Katakana
            r'\u3130-\u318f'           # Hangul Compatibility Jamo
            r'\u3190-\u31ef'           # Kanbun, Bopomofo Extended, CJK Strokes
            r'\u3200-\u33ff'           # Enclosed CJK + CJK Compatibility
            r'\u3400-\u4dbf'           # CJK Unified Ideographs Extension A
            r'\u4e00-\u9fff'           # CJK Unified Ideographs (main block)
            r'\ua000-\ua4cf'           # Yi Syllables + Yi Radicals
            r'\ua960-\ua97f'           # Hangul Jamo Extended-A
            r'\uf900-\ufaff'           # CJK Compatibility Ideographs
            r'\uac00-\ud7af'           # Hangul Syllables
            r'\ud7b0-\ud7ff'           # Hangul Jamo Extended-B
            r'\uff00-\uffef'           # Halfwidth / Fullwidth forms
            r']+'
            ,
            '', s
        )
        # 4a. Named Unicode punctuation replacements
        s = s.replace('\u2013', '-').replace('\u2014', '-').replace('\u2015', '-')
        s = s.replace('\u2018', "'").replace('\u2019', "'")
        s = s.replace('\u201c', '"').replace('\u201d', '"')
        s = s.replace('\u2026', '...')
        s = s.replace('\u2022', '-')
        s = s.replace('\u2212', '-')
        s = s.replace('\u30fb', '\u00b7')
        s = s.replace('\u301c', '~').replace('\uff5e', '~')
        # 4b. Map ALL remaining non-Latin-1 Unicode dashes/connectors (Pd/Pc
        #     category) to ASCII hyphen and non-Latin-1 space separators (Zs)
        #     to ASCII space.  This catches U+2010 HYPHEN, U+2011 NB-HYPHEN,
        #     U+2012 FIGURE DASH and any other LLM-generated separator that
        #     would otherwise survive as a ``?`` after Latin-1 encoding.
        #     Only characters outside Latin-1 (code-point > U+00FF) that survived
        #     the earlier CJK/named-replacement passes are visited here, so the
        #     ord() guard means _ud.category() is never called for Latin-1 chars.
        def _map_non_latin1(m):
            c = m.group(0)
            cat = _ud.category(c)
            if cat in ('Pd', 'Pc'):
                return '-'
            if cat == 'Zs':
                return ' '
            return c  # keep; encode() will drop silently if not Latin-1

        s = _re.sub(r'[\u0100-\uffff]', _map_non_latin1, s)
        return s.encode("latin-1", errors="ignore").decode("latin-1")

    def _sp(t):
        """Return text safe for Platypus Paragraph (Latin-1 + XML-escaped).

        Platypus Paragraph parses its input as XML, so literal ``<``, ``>``
        and ``&`` characters must be escaped to prevent the XML parser from
        failing or misrendering the text.
        """
        from xml.sax.saxutils import escape as _xmlesc
        return _xmlesc(_s(t))

    def _spwrap(t, max_chars=100):
        """Return XML-safe text with hard line breaks every ``max_chars`` chars.

        Honors existing ``\\n`` line breaks (intentional multi-line content)
        and additionally enforces the ``max_chars`` limit on any line that
        still exceeds it.  Words are preserved wherever possible; a word that
        exceeds ``max_chars`` by itself is hard-split at the character boundary.
        Lines are joined with Platypus ``<br/>`` tags so the paragraph renderer
        respects each visual line.
        """
        from xml.sax.saxutils import escape as _xmlesc
        raw = _s(t)
        # Split on existing newlines first (intentional line breaks), then
        # enforce max_chars within each segment.
        segments = raw.split('\n')
        final_lines = []
        for segment in segments:
            segment = segment.strip()
            if not segment:
                continue
            if len(segment) <= max_chars:
                final_lines.append(segment)
                continue
            # Re-wrap this segment at word boundaries
            words = segment.split(' ')
            current = ''
            for word in words:
                if not word:
                    continue
                if len(word) > max_chars:
                    if current:
                        final_lines.append(current)
                        current = ''
                    for ch in range(0, len(word), max_chars):
                        final_lines.append(word[ch:ch + max_chars])
                    continue
                candidate = (current + ' ' + word).lstrip() if current else word
                if len(candidate) > max_chars:
                    final_lines.append(current)
                    current = word
                else:
                    current = candidate
            if current:
                final_lines.append(current)
        return '<br/>'.join(_xmlesc(line) for line in final_lines)

    # ── Try reportlab.platypus (primary path) ──────────────────────────────
    try:
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            HRFlowable,
        )
        try:
            from reportlab.platypus import ListFlowable as _ListFlowable
            from reportlab.platypus import ListItem as _ListItem
        except ImportError:
            _ListFlowable = None
            _ListItem = None
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import pt
        from reportlab.lib.colors import HexColor, white, black
        import io as _io

        AZURE    = HexColor('#073679')
        COOL     = HexColor('#4c82b8')
        ROBIN    = HexColor('#6deaf9')
        ROBIN_BG = HexColor('#edf9ff')
        DARK     = HexColor('#1a1a1a')
        GRAY     = HexColor('#5a5a5a')

        PAGE_W, PAGE_H = A4
        ML, MR, MT, MB = 45, 45, 20, 30   # margins (pt); top is extra for header

        # Header band height (drawn in page callbacks)
        HDR_H = 60

        def _hdr_name(name_text, headline_text):
            """Return rough header band height needed for name + headline.

            Estimates line count using approximate chars-per-line values:
              * 60 chars/line for name (20pt bold)
              * 80 chars/line for headline (10pt regular)
            Each line adds ~20pt of height; 20pt base padding each side.
            """
            _NAME_CHARS_PER_LINE = 60
            _HL_CHARS_PER_LINE   = 80
            _LINE_HEIGHT_PT      = 20  # vertical step per text line
            _MIN_BAND_HEIGHT     = 60  # never shorter than this (pt)
            lines = 1 + (len(name_text) // _NAME_CHARS_PER_LINE)
            if headline_text:
                lines += 1 + (len(headline_text) // _HL_CHARS_PER_LINE)
            return max(_MIN_BAND_HEIGHT, lines * _LINE_HEIGHT_PT + _LINE_HEIGHT_PT)

        name     = _s(data.get("name") or "Unknown")
        headline = _s(data.get("headline") or "")
        dynamic_hdr_h = _hdr_name(name, headline)

        def _draw_header(canvas_obj, doc_obj):
            canvas_obj.saveState()
            # Azure band
            canvas_obj.setFillColor(AZURE)
            canvas_obj.rect(0, PAGE_H - dynamic_hdr_h, PAGE_W, dynamic_hdr_h,
                            fill=1, stroke=0)
            # Robin's egg bottom stripe
            canvas_obj.setFillColor(ROBIN)
            canvas_obj.rect(0, PAGE_H - dynamic_hdr_h, PAGE_W, 3, fill=1, stroke=0)
            # Name: compute target font size to fit available header width.
            _MAX_HDR_W = PAGE_W - ML - MR
            canvas_obj.setFillColor(white)
            # Calculate fitting font size directly instead of looping.
            # stringWidth scales linearly with font size, so:
            #   target_sz = floor(max_sz * max_w / current_w)  (clamped to min)
            _w_at_20 = canvas_obj.stringWidth(name, "Helvetica-Bold", 20)
            if _w_at_20 > 0:
                _name_sz = max(10, min(20, int(20 * _MAX_HDR_W / _w_at_20)))
            else:
                _name_sz = 20
            canvas_obj.setFont("Helvetica-Bold", _name_sz)
            _NAME_FONT_HALF = _name_sz + 2
            name_y = PAGE_H - dynamic_hdr_h + (dynamic_hdr_h - _NAME_FONT_HALF) // 2 + \
                     (18 if headline else 0)
            canvas_obj.drawString(ML, name_y, name)
            # Headline: compute target font size similarly
            if headline:
                _w_hl_at_10 = canvas_obj.stringWidth(headline, "Helvetica", 10)
                if _w_hl_at_10 > 0:
                    _hl_sz = max(7, min(10, int(10 * _MAX_HDR_W / _w_hl_at_10)))
                else:
                    _hl_sz = 10
                canvas_obj.setFont("Helvetica", _hl_sz)
                # Ensure the headline baseline sits at least 10pt above the Robin
                # stripe (which occupies the bottom 3pt of the azure band) so it
                # is always clearly separated from the decorative line.
                _stripe_top = PAGE_H - dynamic_hdr_h + 3
                _headline_y = max(name_y - _hl_sz - 4, _stripe_top + 10)
                canvas_obj.drawString(ML, _headline_y, headline)
            # Footer
            canvas_obj.setFillColor(AZURE)
            canvas_obj.rect(0, 0, PAGE_W, 16, fill=1, stroke=0)
            canvas_obj.setFillColor(ROBIN)
            canvas_obj.setFont("Helvetica", 7)
            canvas_obj.drawCentredString(
                PAGE_W / 2, 4,
                _s("Generated by FIOE Recruiting Platform")
            )
            canvas_obj.restoreState()

        def _draw_later(canvas_obj, doc_obj):
            canvas_obj.saveState()
            canvas_obj.setFillColor(AZURE)
            canvas_obj.rect(0, 0, PAGE_W, 16, fill=1, stroke=0)
            canvas_obj.setFillColor(ROBIN)
            canvas_obj.setFont("Helvetica", 7)
            canvas_obj.drawCentredString(
                PAGE_W / 2, 4,
                _s("Generated by FIOE Recruiting Platform")
            )
            canvas_obj.restoreState()

        buf = _io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=ML, rightMargin=MR,
            topMargin=dynamic_hdr_h + MT,
            bottomMargin=MB + 16,
        )

        # ── Styles ─────────────────────────────────────────────────────────
        contact_style = ParagraphStyle(
            'FIOEContact',
            fontName='Helvetica', fontSize=8.5,
            textColor=AZURE, leading=13,
            backColor=ROBIN_BG,
            leftIndent=8, rightIndent=8,
            spaceBefore=0, spaceAfter=0,
        )
        section_hdr_style = ParagraphStyle(
            'FIOESectionHdr',
            fontName='Helvetica-Bold', fontSize=10,
            textColor=AZURE,
            spaceBefore=12, spaceAfter=2,
            leading=14,
        )
        body_style = ParagraphStyle(
            'FIOEBody',
            fontName='Helvetica', fontSize=9,
            textColor=DARK, leading=13,
            spaceBefore=0, spaceAfter=2,
        )
        title_style = ParagraphStyle(
            'FIOETitle',
            fontName='Helvetica-Bold', fontSize=10,
            textColor=AZURE, leading=14,
            spaceBefore=6, spaceAfter=1,
        )
        sub_style = ParagraphStyle(
            'FIOESub',
            fontName='Helvetica', fontSize=8.5,
            textColor=COOL, leading=12,
            spaceBefore=0, spaceAfter=2,
        )
        skill_style = ParagraphStyle(
            'FIOESkill',
            fontName='Helvetica', fontSize=9,
            textColor=DARK, leading=13,
            spaceBefore=0, spaceAfter=2,
        )
        bullet_style = ParagraphStyle(
            'FIOEBullet',
            fontName='Helvetica', fontSize=9,
            textColor=DARK, leading=13,
            leftIndent=15, firstLineIndent=-8,
            spaceBefore=1, spaceAfter=1,
        )

        story = []
        CW = PAGE_W - ML - MR   # usable content width

        def _section(title):
            story.append(Paragraph(_sp(title).upper(), section_hdr_style))
            story.append(HRFlowable(
                width='100%', thickness=1.5,
                color=ROBIN, spaceAfter=4,
            ))

        def _desc_to_flowables(raw_text):
            """Convert a description string to Platypus flowables.

            Priority order for bullet detection:
            1. Newline-separated lines — only treated as distinct bullet items
               when each line plausibly begins a new sentence (starts with an
               uppercase letter after stripping leading bullet/dash characters)
               OR there are at least 2 lines and every line ends with
               punctuation.  Lines that look like LLM word-wrap artefacts
               (continuation fragments that start mid-sentence) are rejoined
               with a space before bullet processing.
            2. Sentence-boundary split — applied when the sanitised text is a
               single continuous paragraph of >= 80 chars; splits on ". " that
               follows at least two lowercase letters (avoids splitting "Dr. X"
               or "e.g. foo") to produce readable bullet points from prose.
               Requires at least 2 sentences.
            3. Plain Paragraph with <br/> word-wrap as final fallback.

            Bullets are rendered as "- text" Paragraphs with a hanging indent
            so they display correctly in any ReportLab version without needing
            the optional ListFlowable / ListItem classes.
            """
            import re as _re2
            from xml.sax.saxutils import escape as _xmlesc
            if not raw_text or not raw_text.strip():
                return []
            raw = _s(raw_text)

            # Split on explicit newlines.
            raw_segments = [seg.strip() for seg in raw.split('\n') if seg.strip()]

            # Re-join segments that are clearly word-wrap artefacts: a
            # continuation fragment starts with a lowercase letter (after
            # stripping any leading bullet characters) meaning it doesn't begin
            # a new sentence/point.
            _BULLET_LEAD_CHARS = '-*\u2022\u00b7\u25aa\u25ab\u25cf\u25e6\u2023\u2043\u2219 '
            merged = []
            for seg in raw_segments:
                seg_stripped = seg.lstrip(_BULLET_LEAD_CHARS)
                if merged and seg_stripped and seg_stripped[0].islower():
                    # Continuation of the previous segment — rejoin with space.
                    merged[-1] = merged[-1].rstrip() + ' ' + seg_stripped
                else:
                    merged.append(seg)
            segments = merged

            # If no explicit bullet-style segments, try sentence-based splitting
            # on long prose descriptions so each sentence becomes a bullet point.
            if len(segments) == 1 and len(raw) >= 80:
                # Split only when a period follows >=2 lowercase letters
                # (guards against abbreviations like "Dr." or "e.g.") and is
                # followed by whitespace + an uppercase letter.
                sent_segs = [s.strip() for s in
                             _re2.split(r'(?<=[a-z][a-z])\.\s+(?=[A-Z])', segments[0])
                             if s.strip()]
                if len(sent_segs) >= 2:
                    segments = sent_segs
            if len(segments) <= 1:
                return [Paragraph(_spwrap(segments[0] if segments else raw_text), body_style)]
            # Multiple segments → dash-prefixed Paragraphs with hanging indent.
            # Maximum characters per individual bullet segment.
            _MAX_SEG_CHARS = 300
            result = []
            for seg in segments:
                seg_clean = seg.lstrip(_BULLET_LEAD_CHARS)
                if not seg_clean:
                    continue
                result.append(
                    Paragraph(
                        "- " + _xmlesc(seg_clean[:_MAX_SEG_CHARS]),
                        bullet_style,
                    )
                )
            if not result:
                return [Paragraph(_spwrap(raw_text), body_style)]
            return result

        # ── Contact strip ──────────────────────────────────────────────────
        loc   = _sp(data.get("location") or "")
        lnkd  = _sp(data.get("linkedin_url") or "")
        email = _sp(data.get("email") or "")

        contact_items = []
        if loc:   contact_items.append(f"Location: {loc}")
        if lnkd:  contact_items.append(f"LinkedIn: {lnkd}")
        if email: contact_items.append(f"Email: {email}")

        if contact_items:
            for ci in contact_items:
                story.append(Paragraph(ci, contact_style))
            story.append(Spacer(1, 8))

        # ── Professional Summary ───────────────────────────────────────────
        summary = _spwrap(data.get("summary") or "")
        if summary:
            _section("Professional Summary")
            story.append(Paragraph(summary, body_style))
            story.append(Spacer(1, 6))

        # ── Experience ────────────────────────────────────────────────────
        exp_list = data.get("experience") or []
        if exp_list:
            _section("Experience")
            for exp in exp_list:
                if not isinstance(exp, dict):
                    continue
                title   = _sp(exp.get("title") or "")
                company = _sp(exp.get("company") or "")
                dates   = _sp(exp.get("dates") or "")
                raw_desc = exp.get("description") or ""
                if not any([title, company, dates, raw_desc.strip()]):
                    continue
                if title:
                    story.append(Paragraph(title, title_style))
                sub_parts = []
                if company: sub_parts.append(company)
                if dates:   sub_parts.append(dates)
                if sub_parts:
                    story.append(Paragraph("  |  ".join(sub_parts), sub_style))
                for fl in _desc_to_flowables(raw_desc):
                    story.append(fl)
                story.append(Spacer(1, 4))

        # ── Education ─────────────────────────────────────────────────────
        edu_list = data.get("education") or []
        if edu_list:
            _section("Education")
            for edu in edu_list:
                if not isinstance(edu, dict):
                    continue
                school = _sp(edu.get("school") or "")
                degree = _sp(edu.get("degree") or "")
                dates  = _sp(edu.get("dates") or "")
                if not any([school, degree, dates]):
                    continue
                if school:
                    story.append(Paragraph(school, title_style))
                sub_parts = []
                if degree: sub_parts.append(degree)
                if dates:  sub_parts.append(dates)
                if sub_parts:
                    story.append(Paragraph("  |  ".join(sub_parts), sub_style))
                story.append(Spacer(1, 4))

        # ── Skills ────────────────────────────────────────────────────────
        skills = data.get("skills") or []
        if skills:
            _section("Skills")
            # Group skills into rows of 3 using a Table for clean layout
            CHUNK = 3
            tbl_data = []
            for i in range(0, len(skills), CHUNK):
                row_cells = [cell for cell in (_sp(s) for s in skills[i:i + CHUNK]) if cell.strip()]
                # Pad to CHUNK columns
                while len(row_cells) < CHUNK:
                    row_cells.append("")
                tbl_data.append([
                    Paragraph(cell, skill_style) for cell in row_cells
                ])
            if tbl_data:
                col_w = CW / CHUNK
                skills_tbl = Table(tbl_data, colWidths=[col_w] * CHUNK)
                skills_tbl.setStyle(TableStyle([
                    ('VALIGN',     (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING',  (0, 0), (-1, -1), 2),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 2),
                    ('TOPPADDING',   (0, 0), (-1, -1), 2),
                    ('BOTTOMPADDING',(0, 0), (-1, -1), 2),
                ]))
                story.append(skills_tbl)

        if not story:
            story.append(Paragraph("No profile data available.", body_style))

        try:
            doc.build(story,
                      onFirstPage=_draw_header,
                      onLaterPages=_draw_later)
            return buf.getvalue()
        except Exception as _build_err:
            logger.warning(
                "[PDF] Platypus doc.build() failed (%s: %s) — falling back to plain-text renderer",
                type(_build_err).__name__, _build_err,
            )
            # Fall through to the plain-text fallback below

    except ImportError:
        pass

    # ── Fallback: plain text via _lines_to_pdf_bytes ───────────────────────
    lines = []
    if data.get("name"):
        lines.append(("title", str(data["name"])))
    if data.get("headline"):
        lines.append(("key", str(data["headline"])))
    lines.append(("gap", ""))
    for label, val in [("Location", data.get("location")),
                       ("Email",    data.get("email")),
                       ("LinkedIn", data.get("linkedin_url"))]:
        if val:
            lines.append(("item", f"{label}: {val}"))
    lines.append(("gap", ""))
    if data.get("summary"):
        lines.append(("section", "SUMMARY"))
        lines.append(("item", str(data["summary"])))
        lines.append(("gap", ""))
    if data.get("experience"):
        lines.append(("section", "EXPERIENCE"))
    for exp in (data.get("experience") or []):
        h  = str(exp.get("title") or "")
        co = str(exp.get("company") or "")
        dt = str(exp.get("dates") or "")
        dc = str(exp.get("description") or "")
        if co:
            h = f"{h}  \u00b7  {co}" if h else co
        if dt:
            h = f"{h}  |  {dt}" if h else dt
        if h:
            lines.append(("key", h))
        if dc:
            # Split multi-sentence descriptions into dash-prefixed bullet items
            # so the fallback canvas renderer produces readable output.
            import re as _fb_re
            dc_clean = dc.replace('\r', '')
            raw_paras = [p.strip() for p in dc_clean.split('\n') if p.strip()]
            # Re-join continuation fragments (lines starting with a lowercase
            # letter) to undo any LLM word-wrap artefacts before bullet split.
            merged_paras = []
            _bl_chars = '-*\u2022\u00b7 '
            for para in raw_paras:
                stripped = para.lstrip(_bl_chars)
                if merged_paras and stripped and stripped[0].islower():
                    merged_paras[-1] = merged_paras[-1].rstrip() + ' ' + stripped
                else:
                    merged_paras.append(para)
            dc_paras = merged_paras
            if len(dc_paras) <= 1:
                # Try sentence-based split when text is a single prose paragraph
                sents = [s.strip() for s in
                         _fb_re.split(r'(?<=[a-z][a-z])\.\s+(?=[A-Z])', dc_clean)
                         if s.strip()]
                if len(sents) >= 2:
                    dc_paras = sents
            if len(dc_paras) > 1:
                for para in dc_paras:
                    # Strip any leading bullet/dash chars the LLM may have added
                    para = para.lstrip('-*\u2022\u00b7 ')
                    if para:
                        lines.append(("item", f"- {para}"))
            else:
                lines.append(("item", dc_clean))
        lines.append(("gap", ""))
    if data.get("education"):
        lines.append(("section", "EDUCATION"))
        for edu in (data.get("education") or []):
            s  = str(edu.get("school") or "")
            d  = str(edu.get("degree") or "")
            dt = str(edu.get("dates") or "")
            if d:
                s = f"{s}  \u00b7  {d}" if s else d
            if dt:
                s = f"{s}  ({dt})" if s else dt
            if s:
                lines.append(("item", s))
        lines.append(("gap", ""))
    if data.get("skills"):
        lines.append(("section", "SKILLS"))
        skls = data.get("skills") or []
        for i in range(0, len(skls), 5):
            row = [str(x) for x in skls[i:i + 5] if str(x).strip()]
            if row:
                lines.append(("item", "  \u00b7  ".join(row)))
    return _lines_to_pdf_bytes(lines)


@app.get("/api/linkdapi/profile-to-pdf")
@_require_session
def linkdapi_profile_to_pdf():
    """Convert a saved GP profile JSON to a FIOE-branded PDF and save it to disk.

    Query parameters:
      linkedin_url – the LinkedIn profile URL (required); username is extracted
                     from the URL path.

    Reads the previously-saved JSON from ``LINKDAPI_PROFILE_OUTPUT_DIR``,
    converts it to a FIOE-branded A4 PDF, and saves it alongside the JSON in
    the same directory as ``<slug>_<user>.pdf``.  Returns JSON:
      ``{"saved_pdf_filename": "...", "status": "saved"}``
    """
    linkedin_url = (request.args.get("linkedin_url") or "").strip()
    if not linkedin_url:
        return jsonify({"error": "linkedin_url is required"}), 400

    _m = re.search(r"/in/([A-Za-z0-9_%-]+)", linkedin_url)
    if not _m:
        return jsonify({"error": "Could not extract LinkedIn username from URL"}), 400
    username = _m.group(1)

    def _safe_slug(value: str, fallback: str) -> str:
        slug = re.sub(r'[^A-Za-z0-9_-]+', '_', value or "")
        slug = re.sub(r'_+', '_', slug).strip('_')
        return slug or fallback

    active_username = getattr(request, "_session_user", "") or ""
    safe_active_username = _safe_slug(active_username, "unknown")
    safe_profile_username = _safe_slug(username, "linkdapi_profile")
    json_filename = f"{safe_profile_username}_{safe_active_username}.json"

    if os.sep in json_filename or (os.altsep and os.altsep in json_filename):
        return jsonify({"error": "Invalid filename"}), 400

    out_dir = os.path.abspath(LINKDAPI_PROFILE_OUTPUT_DIR)

    try:
        existing = set(os.listdir(out_dir))
    except FileNotFoundError:
        return jsonify({"error": "GP profiles directory does not exist"}), 404
    except OSError as exc:
        logger.exception("[linkdapi PDF] cannot list profiles dir: %s", exc)
        return jsonify({"error": "Cannot read profiles directory"}), 500

    if json_filename not in existing:
        return jsonify({"error": "GP profile not found. Click the GP button first to fetch the profile."}), 404

    # Use the entry from the directory listing (OS-sourced, not user-derived)
    # to construct the read path — this breaks the taint chain from user input.
    matched_json = next(f for f in existing if f == json_filename)
    json_path = os.path.join(out_dir, matched_json)

    try:
        with open(json_path, "r", encoding="utf-8") as fh:
            profile_data = json.load(fh)
    except json.JSONDecodeError:
        return jsonify({"error": "Saved GP profile JSON is corrupted"}), 500
    except Exception as exc:
        logger.exception("[linkdapi PDF] failed to read profile JSON: %s", exc)
        return jsonify({"error": "Failed to read GP profile JSON"}), 500

    try:
        pdf_bytes = _linkdapi_json_to_pdf_bytes(profile_data)
    except Exception as exc:
        logger.exception("[linkdapi PDF] PDF generation failed: %s", exc)
        return jsonify({"error": "PDF generation failed"}), 500

    # Derive the PDF filename from the OS-sourced JSON directory entry
    # (replacing the .json extension with .pdf) so the write path is not
    # tainted by any user-provided value.
    if not matched_json.endswith(".json"):
        return jsonify({"error": "Unexpected JSON filename format"}), 500
    pdf_filename = matched_json[:-5] + ".pdf"
    pdf_path = os.path.join(out_dir, pdf_filename)
    try:
        with open(pdf_path, "wb") as fh:
            fh.write(pdf_bytes)
        logger.info("[linkdapi PDF] saved %s (%d bytes)", pdf_filename, len(pdf_bytes))
    except Exception as exc:
        logger.exception("[linkdapi PDF] failed to save PDF: %s", exc)
        return jsonify({"error": "Failed to save PDF to profiles directory"}), 500

    return jsonify({
        "saved_pdf_filename": pdf_filename,
        "status": "saved",
    })


def _gp_check_profile_pdf_impl():
    """Shared implementation for checking whether a saved GP profile PDF exists.

    Query parameters:
      linkedin_url – the LinkedIn profile URL (required)

    Returns ``{"exists": true, "filename": "<slug>_<user>.pdf"}`` if found,
    or ``{"exists": false, "filename": ""}`` if not.

    Called by all three service endpoints (BrightData, Scrapingdog, Linkdapi)
    since they all write to the same shared profiles directory.  No external
    API call is made — this is a pure filesystem existence check.
    """
    linkedin_url = (request.args.get("linkedin_url") or "").strip()
    if not linkedin_url:
        return jsonify({"exists": False, "filename": ""}), 200

    _m = re.search(r"/in/([A-Za-z0-9_%-]+)", linkedin_url)
    if not _m:
        return jsonify({"exists": False, "filename": ""}), 200
    username = _m.group(1)

    def _safe_slug(value: str, fallback: str) -> str:
        slug = re.sub(r'[^A-Za-z0-9_-]+', '_', value or "")
        slug = re.sub(r'_+', '_', slug).strip('_')
        return slug or fallback

    active_username = getattr(request, "_session_user", "") or ""
    safe_active = _safe_slug(active_username, "unknown")
    safe_profile = _safe_slug(username, "gp_profile")
    pdf_filename = f"{safe_profile}_{safe_active}.pdf"

    out_dir = os.path.abspath(LINKDAPI_PROFILE_OUTPUT_DIR)
    try:
        exists = os.path.isfile(os.path.join(out_dir, pdf_filename))
    except Exception:
        return jsonify({"exists": False, "filename": ""}), 200

    return jsonify({"exists": exists, "filename": pdf_filename if exists else ""}), 200


@app.get("/api/linkdapi/check-profile-pdf")
@_require_session
def linkdapi_check_profile_pdf():
    """Check whether a saved GP profile PDF exists (Linkdapi variant — last resort).

    Delegates to the shared implementation.  The check sequence is:
    BrightData (primary) → Scrapingdog (fallback) → Linkdapi (last resort).
    """
    return _gp_check_profile_pdf_impl()


@app.get("/api/linkdapi/get-profile-pdf")
@_require_session
def linkdapi_get_profile_pdf():
    """Return the saved GP profile PDF bytes for a given LinkedIn URL.

    Used by the Assess button to fetch the PDF before submitting it to
    ``/process/upload_multiple_cvs`` (the same bulk-upload pipeline used by
    the Bulk Assessment button).

    Query parameters:
      linkedin_url – the LinkedIn profile URL (required)

    Returns the PDF as ``application/pdf`` with header
    ``X-PDF-Filename`` carrying the safe filename (used by the frontend
    when constructing the FormData entry).
    """
    from flask import Response as _FlaskResponse

    linkedin_url = (request.args.get("linkedin_url") or "").strip()
    if not linkedin_url:
        return jsonify({"error": "linkedin_url is required"}), 400

    _m = re.search(r"/in/([A-Za-z0-9_%-]+)", linkedin_url)
    if not _m:
        return jsonify({"error": "Could not extract LinkedIn username from URL"}), 400
    username = _m.group(1)

    def _safe_slug(value: str, fallback: str) -> str:
        slug = re.sub(r'[^A-Za-z0-9_-]+', '_', value or "")
        slug = re.sub(r'_+', '_', slug).strip('_')
        return slug or fallback

    active_username = getattr(request, "_session_user", "") or ""
    safe_active_username = _safe_slug(active_username, "unknown")
    safe_profile_username = _safe_slug(username, "linkdapi_profile")
    # The PDF filename encodes both the profile slug and the active session
    # username — so each user can only retrieve PDFs they generated themselves
    # (implicit per-user access control via filename convention).
    pdf_filename = f"{safe_profile_username}_{safe_active_username}.pdf"

    out_dir = os.path.abspath(LINKDAPI_PROFILE_OUTPUT_DIR)

    try:
        existing = set(os.listdir(out_dir))
    except FileNotFoundError:
        return jsonify({"error": "GP profiles directory does not exist"}), 404
    except OSError as exc:
        logger.exception("[linkdapi get-pdf] cannot list profiles dir: %s", exc)
        return jsonify({"error": "Cannot read profiles directory"}), 500

    if pdf_filename not in existing:
        return jsonify({
            "error": "GP profile PDF not found. Click the GP button first to generate the PDF."
        }), 404

    # Use OS-sourced entry (breaks taint chain from user input)
    matched_pdf = next((f for f in existing if f == pdf_filename), None)
    if not matched_pdf:
        return jsonify({"error": "GP profile PDF not found on disk"}), 404
    pdf_path = os.path.join(out_dir, matched_pdf)

    try:
        with open(pdf_path, "rb") as fh:
            pdf_bytes = fh.read()
    except Exception as exc:
        logger.exception("[linkdapi get-pdf] failed to read PDF: %s", exc)
        return jsonify({"error": "Failed to read GP profile PDF"}), 500

    resp = _FlaskResponse(pdf_bytes, mimetype="application/pdf")
    resp.headers["Content-Disposition"] = f"inline; filename=\"{matched_pdf}\""
    resp.headers["X-PDF-Filename"] = matched_pdf
    return resp


def _gp_list_profile_pdfs_impl():
    """Shared implementation for listing saved GP profile PDFs.

    Returns ``{"files": ["slug_user.pdf", ...]}`` — only files whose name ends
    with ``_<session_username>.pdf`` are included, so each user sees only their
    own generated PDFs.

    Called by all three service endpoints (BrightData, Scrapingdog, Linkdapi)
    since they all write to the same shared profiles directory.
    """
    active_username = getattr(request, "_session_user", "") or ""

    def _safe_slug(value: str, fallback: str) -> str:
        slug = re.sub(r'[^A-Za-z0-9_-]+', '_', value or "")
        slug = re.sub(r'_+', '_', slug).strip('_')
        return slug or fallback

    safe_active = _safe_slug(active_username, "unknown")
    user_suffix = f"_{safe_active}.pdf"

    out_dir = os.path.abspath(LINKDAPI_PROFILE_OUTPUT_DIR)
    try:
        all_entries = os.listdir(out_dir)
    except FileNotFoundError:
        return jsonify({"files": []}), 200
    except OSError as exc:
        logger.exception("[gp list-pdfs] cannot list profiles dir: %s", exc)
        return jsonify({"error": "Cannot read profiles directory"}), 500

    user_pdfs = [f for f in all_entries if f.endswith(user_suffix)]
    return jsonify({"files": user_pdfs}), 200


@app.get("/api/linkdapi/list-profile-pdfs")
@_require_session
def linkdapi_list_profile_pdfs():
    """List saved GP profile PDFs for the current session user (Linkdapi variant).

    Delegates to the shared implementation.  BrightData/Serp AI is the canonical
    handler for profile-directory listing across all GP services.
    """
    return _gp_list_profile_pdfs_impl()


@app.get("/api/linkdapi/get-pdf-by-filename")
@_require_session
def linkdapi_get_pdf_by_filename():
    """Serve a saved GP profile PDF by its filename.

    Only PDFs that end with ``_<session_username>.pdf`` may be retrieved —
    this enforces per-user access control (a user cannot fetch another user's
    PDFs by guessing a filename).

    Query parameters:
      filename – the PDF filename (e.g. ``jdoe_admin.pdf``); must belong to
                 the current session user.

    Returns the PDF as ``application/pdf`` with header ``X-PDF-Filename``.
    """
    from flask import Response as _FlaskResponse

    filename = (request.args.get("filename") or "").strip()
    if not filename:
        return jsonify({"error": "filename is required"}), 400

    # Reject any path-traversal or separator characters
    if any(c in filename for c in ('/', '\\', '\0')):
        return jsonify({"error": "Invalid filename"}), 400
    if not filename.lower().endswith('.pdf'):
        return jsonify({"error": "Only PDF files are supported"}), 400

    active_username = getattr(request, "_session_user", "") or ""

    def _safe_slug(value: str, fallback: str) -> str:
        slug = re.sub(r'[^A-Za-z0-9_-]+', '_', value or "")
        slug = re.sub(r'_+', '_', slug).strip('_')
        return slug or fallback

    safe_active = _safe_slug(active_username, "unknown")
    expected_suffix = f"_{safe_active}.pdf"

    # Access control: filename must end with _<session_user>.pdf
    if not filename.endswith(expected_suffix):
        return jsonify({"error": "Access denied"}), 403

    out_dir = os.path.abspath(LINKDAPI_PROFILE_OUTPUT_DIR)
    try:
        existing = set(os.listdir(out_dir))
    except FileNotFoundError:
        return jsonify({"error": "Profiles directory does not exist"}), 404
    except OSError as exc:
        logger.exception("[linkdapi get-pdf-by-filename] cannot list dir: %s", exc)
        return jsonify({"error": "Cannot read profiles directory"}), 500

    if filename not in existing:
        return jsonify({"error": "PDF file not found"}), 404

    # Use OS-sourced entry to break taint chain from user input
    matched = next((f for f in existing if f == filename), None)
    if not matched:
        return jsonify({"error": "PDF file not found"}), 404

    pdf_path = os.path.join(out_dir, matched)
    try:
        with open(pdf_path, "rb") as fh:
            pdf_bytes = fh.read()
    except Exception as exc:
        logger.exception("[linkdapi get-pdf-by-filename] failed to read PDF: %s", exc)
        return jsonify({"error": "Failed to read PDF"}), 500

    resp = _FlaskResponse(pdf_bytes, mimetype="application/pdf")
    resp.headers["Content-Disposition"] = f"inline; filename=\"{matched}\""
    resp.headers["X-PDF-Filename"] = matched
    return resp


@app.route("/api/linkdapi/upload-profile-pdf", methods=["POST"])
@_require_session
def linkdapi_upload_profile_pdf():
    """Upload a pre-saved GP profile PDF from the profiles directory into the DB.

    Request body (JSON):
      linkedin_url – the LinkedIn profile URL (required)
      name         – candidate name (optional, for name-validation skip)

    Finds the saved ``<slug>_<user>.pdf`` in ``LINKDAPI_PROFILE_OUTPUT_DIR``,
    reads it, stores it in the ``process.cv`` column (same as
    ``/process/upload_cv``), then triggers background CV parsing via
    ``analyze_cv_background``.  Returns the parsed CV data dict on success.
    """
    body = request.get_json(silent=True) or {}
    linkedin_url = (body.get("linkedin_url") or "").strip()
    if not linkedin_url:
        return jsonify({"error": "linkedin_url is required"}), 400

    _m = re.search(r"/in/([A-Za-z0-9_%-]+)", linkedin_url)
    if not _m:
        return jsonify({"error": "Could not extract LinkedIn username from URL"}), 400
    username = _m.group(1)

    def _safe_slug(value: str, fallback: str) -> str:
        slug = re.sub(r'[^A-Za-z0-9_-]+', '_', value or "")
        slug = re.sub(r'_+', '_', slug).strip('_')
        return slug or fallback

    active_username = getattr(request, "_session_user", "") or ""
    safe_active_username = _safe_slug(active_username, "unknown")
    safe_profile_username = _safe_slug(username, "linkdapi_profile")
    pdf_filename = f"{safe_profile_username}_{safe_active_username}.pdf"

    out_dir = os.path.abspath(LINKDAPI_PROFILE_OUTPUT_DIR)

    try:
        existing = set(os.listdir(out_dir))
    except FileNotFoundError:
        return jsonify({"error": "GP profiles directory does not exist"}), 404
    except OSError as exc:
        logger.exception("[linkdapi upload-pdf] cannot list profiles dir: %s", exc)
        return jsonify({"error": "Cannot read profiles directory"}), 500

    if pdf_filename not in existing:
        return jsonify({
            "error": "GP profile PDF not found. Click the GP button first to generate the PDF."
        }), 404

    # Use OS-sourced entry to build the read path (breaks taint chain)
    matched_pdf = next((f for f in existing if f == pdf_filename), None)
    if not matched_pdf:
        return jsonify({"error": "GP profile PDF not found on disk"}), 404
    pdf_path = os.path.join(out_dir, matched_pdf)

    try:
        with open(pdf_path, "rb") as fh:
            pdf_bytes = fh.read()
    except Exception as exc:
        logger.exception("[linkdapi upload-pdf] failed to read PDF: %s", exc)
        return jsonify({"error": "Failed to read GP profile PDF"}), 500

    # Store PDF in process table (same logic as process_upload_cv)
    try:
        import psycopg2
        pg_host     = os.getenv("PGHOST", "localhost")
        pg_port     = int(os.getenv("PGPORT", "5432"))
        pg_user     = os.getenv("PGUSER", "postgres")
        pg_password = os.getenv("PGPASSWORD", "")
        pg_db       = os.getenv("PGDATABASE", "candidate_db")
        conn = _pg_connect()
        cur = conn.cursor()
        binary_cv = psycopg2.Binary(pdf_bytes)

        # Try by exact URL first, then normalised partial match.
        # Track the process table primary key so we can pass it to
        # analyze_cv_background, ensuring it updates the correct row
        # (avoids the sourcing-table ID mismatch that caused only-skillset scoring).
        _proc_id_for_bg = None
        from webbridge_cv import _normalize_linkedin_to_path  # type: ignore
        from psycopg2 import sql as _pgsql
        normalized = _normalize_linkedin_to_path(linkedin_url)
        cur.execute(
            "UPDATE process SET cv = %s WHERE linkedinurl = %s",
            (binary_cv, linkedin_url)
        )
        _cv_rowcount = cur.rowcount
        if _cv_rowcount > 0:
            cur.execute(
                "SELECT id FROM process WHERE linkedinurl = %s LIMIT 1",
                (linkedin_url,)
            )
            _id_row = cur.fetchone()
            if _id_row:
                _proc_id_for_bg = _id_row[0]
        if _cv_rowcount == 0 and normalized:
            cur.execute(
                "UPDATE process SET cv = %s WHERE LOWER(linkedinurl) LIKE %s",
                (binary_cv, f"%{normalized}%")
            )
            _cv_rowcount = cur.rowcount
            if _cv_rowcount > 0:
                cur.execute(
                    "SELECT id FROM process WHERE LOWER(linkedinurl) LIKE %s LIMIT 1",
                    (f"%{normalized}%",)
                )
                _id_row = cur.fetchone()
                if _id_row:
                    _proc_id_for_bg = _id_row[0]

        # If no existing process row was found, INSERT a minimal stub so that
        # bulk_assess can locate the record and the background CV parse can
        # write profile fields back to it.
        if _cv_rowcount == 0:
            cand_name   = (body.get("name") or "").strip()
            active_user = getattr(request, "_session_user", "") or ""
            try:
                # Discover which columns exist so we only insert what is present.
                cur.execute("""
                    SELECT column_name FROM information_schema.columns
                    WHERE table_schema = 'public' AND table_name = 'process'
                """)
                _existing_cols = {r[0].lower() for r in cur.fetchall()}

                _ins_fields: list = ["linkedinurl", "cv"]
                _ins_vals:   list = [linkedin_url, binary_cv]
                if "name" in _existing_cols:
                    _ins_fields.append("name")
                    _ins_vals.append(cand_name or linkedin_url)
                if "username" in _existing_cols and active_user:
                    _ins_fields.append("username")
                    _ins_vals.append(active_user)

                # Use RETURNING id so we can pass the exact row PK to
                # analyze_cv_background (avoids sourcing-ID mismatch).
                _ins_sql = _pgsql.SQL(
                    "INSERT INTO process ({}) VALUES ({}) RETURNING id"
                ).format(
                    _pgsql.SQL(", ").join(_pgsql.Identifier(f) for f in _ins_fields),
                    _pgsql.SQL(", ").join(_pgsql.Placeholder() for _ in _ins_fields),
                )
                cur.execute(_ins_sql, _ins_vals)
                _ins_ret = cur.fetchone()
                if _ins_ret:
                    _proc_id_for_bg = _ins_ret[0]
                logger.info(
                    "[linkdapi upload-pdf] Inserted new process row id=%s for %s",
                    _proc_id_for_bg, linkedin_url,
                )
            except Exception as _ins_exc:
                logger.warning(
                    "[linkdapi upload-pdf] Could not insert process row (non-fatal): %s",
                    _ins_exc,
                )
                try:
                    conn.rollback()
                except Exception as _rb_exc:
                    logger.debug(
                        "[linkdapi upload-pdf] rollback after failed INSERT: %s", _rb_exc
                    )

        conn.commit()
        cur.close()
        conn.close()
    except Exception as exc:
        logger.exception("[linkdapi upload-pdf] DB write failed: %s", exc)
        return jsonify({"error": "Failed to store GP profile PDF in database"}), 500

    # Run synchronous CV parse so the result can be returned directly to the
    # frontend (avoids the separate /process/parse_cv_and_update round-trip).
    parse_result = {}
    try:
        from webbridge_cv import _analyze_cv_bytes_sync, analyze_cv_background  # type: ignore
        obj = _analyze_cv_bytes_sync(pdf_bytes)
        if obj:
            parse_result = {
                "skillset":        obj.get("skillset", []),
                "total_years":     obj.get("total_experience_years", 0),
                "tenure":          obj.get("tenure", 0.0),
                "experience":      obj.get("experience", []),
                "education":       obj.get("education", []),
                "product":         obj.get("product_list", []),
                "company":         obj.get("company", ""),
                "job_title":       obj.get("job_title", ""),
                "country":         obj.get("country", ""),
                "job_family":      obj.get("job_family", ""),
                "seniority":       obj.get("seniority", ""),
                "experience_text": "\n".join(obj.get("experience", [])),
                "education_text":  "\n".join(obj.get("education", [])),
            }

        # Run ML enrichment (sector, company normalisation, product/job_family
        # extraction) SYNCHRONOUSLY using the already-parsed obj so all fields
        # are committed to DB before this HTTP response returns.  Passing
        # pre_parsed_obj avoids a second Gemini call; skip_auto_assess=True
        # because the frontend will immediately call performCombinedAssessment
        # (L2) which is richer than the L1 auto-assessment.
        # process_id targets the exact process row so analyze_cv_background
        # never accidentally updates a different row via sourcing-table ID lookup.
        analyze_cv_background(
            linkedin_url, pdf_bytes,
            pre_parsed_obj=obj,
            skip_auto_assess=True,
            process_id=_proc_id_for_bg,
        )
        logger.info(
            "[linkdapi upload-pdf] ML enrichment complete for %s; all fields in DB",
            linkedin_url,
        )
    except ImportError:
        pass
    except Exception as exc:
        logger.warning("[linkdapi upload-pdf] sync parse failed (non-fatal): %s", exc)

    return jsonify({
        "status": "uploaded",
        "pdf_filename": pdf_filename,
        **parse_result,
    })


# ── Scrapingdog API endpoints ────────────────────────────────────────────────

@app.get("/api/scrapingdog/get-profile")
@_require_session
@_check_gp_rate_limit()
def scrapingdog_get_profile():
    """Fetch a LinkedIn profile via Scrapingdog for a given LinkedIn URL.

    Query parameters:
      linkedin_url – the LinkedIn profile URL (required)

    Calls the Scrapingdog API to retrieve the profile, saves the JSON,
    converts to PDF via Gemini, and saves the PDF.
    """
    linkedin_url = (request.args.get("linkedin_url") or "").strip()
    if not linkedin_url:
        return jsonify({"error": "linkedin_url is required"}), 400

    _m = re.search(r"/in/([A-Za-z0-9_%-]+)", linkedin_url)
    if not _m:
        return jsonify({"error": "Could not extract LinkedIn username from URL"}), 400
    username = _m.group(1)

    # Normalise to www.linkedin.com (handles regional subdomains like cn., jp., de., …)
    # Scrapingdog's id parameter requires the full canonical LinkedIn URL.
    normalized_linkedin_url = f"https://www.linkedin.com/in/{username}"

    # Per-user GP key (VIP admin assignment or user's own api_porting config) takes priority.
    _req_user = getattr(request, '_session_user', None) or (request.cookies.get('username') or '').strip()
    _user_gp = _load_user_gp_cfg(_req_user)
    if _user_gp.get('provider') == 'scrapingdog':
        api_key = (_user_gp.get('GP_SCRAPINGDOG_API_KEY') or '').strip()
        if not api_key:
            return jsonify({"error": "ScrapingDog API key is not configured for your account"}), 503
    else:
        gp_cfg = _load_get_profiles_config()
        sd = gp_cfg.get("scrapingdog", {})
        if sd.get("enabled") != "enabled":
            return jsonify({"error": "scrapingdog is not enabled"}), 503
        api_key = (sd.get("api_key") or "").strip()
        if not api_key:
            return jsonify({"error": "Scrapingdog API key is not configured"}), 503

    logger.info("[scrapingdog] fetching profile for %s (normalized: %s)", linkedin_url, normalized_linkedin_url)
    body_str, status_code = _scrapingdog_fetch_profile(normalized_linkedin_url, api_key)

    if status_code == 401:
        return jsonify({"error": "Scrapingdog authentication failed (HTTP 401). Check your API key."}), 401
    if status_code == 403:
        return jsonify({"error": "Scrapingdog returned HTTP 403 — quota may be exceeded or key restricted"}), 403
    if status_code == 404:
        return jsonify({"error": "Profile not found (HTTP 404)"}), 404
    if status_code == 202:
        return jsonify({
            "error": "Job accepted. Results will be available in 2-3 minutes.",
            "status": "processing",
        }), 202
    if status_code >= 400:
        logger.warning("[scrapingdog] upstream HTTP %d; body snippet: %.200s",
                       status_code, body_str[:200] if body_str else "(empty)")
        return jsonify({"error": f"Scrapingdog returned an error (HTTP {status_code})"}), 502
    if status_code == 0:
        return jsonify({"error": "Failed to reach Scrapingdog service"}), 502

    try:
        profile_data = json.loads(body_str)
    except (json.JSONDecodeError, ValueError) as je:
        logger.warning("[scrapingdog] invalid JSON: %s; body snippet: %.200s",
                       je, body_str[:200] if body_str else "(empty)")
        return jsonify({"error": "Invalid JSON from Scrapingdog"}), 502

    # Scrapingdog may wrap the profile in a list; extract the first element.
    if isinstance(profile_data, list):
        if profile_data and isinstance(profile_data[0], dict):
            profile_data = profile_data[0]
        elif not profile_data:
            logger.warning("[scrapingdog] empty list response")
            return jsonify({"error": "No profile data returned (empty list)"}), 502
        else:
            logger.warning("[scrapingdog] list response but first element is %s, not dict",
                           type(profile_data[0]).__name__)
            return jsonify({"error": "No profile data returned (unexpected format)"}), 502

    if not profile_data or not isinstance(profile_data, dict):
        logger.warning("[scrapingdog] unexpected response type %s; snippet: %.200s",
                       type(profile_data).__name__,
                       body_str[:200] if body_str else "(empty)")
        return jsonify({"error": "No profile data returned"}), 502

    def _safe_slug(value: str, fallback: str) -> str:
        slug = re.sub(r'[^A-Za-z0-9_-]+', '_', value or "")
        slug = re.sub(r'_+', '_', slug).strip('_')
        return slug or fallback

    active_username = getattr(request, "_session_user", "") or ""
    safe_active_username = _safe_slug(active_username, "unknown")
    safe_profile_username = _safe_slug(username, "sd_profile")
    out_filename = f"{safe_profile_username}_{safe_active_username}.json"
    out_dir = os.path.abspath(LINKDAPI_PROFILE_OUTPUT_DIR)

    # Reject filenames containing path separators (defence-in-depth)
    if os.sep in out_filename or (os.altsep and os.altsep in out_filename):
        return jsonify({"error": "Invalid filename"}), 400

    out_path = os.path.join(out_dir, out_filename)

    try:
        real_out_dir = os.path.realpath(out_dir)
        real_out_path = os.path.realpath(out_path)
        if os.path.commonpath([real_out_dir, real_out_path]) != real_out_dir:
            return jsonify({"error": "Invalid output path"}), 400
    except ValueError:
        return jsonify({"error": "Invalid output path"}), 400

    try:
        os.makedirs(out_dir, exist_ok=True)
        with open(out_path, "w", encoding="utf-8") as fh:
            json.dump(profile_data, fh, ensure_ascii=False, indent=2)
    except Exception as exc:
        logger.exception("[scrapingdog] failed to save profile JSON: %s", exc)
        return jsonify({"error": "Failed to save profile JSON output file"}), 500

    # Step 2: Convert to PDF using the same pipeline as linkdapi
    try:
        pdf_bytes = _linkdapi_json_to_pdf_bytes(profile_data)
    except Exception as exc:
        logger.exception("[scrapingdog PDF] PDF generation failed: %s", exc)
        return jsonify({"error": "PDF generation failed"}), 500

    pdf_filename = out_filename[:-5] + ".pdf" if out_filename.endswith(".json") else out_filename + ".pdf"
    # pdf_filename derives from out_filename which was already validated above
    pdf_path = os.path.join(out_dir, pdf_filename)
    try:
        with open(pdf_path, "wb") as fh:
            fh.write(pdf_bytes)
        logger.info("[scrapingdog PDF] saved %s (%d bytes)", pdf_filename, len(pdf_bytes))
    except Exception as exc:
        logger.exception("[scrapingdog PDF] failed to save PDF: %s", exc)
        return jsonify({"error": "Failed to save PDF to profiles directory"}), 500

    return jsonify({
        "profile": profile_data,
        "saved_filename": out_filename,
        "saved_pdf_filename": pdf_filename,
        "saved_for_user": safe_active_username,
    })


# ── BrightData API endpoints ──────────────────────────────────────────────────

@app.get("/api/brightdata/get-profile")
@_require_session
@_check_gp_rate_limit()
def brightdata_get_profile():
    """Fetch a LinkedIn profile via BrightData for a given LinkedIn URL.

    Query parameters:
      linkedin_url – the LinkedIn profile URL (required)

    Triggers a BrightData dataset collection job, polls for the result, saves
    the profile JSON, converts to PDF via Gemini, and saves the PDF.
    """
    linkedin_url = (request.args.get("linkedin_url") or "").strip()
    if not linkedin_url:
        return jsonify({"error": "linkedin_url is required"}), 400

    _m = re.search(r"/in/([A-Za-z0-9_%-]+)", linkedin_url)
    if not _m:
        return jsonify({"error": "Could not extract LinkedIn username from URL"}), 400
    username = _m.group(1)

    # Candidate name forwarded from the frontend (used to build the SERP search query)
    candidate_name = request.args.get("name", "").strip()

    # Build a Google SERP URL so BrightData's SERP zone can locate the profile.
    # Include the candidate name (if available) to improve result precision, and
    # constrain the search to the specific profile path.
    import urllib.parse as _urlparse  # noqa: PLC0415
    if candidate_name:
        search_query = f'"{candidate_name}" site:linkedin.com/in/{username}'
    else:
        search_query = f'site:linkedin.com/in/{username}'
    serp_url = f"https://www.google.com/search?q={_urlparse.quote(search_query)}"

    # Per-user GP key (VIP admin assignment or user's own api_porting config) takes priority.
    _req_user = getattr(request, '_session_user', None) or (request.cookies.get('username') or '').strip()
    _user_gp = _load_user_gp_cfg(_req_user)
    if _user_gp.get('provider') == 'brightdata':
        api_key = (_user_gp.get('GP_BRIGHTDATA_API_KEY') or '').strip()
        zone    = (_user_gp.get('GP_BRIGHTDATA_ZONE') or '').strip()
        if not api_key:
            return jsonify({"error": "BrightData API key is not configured for your account"}), 503
        if not zone:
            return jsonify({"error": "BrightData zone is not configured for your account"}), 503
    else:
        gp_cfg = _load_get_profiles_config()
        bd = gp_cfg.get("brightdata", {})
        if bd.get("enabled") != "enabled":
            return jsonify({"error": "BrightData is not enabled"}), 503
        api_key = (bd.get("api_key") or "").strip()
        if not api_key:
            return jsonify({"error": "BrightData API key is not configured"}), 503
        zone = (bd.get("zone") or "").strip()
        if not zone:
            return jsonify({"error": "BrightData zone is not configured"}), 503

    logger.info("[brightdata] fetching profile for %s via SERP URL: %s", linkedin_url, serp_url)
    body_str, status_code = _brightdata_fetch_profile(serp_url, api_key, zone)

    if status_code == 401:
        return jsonify({"error": "BrightData authentication failed (HTTP 401). Check your API key."}), 401
    if status_code == 403:
        return jsonify({"error": "BrightData returned HTTP 403 — quota may be exceeded or key restricted"}), 403
    if status_code >= 400:
        logger.warning("[brightdata] upstream HTTP %d; body snippet: %.200s",
                       status_code, body_str[:200] if body_str else "(empty)")
        return jsonify({"error": f"BrightData returned an error (HTTP {status_code})"}), 502
    if status_code == 0:
        return jsonify({"error": "Failed to reach BrightData service"}), 502

    try:
        raw_data = json.loads(body_str)
    except ValueError as je:
        logger.warning("[brightdata] invalid JSON: %s; body snippet: %.200s",
                       je, body_str[:200] if body_str else "(empty)")
        return jsonify({"error": "Invalid JSON from BrightData"}), 502

    # BrightData returns a JSON array; extract the first profile object
    if isinstance(raw_data, list):
        if raw_data and isinstance(raw_data[0], dict):
            profile_data = raw_data[0]
        elif not raw_data:
            logger.warning("[brightdata] empty list response")
            return jsonify({"error": "No profile data returned (empty list)"}), 502
        else:
            logger.warning("[brightdata] list response but first element is %s, not dict",
                           type(raw_data[0]).__name__)
            return jsonify({"error": "No profile data returned (unexpected format)"}), 502
    elif isinstance(raw_data, dict):
        profile_data = raw_data
    else:
        logger.warning("[brightdata] unexpected response type %s; snippet: %.200s",
                       type(raw_data).__name__, body_str[:200] if body_str else "(empty)")
        return jsonify({"error": "No profile data returned"}), 502

    if not profile_data:
        return jsonify({"error": "No profile data returned"}), 502

    def _safe_slug(value: str, fallback: str) -> str:
        slug = re.sub(r'[^A-Za-z0-9_-]+', '_', value or "")
        slug = re.sub(r'_+', '_', slug).strip('_')
        return slug or fallback

    active_username = getattr(request, "_session_user", "") or ""
    safe_active_username = _safe_slug(active_username, "unknown")
    safe_profile_username = _safe_slug(username, "bd_profile")
    out_filename = f"{safe_profile_username}_{safe_active_username}.json"
    out_dir = os.path.abspath(LINKDAPI_PROFILE_OUTPUT_DIR)

    # Reject filenames containing path separators (defence-in-depth)
    if os.sep in out_filename or (os.altsep and os.altsep in out_filename):
        return jsonify({"error": "Invalid filename"}), 400

    out_path = os.path.join(out_dir, out_filename)

    try:
        real_out_dir = os.path.realpath(out_dir)
        real_out_path = os.path.realpath(out_path)
        if os.path.commonpath([real_out_dir, real_out_path]) != real_out_dir:
            return jsonify({"error": "Invalid output path"}), 400
    except ValueError:
        return jsonify({"error": "Invalid output path"}), 400

    try:
        os.makedirs(out_dir, exist_ok=True)
        with open(real_out_path, "w", encoding="utf-8") as fh:
            json.dump(profile_data, fh, ensure_ascii=False, indent=2)
    except Exception as exc:
        logger.exception("[brightdata] failed to save profile JSON: %s", exc)
        return jsonify({"error": "Failed to save profile JSON output file"}), 500

    # Convert to PDF using the same pipeline as linkdapi / scrapingdog
    try:
        pdf_bytes = _linkdapi_json_to_pdf_bytes(profile_data)
    except Exception as exc:
        logger.exception("[brightdata PDF] PDF generation failed: %s", exc)
        return jsonify({"error": "PDF generation failed"}), 500

    pdf_filename = out_filename[:-5] + ".pdf" if out_filename.endswith(".json") else out_filename + ".pdf"
    pdf_path = os.path.join(out_dir, pdf_filename)
    try:
        real_pdf_path = os.path.realpath(pdf_path)
        if os.path.commonpath([real_out_dir, real_pdf_path]) != real_out_dir:
            return jsonify({"error": "Invalid output path"}), 400
    except ValueError:
        return jsonify({"error": "Invalid output path"}), 400
    try:
        with open(real_pdf_path, "wb") as fh:
            fh.write(pdf_bytes)
        logger.info("[brightdata PDF] saved %s (%d bytes)", pdf_filename, len(pdf_bytes))
    except Exception as exc:
        logger.exception("[brightdata PDF] failed to save PDF: %s", exc)
        return jsonify({"error": "Failed to save PDF to profiles directory"}), 500

    return jsonify({
        "profile": profile_data,
        "saved_filename": out_filename,
        "saved_pdf_filename": pdf_filename,
        "saved_for_user": safe_active_username,
    })


# ── Service-agnostic profile-PDF check / retrieve endpoints ──────────────────
# Scrapingdog and BrightData save their PDFs to the same LINKDAPI_PROFILE_OUTPUT_DIR
# using the same filename convention as Linkdapi.  These thin aliases let the
# frontend call the endpoint that matches the active service rather than always
# routing through the Linkdapi namespace.

@app.get("/api/scrapingdog/list-profile-pdfs")
@_require_session
def scrapingdog_list_profile_pdfs():
    """List saved GP profile PDFs for the current session user (Scrapingdog variant).

    Delegates to the shared implementation used by all GP profile services
    since they all write to the same profiles directory.
    """
    return _gp_list_profile_pdfs_impl()


@app.get("/api/brightdata/list-profile-pdfs")
@_require_session
def brightdata_list_profile_pdfs():
    """List saved GP profile PDFs for the current session user (BrightData/Serp AI variant).

    This is the canonical handler for profile-directory listing across all GP
    services.  Linkdapi and Scrapingdog variants delegate here or to the shared
    implementation; all services write to the same profiles directory.
    """
    return _gp_list_profile_pdfs_impl()


@app.get("/api/scrapingdog/check-profile-pdf")
@_require_session
def scrapingdog_check_profile_pdf():
    """Check whether a saved GP profile PDF exists (Scrapingdog variant — fallback).

    Delegates to the shared implementation used by all GP profile services
    since they all write to the same profiles directory.  The check sequence
    is: BrightData (primary) → Scrapingdog (fallback) → Linkdapi (last resort).
    """
    return _gp_check_profile_pdf_impl()


@app.get("/api/brightdata/check-profile-pdf")
@_require_session
def brightdata_check_profile_pdf():
    """Check whether a saved GP profile PDF exists (BrightData/Serp AI variant — primary).

    This is the canonical handler for profile-existence checks across all GP
    services.  The check sequence is: BrightData (primary) → Scrapingdog
    (fallback) → Linkdapi (last resort).  No BrightData API key is required —
    this is a pure filesystem existence check.
    """
    return _gp_check_profile_pdf_impl()


@app.get("/api/scrapingdog/get-profile-pdf")
@_require_session
def scrapingdog_get_profile_pdf():
    """Return the saved GP profile PDF bytes (Scrapingdog variant).

    Delegates to the shared implementation used by all GP profile services
    since they all write to the same profiles directory.
    """
    return linkdapi_get_profile_pdf()


@app.get("/api/brightdata/get-profile-pdf")
@_require_session
def brightdata_get_profile_pdf():
    """Return the saved GP profile PDF bytes (BrightData variant).

    Delegates to the shared implementation used by all GP profile services
    since they all write to the same profiles directory.
    """
    return linkdapi_get_profile_pdf()


@_require_session
def process_update_tenure():
    """Best-effort endpoint to update the tenure field in the process table.

    Request body (JSON):
      linkedinurl – LinkedIn profile URL
      tenure      – computed tenure (float, years)

    Returns ``{"ok": true}`` on success.  Non-critical; callers should
    swallow errors.
    """
    body = request.get_json(silent=True) or {}
    linkedin_url = (body.get("linkedinurl") or "").strip()
    tenure = body.get("tenure")
    if not linkedin_url or tenure is None:
        return jsonify({"ok": False, "error": "missing params"}), 400
    try:
        tenure_val = float(tenure)
    except (TypeError, ValueError):
        return jsonify({"ok": False, "error": "invalid tenure"}), 400

    try:
        conn = _pg_connect()
        cur = conn.cursor()
        cur.execute(
            "UPDATE process SET tenure = %s WHERE linkedinurl = %s",
            (tenure_val, linkedin_url),
        )
        if cur.rowcount == 0:
            from webbridge_cv import _normalize_linkedin_to_path  # type: ignore
            normalized = _normalize_linkedin_to_path(linkedin_url)
            if normalized:
                # Escape SQL wildcard chars to prevent unintended pattern matches
                safe_norm = normalized.replace("\\", "\\\\").replace("%", "\\%").replace("_", "\\_")
                cur.execute(
                    "UPDATE process SET tenure = %s WHERE LOWER(linkedinurl) LIKE %s ESCAPE '\\'",
                    (tenure_val, f"%{safe_norm}%"),
                )
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"ok": True})
    except ImportError:
        return jsonify({"ok": False, "error": "psycopg2 not available"}), 500
    except Exception as exc:
        logger.warning("[process/update] tenure update failed (non-fatal): %s", exc)
        return jsonify({"ok": False, "error": "tenure update failed"}), 500


@app.post("/api/user-service-config/activate")
def user_svc_config_activate():
    """Store per-user service config. Encrypts when PORTING_SECRET is set,
    otherwise writes plaintext JSON (matching server.js writeUserServiceConfig)."""
    username, err = _porting_login_required()
    if err:
        return err
    try:
        body = request.get_json(silent=True) or {}
        cfg = {
            'username': username,
            'search': body.get('search') or {},
            'llm': body.get('llm') or {},
            'email_verif': body.get('email_verif') or {},
            'contact_gen': body.get('contact_gen') or {},
        }
        raw = json.dumps(cfg).encode('utf-8')
        if os.getenv("PORTING_SECRET", "").strip():
            # Encrypted storage — server.js can decrypt with the same PORTING_SECRET
            encrypted = _svc_config_encrypt(raw)
            dest = _svc_config_path(username)
            with open(dest, 'wb') as fh:
                fh.write(encrypted)
        else:
            # Plaintext JSON fallback — matches server.js behavior when no PORTING_SECRET
            dest = _svc_config_json_path(username)
            with open(dest, 'w', encoding='utf-8') as fh:
                fh.write(raw.decode('utf-8'))
        log_infrastructure("user_svc_config_activated", username=username,
                           detail="Per-user service config activated", status="success")
        return jsonify({"ok": True, "active": True})
    except Exception as exc:
        logger.exception("[user-service-config/activate]")
        return jsonify({"error": "Activation failed"}), 500


@app.delete("/api/user-service-config/deactivate")
def user_svc_config_deactivate():
    """Delete the per-user service config (keys are wiped)."""
    username, err = _porting_login_required()
    if err:
        return err
    try:
        # Remove both .enc and .json files (matches server.js deleteUserServiceConfig)
        for fp in (_svc_config_path(username), _svc_config_json_path(username)):
            try:
                if os.path.isfile(fp):
                    os.remove(fp)
            except OSError:
                pass
        log_infrastructure("user_svc_config_deactivated", username=username,
                           detail="Per-user service config removed", status="success")
        return jsonify({"ok": True, "active": False})
    except Exception as exc:
        logger.exception("[user-service-config/deactivate]")
        return jsonify({"error": "Deactivation failed"}), 500


# ── Admin VIP endpoints — manage user/access-level service configs ────────────

def _vip_read_user_svc(username: str):
    """Read and return the parsed service config dict for *username*, or None.
    *username* must already be validated by the caller (e.g. via _VIP_USERNAME_RE)."""
    stored = None
    safe   = _porting_safe_name(username)
    svc_dir = os.path.realpath(os.path.join(_PORTING_INPUT_DIR, 'user-services'))
    enc_path  = _svc_config_path(safe)
    json_path = _svc_config_json_path(safe)
    # Confinement: ensure both resolved paths stay within the expected directory.
    if (not os.path.realpath(enc_path).startswith(svc_dir + os.sep) and
            os.path.realpath(enc_path) != svc_dir):
        raise ValueError("Path traversal detected for enc_path")
    if (not os.path.realpath(json_path).startswith(svc_dir + os.sep) and
            os.path.realpath(json_path) != svc_dir):
        raise ValueError("Path traversal detected for json_path")
    if os.path.isfile(enc_path):
        try:
            with open(enc_path, 'rb') as fh:
                raw = fh.read()
            stored = json.loads(_svc_config_decrypt(raw).decode('utf-8'))
        except Exception:
            logger.warning("[vip] .enc decrypt failed for %s", safe, exc_info=True)
    if stored is None and os.path.isfile(json_path):
        try:
            with open(json_path, 'r', encoding='utf-8') as fh:
                stored = json.load(fh)
        except Exception:
            logger.warning("[vip] .json parse failed for %s", safe, exc_info=True)
    return stored


def _vip_write_user_svc(username: str, cfg: dict) -> None:
    """Write *cfg* as the service config for *username* (encrypts when PORTING_SECRET set).
    *username* must already be validated by the caller (e.g. via _VIP_USERNAME_RE)."""
    safe    = _porting_safe_name(username)
    svc_dir = os.path.realpath(os.path.join(_PORTING_INPUT_DIR, 'user-services'))
    enc_path  = _svc_config_path(safe)
    json_path = _svc_config_json_path(safe)
    # Confinement: ensure resolved paths stay within the expected directory.
    if (not os.path.realpath(enc_path).startswith(svc_dir + os.sep) and
            os.path.realpath(enc_path) != svc_dir):
        raise ValueError("Path traversal detected for enc_path")
    if (not os.path.realpath(json_path).startswith(svc_dir + os.sep) and
            os.path.realpath(json_path) != svc_dir):
        raise ValueError("Path traversal detected for json_path")
    cfg['username'] = safe
    raw = json.dumps(cfg).encode('utf-8')
    if os.getenv("PORTING_SECRET", "").strip():
        with open(enc_path, 'wb') as fh:
            fh.write(_svc_config_encrypt(raw))
    else:
        with open(json_path, 'w', encoding='utf-8') as fh:
            fh.write(raw.decode('utf-8'))


def _vip_mask_keys(cfg: dict) -> dict:
    """Return a copy of cfg with sensitive key values replaced by '***'."""
    SENSITIVE = {
        'SERPER_API_KEY', 'DATAFORSEO_PASSWORD', 'LINKEDIN_API_KEY',
        'OPENAI_API_KEY', 'ANTHROPIC_API_KEY', 'NEVERBOUNCE_API_KEY',
        'ZEROBOUNCE_API_KEY', 'BOUNCER_API_KEY', 'CONTACTOUT_API_KEY',
        'APOLLO_API_KEY', 'ROCKETREACH_API_KEY',
        # Get Profile provider keys
        'GP_LINKDAPI_API_KEY', 'GP_BRIGHTDATA_API_KEY', 'GP_SCRAPINGDOG_API_KEY',
    }
    result = {}
    for section, val in cfg.items():
        if isinstance(val, dict):
            masked = {}
            for k, v in val.items():
                masked[k] = '***' if k in SENSITIVE and v else v
            result[section] = masked
        else:
            result[section] = val
    return result


# Regex that matches valid usernames accepted by _porting_safe_name (only alphanumerics + _-@.)
_VIP_USERNAME_RE = _re.compile(r'^[a-zA-Z0-9_@.\-]{1,128}$')


@app.get("/admin/vip/user-service-config")
@_require_admin
def admin_vip_get_user_svc_config():
    """Return the current service-config status for a target user (keys masked)."""
    target = (request.args.get("username") or "").strip()
    if not target or not _VIP_USERNAME_RE.match(target):
        return jsonify({"error": "username is required and must contain only letters, numbers, and _@.-"}), 400
    try:
        stored = _vip_read_user_svc(target)
        if stored is None:
            return jsonify({"active": False, "config": {}})
        providers = {
            'search':      stored.get('search',      {}).get('provider', 'google_cse'),
            'llm':         stored.get('llm',         {}).get('provider', 'gemini'),
            'email_verif': stored.get('email_verif', {}).get('provider', 'default'),
            'contact_gen': stored.get('contact_gen', {}).get('provider', 'gemini'),
            'get_profile': stored.get('get_profile', {}).get('provider', 'platform'),
        }
        return jsonify({"active": True, "providers": providers, "config": _vip_mask_keys(stored)})
    except Exception:
        logger.exception("[admin/vip/user-service-config GET]")
        return jsonify({"error": "Could not read user service config"}), 500


@app.post("/admin/vip/user-service-config")
@_csrf_required
@_require_admin
def admin_vip_set_user_svc_config():
    """Admin: write service-config for a target user on their behalf."""
    body = request.get_json(force=True, silent=True) or {}
    target = (body.get("username") or "").strip()
    if not target or not _VIP_USERNAME_RE.match(target):
        return jsonify({"error": "username is required and must contain only letters, numbers, and _@.-"}), 400
    try:
        cfg = {
            'search':      body.get('search')      or {},
            'llm':         body.get('llm')         or {},
            'email_verif': body.get('email_verif') or {},
            'contact_gen': body.get('contact_gen') or {},
            'get_profile': body.get('get_profile') or {},
        }
        # Preserve existing key values for fields left blank (placeholder kept)
        stored = _vip_read_user_svc(target) or {}
        for section in ('search', 'llm', 'email_verif', 'contact_gen', 'get_profile'):
            existing = stored.get(section) or {}
            new_sec  = cfg[section]
            merged = dict(existing)
            merged.update({k: v for k, v in new_sec.items() if v})
            cfg[section] = merged
        _vip_write_user_svc(target, cfg)
        log_infrastructure("admin_vip_user_svc_set", username=target,
                           detail="Admin set service config for user", status="success")
        return jsonify({"ok": True})
    except Exception:
        logger.exception("[admin/vip/user-service-config POST]")
        return jsonify({"error": "Save failed"}), 500


@app.delete("/admin/vip/user-service-config")
@_csrf_required
@_require_admin
def admin_vip_delete_user_svc_config():
    """Admin: delete service-config for a target user."""
    body = request.get_json(force=True, silent=True) or {}
    target = (body.get("username") or "").strip()
    if not target or not _VIP_USERNAME_RE.match(target):
        return jsonify({"error": "username is required and must contain only letters, numbers, and _@.-"}), 400
    try:
        safe    = _porting_safe_name(target)
        svc_dir = os.path.realpath(os.path.join(_PORTING_INPUT_DIR, 'user-services'))
        for fp in (_svc_config_path(safe), _svc_config_json_path(safe)):
            # Confinement check before removal
            if (not os.path.realpath(fp).startswith(svc_dir + os.sep) and
                    os.path.realpath(fp) != svc_dir):
                continue
            try:
                if os.path.isfile(fp):
                    os.remove(fp)
            except OSError:
                pass
        log_infrastructure("admin_vip_user_svc_deleted", username=safe,
                           detail="Admin cleared service config for user", status="success")
        return jsonify({"ok": True})
    except Exception:
        logger.exception("[admin/vip/user-service-config DELETE]")
        return jsonify({"error": "Clear failed"}), 500


# ── Access-level service configs (stored in rate_limits.json) ──────────────

_AL_SVC_KEY = "access_level_service_configs"


@app.get("/admin/vip/access-level-service-config")
@_require_admin
def admin_vip_get_al_svc_config():
    """Return the stored service-config for an access level."""
    level = (request.args.get("level") or "").strip()
    if not level:
        return jsonify({"error": "level is required"}), 400
    try:
        cfg = _load_rate_limits()
        al_cfgs = cfg.get(_AL_SVC_KEY) or {}
        lvl_cfg = al_cfgs.get(level)
        if not lvl_cfg:
            return jsonify({"active": False, "config": {}})
        return jsonify({"active": True, "config": _vip_mask_keys(lvl_cfg)})
    except Exception:
        logger.exception("[admin/vip/access-level-service-config GET]")
        return jsonify({"error": "Could not read access-level service config"}), 500


@app.post("/admin/vip/access-level-service-config")
@_csrf_required
@_require_admin
def admin_vip_set_al_svc_config():
    """Admin: save a service-config template for an access level."""
    body = request.get_json(force=True, silent=True) or {}
    level = (body.get("level") or "").strip()
    if not level:
        return jsonify({"error": "level is required"}), 400
    try:
        cfg = _load_rate_limits()
        al_cfgs = cfg.get(_AL_SVC_KEY) or {}
        existing = al_cfgs.get(level) or {}
        new_cfg = {}
        for section in ('search', 'llm', 'email_verif', 'contact_gen', 'get_profile'):
            merged = dict(existing.get(section) or {})
            merged.update({k: v for k, v in (body.get(section) or {}).items() if v})
            new_cfg[section] = merged
        al_cfgs[level] = new_cfg
        cfg[_AL_SVC_KEY] = al_cfgs
        _save_rate_limits(cfg)
        log_infrastructure("admin_vip_al_svc_set", username="admin",
                           detail=f"Admin set service config for level {level}", status="success")
        return jsonify({"ok": True})
    except Exception:
        logger.exception("[admin/vip/access-level-service-config POST]")
        return jsonify({"error": "Save failed"}), 500


@app.delete("/admin/vip/access-level-service-config")
@_csrf_required
@_require_admin
def admin_vip_delete_al_svc_config():
    """Admin: remove the service-config template for an access level."""
    body = request.get_json(force=True, silent=True) or {}
    level = (body.get("level") or "").strip()
    if not level:
        return jsonify({"error": "level is required"}), 400
    try:
        cfg = _load_rate_limits()
        al_cfgs = cfg.get(_AL_SVC_KEY) or {}
        al_cfgs.pop(level, None)
        cfg[_AL_SVC_KEY] = al_cfgs
        _save_rate_limits(cfg)
        log_infrastructure("admin_vip_al_svc_deleted", username="admin",
                           detail=f"Admin cleared service config for level {level}", status="success")
        return jsonify({"ok": True})
    except Exception:
        logger.exception("[admin/vip/access-level-service-config DELETE]")
        return jsonify({"error": "Clear failed"}), 500


@app.post("/admin/vip/validate-config")
@_csrf_required
@_require_admin
def admin_vip_validate_config():
    """Admin: validate provider keys in a VIP service config (does NOT store anything).
    Delegates to the existing user-service-config/validate logic for the four standard
    sections (search, llm, email_verif, contact_gen) and adds additional validation for
    the get_profile section (Linkdapi / BrightData / ScrapingDog)."""
    import urllib.request as _ureq3
    import urllib.parse   as _uparse3

    body = request.get_json(force=True, silent=True) or {}

    # ── Probe helper (local to this endpoint) ─────────────────────────────────
    def _probe(url, headers=None, timeout=8):
        req = _ureq3.Request(url, headers=headers or {})
        try:
            with _ureq3.urlopen(req, timeout=timeout) as resp:
                return resp.status, resp.read().decode('utf-8', errors='replace')
        except Exception as exc:
            if hasattr(exc, 'code'):
                try:
                    return exc.code, exc.read().decode('utf-8', errors='replace')
                except Exception:
                    return exc.code, ''
            return None, ''

    results = []

    # ── Standard sections (search / llm / email_verif / contact_gen) ──────────
    std_sections = {k: body[k] for k in ('search', 'llm', 'email_verif', 'contact_gen') if k in body}
    try:
        for r in _run_svc_config_validation(std_sections):
            results.append(r)
    except Exception:
        logger.warning("[admin/vip/validate-config] std validation failed", exc_info=True)

    # ── Get Profile section ────────────────────────────────────────────────────
    gp = body.get('get_profile') or {}
    gp_provider = (gp.get('provider') or 'platform').strip()

    if gp_provider in ('platform', ''):
        results.append({'label': 'Get Profile', 'status': 'ok',
                        'detail': 'Using platform default — no custom key required.'})
    elif gp_provider == 'linkdapi':
        key = (gp.get('GP_LINKDAPI_API_KEY') or '').strip()
        if not key:
            results.append({'label': 'Linkdapi (Get Profile)', 'status': 'error',
                            'detail': 'GP_LINKDAPI_API_KEY is required.'})
        else:
            # linkdapi.com/api/v1/profile/full?username=test — 401 = bad key
            status, body_txt = _probe(
                'https://linkdapi.com/api/v1/profile/full?username=test',
                headers={'Authorization': f'Token {key}'})
            if status == 401:
                results.append({'label': 'Linkdapi', 'status': 'error',
                                'detail': 'Authentication failed (HTTP 401). Check your API key.'})
            elif status in (200, 400, 404, 422):
                results.append({'label': 'Linkdapi', 'status': 'ok',
                                'detail': 'API key accepted.'})
            elif status == 403:
                results.append({'label': 'Linkdapi', 'status': 'warn',
                                'detail': 'HTTP 403 — key may be valid but quota exceeded or account restricted.'})
            else:
                results.append({'label': 'Linkdapi', 'status': 'warn',
                                'detail': f'Unexpected HTTP {status}.' if status else 'Could not reach Linkdapi API.'})
    elif gp_provider == 'brightdata':
        key  = (gp.get('GP_BRIGHTDATA_API_KEY') or '').strip()
        zone = (gp.get('GP_BRIGHTDATA_ZONE') or '').strip()
        if not key:
            results.append({'label': 'BrightData (Get Profile)', 'status': 'error',
                            'detail': 'GP_BRIGHTDATA_API_KEY is required.'})
        elif not zone:
            results.append({'label': 'BrightData (Get Profile)', 'status': 'error',
                            'detail': 'GP_BRIGHTDATA_ZONE is required.'})
        else:
            # BrightData: POST /request with minimal body — 401/403 = bad key/zone
            import urllib.error as _uerr3
            try:
                import json as _json3
                req_data = _json3.dumps({'zone': zone, 'url': 'https://www.google.com', 'format': 'json'}).encode()
                req3 = _ureq3.Request(
                    'https://api.brightdata.com/request',
                    data=req_data,
                    headers={'Content-Type': 'application/json', 'Authorization': f'Bearer {key}'},
                )
                with _ureq3.urlopen(req3, timeout=8) as resp3:
                    bd_status = resp3.status
                    bd_body   = resp3.read().decode('utf-8', errors='replace')
            except _uerr3.HTTPError as e3:
                bd_status = e3.code
                try:
                    bd_body = e3.read().decode('utf-8', errors='replace')
                except Exception:
                    bd_body = ''
            except Exception as e3:
                bd_status = None
                bd_body   = str(e3)
            if bd_status == 401:
                results.append({'label': 'BrightData', 'status': 'error',
                                'detail': 'Authentication failed (HTTP 401). Check your API key.'})
            elif bd_status == 403:
                results.append({'label': 'BrightData', 'status': 'error',
                                'detail': 'Access denied (HTTP 403). Check your zone and API key.'})
            elif bd_status in (200, 202):
                results.append({'label': 'BrightData', 'status': 'ok',
                                'detail': 'API key and zone accepted.'})
            elif bd_status is not None:
                results.append({'label': 'BrightData', 'status': 'warn',
                                'detail': f'HTTP {bd_status} — key may be valid but check your zone/plan.'})
            else:
                results.append({'label': 'BrightData', 'status': 'warn',
                                'detail': 'Could not reach BrightData API — please try again.'})
    elif gp_provider == 'scrapingdog':
        key = (gp.get('GP_SCRAPINGDOG_API_KEY') or '').strip()
        if not key:
            results.append({'label': 'ScrapingDog (Get Profile)', 'status': 'error',
                            'detail': 'GP_SCRAPINGDOG_API_KEY is required.'})
        else:
            # ScrapingDog: GET /profile with a test username and api_key — 401 = bad key
            test_url = (
                f'https://api.scrapingdog.com/profile'
                f'?api_key={_uparse3.quote(key, safe="")}'
                f'&id=https://www.linkedin.com/in/test'
                f'&type=profile&premium=false&webhook=false&fresh=false'
            )
            status, _ = _probe(test_url)
            if status == 401:
                results.append({'label': 'ScrapingDog', 'status': 'error',
                                'detail': 'Authentication failed (HTTP 401). Check your API key.'})
            elif status in (200, 400, 404, 422, 402):
                results.append({'label': 'ScrapingDog', 'status': 'ok',
                                'detail': 'API key accepted.'})
            elif status == 429:
                results.append({'label': 'ScrapingDog', 'status': 'warn',
                                'detail': 'API key valid but rate-limited (HTTP 429).'})
            else:
                results.append({'label': 'ScrapingDog', 'status': 'warn',
                                'detail': f'Unexpected HTTP {status}.' if status else 'Could not reach ScrapingDog API.'})
    else:
        results.append({'label': 'Get Profile', 'status': 'warn',
                        'detail': 'Unknown Get Profile provider selected.'})
    return jsonify({'results': results})


@app.get("/load_search_criteria")
def load_search_criteria():
    """Return the saved search criteria and profile list for the given username and role_tag.

    Query params:
        username  – recruiter username
        role_tag  – recruiter's active role tag
    Returns the criteria JSON object (including profiles list), or 404 if no file exists.
    """
    username = (request.args.get("username") or "").strip()
    role_tag = (request.args.get("role_tag") or "").strip()
    if not username or not role_tag:
        return jsonify({"error": "username and role_tag are required"}), 400
    filepath = _get_criteria_filepath(username, role_tag)
    if not filepath:
        return jsonify({"error": "No criteria file found"}), 404
    try:
        with open(filepath, "r", encoding="utf-8") as fh:
            record = json.load(fh)
    except FileNotFoundError:
        return jsonify({"error": "No criteria file found"}), 404
    except Exception as exc:
        logger.warning(f"[load_search_criteria] Failed to read {filepath}: {exc}")
        return jsonify({"error": "Failed to read criteria file"}), 500
    return jsonify({
        "ok": True,
        "criteria": record.get("criteria") or {},
        "name": record.get("name") or record.get("profiles") or [],
    }), 200


@app.post("/save_search_criteria")
def save_search_criteria():
    """Save the search category breakdown criteria to a JSON file on the server.

    Expected payload:
        username   – recruiter username (e.g. "orlha")
        role_tag   – current role tag    (e.g. "Site Activation Manager")
        criteria   – object with keys: Job Title, Seniority, Sector, Country,
                     Company, Skillset, Tenure
    File is written to  <CRITERIA_OUTPUT_DIR>/<role_tag> <username>.json
    """
    try:
        body = request.get_json(force=True, silent=True) or {}
        username = (body.get("username") or "").strip()
        role_tag = (body.get("role_tag") or "").strip()
        criteria = body.get("criteria") or {}

        if not username or not role_tag:
            return jsonify({"error": "username and role_tag are required"}), 400

        filepath = _get_criteria_filepath(username, role_tag)
        if not filepath:
            return jsonify({"error": "Invalid role_tag or username after sanitization"}), 400
        filename = os.path.basename(filepath)

        os.makedirs(CRITERIA_OUTPUT_DIR, exist_ok=True)

        # Fetch profile names from the sourcing table for this user/role search
        profile_names = []
        try:
            _pconn = _pg_connect()
            try:
                _pcur = _pconn.cursor()
                _pcur.execute(
                    "SELECT DISTINCT name FROM sourcing "
                    "WHERE username=%s AND role_tag=%s AND name IS NOT NULL AND name != ''",
                    (username, role_tag)
                )
                # Strip the "님" honorific suffix and surrounding whitespace
                profile_names = [
                    row[0].replace("님", "").strip()
                    for row in _pcur.fetchall()
                    if row[0]
                ]
                _pcur.close()
            finally:
                _pconn.close()
        except Exception as _pe:
            logger.warning(f"[save_search_criteria] Could not fetch profile names: {_pe}")

        record = {
            "role_tag": role_tag,
            "username": username,
            "saved_at": datetime.utcnow().isoformat() + "Z",
            "name": profile_names,
            "criteria": {
                "Job Title":  criteria.get("Job Title") or [],
                "Seniority":  criteria.get("Seniority") or "",
                "Sector":     criteria.get("Sector") or [],
                "Country":    criteria.get("Country") or "",
                "Company":    criteria.get("Company") or [],
                "Skillset":   criteria.get("Skillset") or [],
                "Tenure":     criteria.get("Tenure"),
            }
        }

        with open(filepath, "w", encoding="utf-8") as fh:
            json.dump(record, fh, ensure_ascii=False, indent=2)

        logger.info(f"[save_search_criteria] Written to {filepath} with {len(profile_names)} profile(s)")
        return jsonify({"ok": True, "file": filename, "name": len(profile_names)}), 200

    except Exception as exc:
        logger.exception("[save_search_criteria]")
        return jsonify({"error": str(exc)}), 500


def _find_criteria_file_for_candidate(candidate_name: str):
    """Scan CRITERIA_OUTPUT_DIR for a JSON file whose 'name' list contains candidate_name.
    Returns the file path and parsed record, or (None, None) if not found.
    """
    if not candidate_name or not os.path.isdir(CRITERIA_OUTPUT_DIR):
        return None, None
    norm = candidate_name.replace("님", "").strip().lower()
    try:
        for fname in os.listdir(CRITERIA_OUTPUT_DIR):
            if not fname.endswith(".json"):
                continue
            fpath = os.path.join(CRITERIA_OUTPUT_DIR, fname)
            try:
                with open(fpath, "r", encoding="utf-8") as fh:
                    rec = json.load(fh)
                names = rec.get("name") or rec.get("profiles") or []
                for n in names:
                    if n.replace("님", "").strip().lower() == norm:
                        return fpath, rec
            except Exception:
                continue
    except Exception:
        pass
    return None, None


def _criteria_record_to_pdf_bytes(record: dict) -> bytes:
    """Convert a criteria JSON record to a minimal PDF.
    Falls back to reportlab if available, otherwise writes a raw minimal PDF.
    """
    # -- Flatten content to lines -----------------------------------------------
    lines = []
    lines.append(("title", f"Search Criteria Report"))
    lines.append(("gap", ""))
    lines.append(("key", f"Role: {record.get('role_tag', '')}"))
    lines.append(("key", f"Generated: {record.get('saved_at', '')}"))
    lines.append(("gap", ""))
    lines.append(("section", "CRITERIA"))
    criteria = record.get("criteria") or {}
    for k, v in criteria.items():
        if isinstance(v, list):
            v_str = ", ".join(str(x) for x in v) if v else "—"
        else:
            v_str = str(v) if v is not None else "—"
        lines.append(("item", f"{k}: {v_str}"))
    lines.append(("gap", ""))
    lines.append(("section", "SOURCED PROFILES"))
    profiles = record.get("name") or record.get("profiles") or []
    for p in profiles:
        lines.append(("item", f"  • {p}"))

    # -- Try reportlab first -----------------------------------------------------
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import io as _io
        buf = _io.BytesIO()
        c = rl_canvas.Canvas(buf, pagesize=A4)
        w, h = A4
        y = h - 50
        for kind, text in lines:
            if kind == "gap":
                y -= 10; continue
            if kind == "title":
                c.setFont("Helvetica-Bold", 16)
            elif kind == "section":
                c.setFont("Helvetica-Bold", 12)
                y -= 4
            else:
                c.setFont("Helvetica", 11)
            # Encode to latin-1; strip non-encodable characters silently
            safe = str(text).encode("latin-1", errors="ignore").decode("latin-1")
            c.drawString(40, y, safe)
            y -= 16
            if y < 60:
                c.showPage()
                y = h - 50
        c.save()
        return buf.getvalue()
    except ImportError:
        pass

    # -- Raw minimal PDF fallback (Latin-1 only) ---------------------------------
    import struct as _struct
    text_parts = []
    for kind, text in lines:
        safe = str(text).encode("latin-1", errors="ignore").decode("latin-1")
        size = 16 if kind == "title" else (12 if kind == "section" else 11)
        bold = "-Bold" if kind in ("title", "section") else ""
        text_parts.append((safe, size, bold))

    def _pdf_str(s):
        return s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    stream_lines = ["BT", "/F1 16 Tf", "40 792 Td"]
    y = 742
    for safe, size, bold in text_parts:
        if not safe.strip():
            y -= 10
            stream_lines.append(f"0 -{10} Td")
            continue
        fname = f"/F{'B' if bold else '1'}"
        stream_lines.append(f"{fname} {size} Tf")
        stream_lines.append(f"({_pdf_str(safe)}) Tj")
        y -= 16
        stream_lines.append(f"0 -{16} Td")
    stream_lines.append("ET")
    stream = "\n".join(stream_lines).encode("latin-1")

    def _obj(n, d, s=None):
        out = f"{n} 0 obj\n{d}\n"
        if s is not None:
            out += f"stream\n"
            out = out.encode("latin-1") + s + b"\nendstream\n"
            return out + b"endobj\n"
        return (out + "endobj\n").encode("latin-1")

    o1 = _obj(1, "<< /Type /Catalog /Pages 2 0 R >>")
    o2 = _obj(2, "<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    o3 = _obj(3, f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Contents 4 0 R /Resources << /Font << /F1 5 0 R /FB 6 0 R >> >> >>")
    slen = len(stream)
    o4 = _obj(4, f"<< /Length {slen} >>", stream)
    o5 = _obj(5, "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    o6 = _obj(6, "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>")

    header = b"%PDF-1.4\n"
    body = o1 + o2 + o3 + o4 + o5 + o6
    offsets = []
    pos = len(header)
    for chunk in (o1, o2, o3, o4, o5, o6):
        offsets.append(pos)
        pos += len(chunk)

    xref_offset = len(header) + len(body)
    xref = f"xref\n0 7\n0000000000 65535 f \n"
    for off in offsets:
        xref += f"{off:010d} 00000 n \n"
    trailer = f"trailer\n<< /Size 7 /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n"
    return header + body + (xref + trailer).encode("latin-1")


@app.get("/sourcing/has_criteria_json")
def has_criteria_json():
    """Check whether a criteria JSON file exists for the given candidate.

    Query params:
        linkedin  – candidate LinkedIn URL
        name      – candidate name (used for lookup)
    Returns { "exists": true/false }
    """
    name = (request.args.get("name") or "").strip()
    linkedin = (request.args.get("linkedin") or "").strip()
    if not name and not linkedin:
        return jsonify({"exists": False}), 200
    # If no name provided, try to look it up from sourcing table
    if not name and linkedin:
        try:
            conn = _pg_connect()
            try:
                cur = conn.cursor()
                cur.execute(
                    "SELECT name FROM sourcing WHERE linkedinurl=%s AND name IS NOT NULL LIMIT 1",
                    (linkedin,)
                )
                row = cur.fetchone()
                cur.close()
                if row and row[0]:
                    name = row[0].replace("님", "").strip()
            finally:
                conn.close()
        except Exception:
            pass
    if not name:
        return jsonify({"exists": False}), 200
    fpath, _ = _find_criteria_file_for_candidate(name)
    return jsonify({"exists": fpath is not None}), 200


@app.get("/sourcing/download_criteria_pdf")
def download_criteria_pdf():
    """Download the criteria JSON file for the given candidate as a PDF.

    Query params:
        linkedin  – candidate LinkedIn URL
        name      – candidate name (used for file lookup)
    Returns the PDF file as an attachment, or 404 if no criteria file found.
    """
    name = (request.args.get("name") or "").strip()
    linkedin = (request.args.get("linkedin") or "").strip()
    if not name and not linkedin:
        return "name or linkedin required", 400
    # Look up name from sourcing table if not supplied
    if not name and linkedin:
        try:
            conn = _pg_connect()
            try:
                cur = conn.cursor()
                cur.execute(
                    "SELECT name FROM sourcing WHERE linkedinurl=%s AND name IS NOT NULL LIMIT 1",
                    (linkedin,)
                )
                row = cur.fetchone()
                cur.close()
                if row and row[0]:
                    name = row[0].replace("님", "").strip()
            finally:
                conn.close()
        except Exception as exc:
            logger.warning(f"[download_criteria_pdf] DB lookup failed: {exc}")
    if not name:
        return "No candidate name found", 404
    fpath, record = _find_criteria_file_for_candidate(name)
    if not record:
        return "No criteria file found for this candidate", 404
    try:
        pdf_bytes = _criteria_record_to_pdf_bytes(record)
    except Exception as exc:
        logger.exception("[download_criteria_pdf] PDF generation failed")
        return f"PDF generation failed: {exc}", 500
    safe_role = re.sub(r'[<>:"/\\|?*]', '_', record.get("role_tag", "criteria"))
    safe_name = re.sub(r'[<>:"/\\|?*]', '_', name)
    fname = f"{safe_role} {safe_name}.pdf"
    from flask import Response as _Response
    return _Response(
        pdf_bytes,
        mimetype="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


# ---------------------------------------------------------------------------
# Assessment Report Generation (criteria JSON + bulk assessment results → PDF)
# ---------------------------------------------------------------------------

def _enrich_assessment_with_db_vskillset(result: dict, linkedin_url: str = "", process_id=None) -> dict:
    """Attach vskillset and/or category_appraisals from the DB when they are missing from result.

    Two enrichment passes:
    1. vskillset — read from the DB `vskillset` column and attach when absent from result.
       This fixes LinkedIn / SourcingVerify reports where the individual assessment file was
       written before vskillset was added by the caller.
    2. category_appraisals / scoring fields — merged from the DB `rating` column when the
       result's category_appraisals is empty or absent (happens when bulk assessment ran
       before CV data was written to DB, causing _core_assess_profile to exit early with no
       active criteria and no weight breakdown).
    """
    if not isinstance(result, dict):
        return result

    needs_vskillset = not result.get("vskillset")
    needs_appraisals = not result.get("category_appraisals")

    if not needs_vskillset and not needs_appraisals:
        return result
    if not linkedin_url and not process_id:
        return result

    try:
        conn = _pg_connect()
        try:
            cur = conn.cursor()
            if process_id:
                cur.execute(
                    "SELECT vskillset, rating FROM process WHERE id = %s LIMIT 1",
                    (process_id,)
                )
            else:
                cur.execute(
                    "SELECT vskillset, rating FROM process WHERE linkedinurl = %s"
                    " ORDER BY rating_updated_at DESC NULLS LAST LIMIT 1",
                    (linkedin_url,)
                )
            row = cur.fetchone()
            cur.close()
            if row:
                result = dict(result)
                # Pass 1: attach vskillset from DB vskillset column
                if needs_vskillset and row[0]:
                    db_vs = row[0]
                    if isinstance(db_vs, str):
                        try:
                            db_vs = json.loads(db_vs)
                        except Exception:
                            db_vs = None
                    if db_vs:
                        result["vskillset"] = db_vs
                # Pass 2: merge scoring fields from DB rating when appraisals are missing.
                # Fields merged: the full scoring breakdown (category_appraisals, criteria,
                # total_score, stars) and display fields (assessment_level, comments,
                # overall_comment). Fields intentionally excluded: vskillset (handled above),
                # is_level2 (re-derived), file (disk path, not relevant to report).
                if needs_appraisals and row[1]:
                    db_rating = row[1]
                    if isinstance(db_rating, str):
                        try:
                            db_rating = json.loads(db_rating)
                        except Exception:
                            db_rating = None
                    if isinstance(db_rating, dict) and db_rating.get("category_appraisals"):
                        for _merge_key in ("category_appraisals", "criteria", "total_score",
                                           "stars", "assessment_level", "comments",
                                           "overall_comment"):
                            if db_rating.get(_merge_key) and not result.get(_merge_key):
                                result[_merge_key] = db_rating[_merge_key]
        finally:
            conn.close()
    except Exception as _enrich_exc:
        logger.debug(f"[enrich_assessment] DB enrichment failed (non-fatal): {_enrich_exc}")
    return result


def _find_assessment_for_candidate(linkedin_url: str):
    """Find the latest assessment result for a candidate.

    Checks (in order):
      1. OUTPUT_DIR/assessments/assessment_{sha256[:16]}*.json  (individual assessment — may include username suffix)
      2. OUTPUT_DIR/bulk_*_results*.json                        (most-recent bulk run)

    Returns the assessment result dict, or None if not found.
    """
    if not linkedin_url:
        return None
    # 1. Individual assessment file (written by gemini_assess_profile / _assess_and_persist)
    # Match both legacy `assessment_{hash}.json` and new `assessment_{hash}_{username}.json`.
    _hash_prefix = "assessment_" + hashlib.sha256(linkedin_url.encode("utf-8")).hexdigest()[:16]
    assess_dir = os.path.join(OUTPUT_DIR, "assessments")
    _assess_path_legacy = os.path.join(assess_dir, _hash_prefix + ".json")
    try:
        # Try exact legacy name first for speed
        if os.path.exists(_assess_path_legacy):
            with open(_assess_path_legacy, "r", encoding="utf-8") as fh:
                result = json.load(fh)
            return _enrich_assessment_with_db_vskillset(result, linkedin_url=linkedin_url)
        # Scan for files with username suffix (new naming)
        if os.path.isdir(assess_dir):
            candidates = [f for f in os.listdir(assess_dir) if f.startswith(_hash_prefix) and f.endswith(".json")]
            if candidates:
                # Pick most recently modified
                candidates.sort(key=lambda f: os.path.getmtime(os.path.join(assess_dir, f)), reverse=True)
                with open(os.path.join(assess_dir, candidates[0]), "r", encoding="utf-8") as fh:
                    result = json.load(fh)
                return _enrich_assessment_with_db_vskillset(result, linkedin_url=linkedin_url)
    except Exception:
        pass
    # 2. Bulk results files — scan newest first
    if not os.path.isdir(OUTPUT_DIR):
        return None
    try:
        bulk_files = sorted(
            [f for f in os.listdir(OUTPUT_DIR) if f.endswith("_results.json")],
            key=lambda f: os.path.getmtime(os.path.join(OUTPUT_DIR, f)),
            reverse=True,
        )
        norm_url = linkedin_url.strip().lower()
        for fname in bulk_files:
            fpath = os.path.join(OUTPUT_DIR, fname)
            try:
                with open(fpath, "r", encoding="utf-8") as fh:
                    records = json.load(fh)
                if not isinstance(records, list):
                    continue
                for rec in records:
                    if isinstance(rec, dict) and rec.get("linkedinurl", "").strip().lower() == norm_url:
                        result = rec.get("result")
                        if result and not result.get("_skipped"):
                            # Attach vskillset from the record sibling if not already present
                            if "vskillset" not in result and rec.get("vskillset"):
                                result = dict(result)
                                result["vskillset"] = rec.get("vskillset")
                            return _enrich_assessment_with_db_vskillset(result, linkedin_url=linkedin_url)
            except Exception:
                continue
    except Exception:
        pass
    # 3. DB fallback — read the `rating` column from the process table.
    #    This covers candidates assessed via the individual path before file-writing
    #    was added, and acts as a safety net when the on-disk file is missing.
    try:
        conn = _pg_connect()
        try:
            cur = conn.cursor()
            cur.execute(
                "SELECT rating FROM process WHERE linkedinurl = %s AND rating IS NOT NULL"
                " ORDER BY rating_updated_at DESC NULLS LAST LIMIT 1",
                (linkedin_url,)
            )
            row = cur.fetchone()
            cur.close()
            if row and row[0]:
                rating = row[0]
                if isinstance(rating, str):
                    try:
                        rating = json.loads(rating)
                    except Exception:
                        rating = None
                if isinstance(rating, dict) and not rating.get("_skipped"):
                    return _enrich_assessment_with_db_vskillset(rating, linkedin_url=linkedin_url)
        finally:
            conn.close()
    except Exception:
        pass
    return None


def _lines_to_pdf_bytes(lines: list) -> bytes:
    """Convert a list of (kind, text) tuples to PDF bytes.

    Kinds: 'title', 'section', 'key', 'item', 'gap', 'table'
    For 'table' kind, text is a list of rows (list of strings). The first row
    is rendered as a bold header.  Cell text is automatically word-wrapped so
    content never overflows the page boundary.

    Tries reportlab first; falls back to a raw minimal PDF.
    """
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4
        import io as _io

        buf = _io.BytesIO()
        c = rl_canvas.Canvas(buf, pagesize=A4)
        page_w, page_h = A4
        MARGIN_L = 40
        MARGIN_R = 30
        CONTENT_W = page_w - MARGIN_L - MARGIN_R  # usable width
        y = page_h - 50

        def _safe(s):
            import re as _re
            s = str(s)
            # Strip CJK and other non-Latin-1 character runs to prevent '?'
            # placeholders when encoding to Latin-1.
            s = _re.sub(
                r'[\u0400-\u04ff\u0600-\u06ff\u0900-\u0dff\u3040-\u30ff'
                r'\u3400-\u4dbf\u4e00-\u9fff\ua000-\ua4cf\uac00-\ud7af'
                r'\uf900-\ufaff\uff00-\uffef]+',
                '', s,
            )
            # Map common Unicode punctuation to Latin-1 equivalents so they
            # display correctly in Helvetica instead of rendering as '?'.
            s = s.replace('\u2022', '-')   # bullet • → dash (ASCII-safe fallback)
            s = s.replace('\u2013', '-').replace('\u2014', '-').replace('\u2015', '-')
            s = s.replace('\u2018', "'").replace('\u2019', "'")
            s = s.replace('\u201c', '"').replace('\u201d', '"')
            s = s.replace('\u2026', '...')
            s = s.replace('\u2212', '-')
            return s.encode("latin-1", errors="ignore").decode("latin-1")

        def _wrap_text(text, font, size, avail_w):
            """Word-wrap text into lines that fit within avail_w pixels."""
            safe = _safe(str(text))
            if not safe.strip():
                return [""]
            words = safe.split()
            result, cur = [], ""
            for word in words:
                test = (cur + " " + word).strip() if cur else word
                if c.stringWidth(test, font, size) <= avail_w:
                    cur = test
                else:
                    if cur:
                        result.append(cur)
                    # Single word too long: truncate to fit
                    while word and c.stringWidth(word, font, size) > avail_w and len(word) > 1:
                        word = word[:-1]
                    cur = word
            if cur:
                result.append(cur)
            return result or [""]

        def _new_page():
            nonlocal y
            c.showPage()
            y = page_h - 50

        def _ensure_space(needed):
            nonlocal y
            if y - needed < 60:
                _new_page()

        def _draw_table(rows):
            """Draw a table with word-wrapped cells and dynamic row heights."""
            nonlocal y
            if not rows:
                return
            col_count = max(len(r) for r in rows)
            if col_count == 0:
                return

            # Column widths by column count
            if col_count == 4:
                # CATEGORY | WEIGHT | RATING/STATUS | COMMENT
                col_w = [105, 45, 82, CONTENT_W - 105 - 45 - 82]
            elif col_count == 3:
                col_w = [100, 75, CONTENT_W - 100 - 75]
            elif col_count == 2:
                col_w = [130, CONTENT_W - 130]
            else:
                col_w = [CONTENT_W / col_count] * col_count

            x_pos = [MARGIN_L]
            for cw in col_w[:-1]:
                x_pos.append(x_pos[-1] + cw)

            FONT_SZ = 9
            LINE_H = FONT_SZ + 3
            PAD_H = 3   # horizontal padding per side
            PAD_V = 4   # vertical padding per side

            # Draw outer border around the entire table
            c.setStrokeColorRGB(0.55, 0.55, 0.55)
            c.setLineWidth(0.7)

            for r_idx, row in enumerate(rows):
                is_hdr = r_idx == 0
                font_n = "Helvetica-Bold" if is_hdr else "Helvetica"

                # Pad row to col_count
                padded = list(row) + [""] * (col_count - len(row))

                # Word-wrap each cell
                cell_wrapped = []
                for ci in range(col_count):
                    avail = col_w[ci] - 2 * PAD_H
                    cell_wrapped.append(_wrap_text(str(padded[ci]), font_n, FONT_SZ, avail))

                n_lines = max(len(cl) for cl in cell_wrapped)
                row_h = n_lines * LINE_H + 2 * PAD_V
                row_h = max(row_h, LINE_H + 2 * PAD_V)

                _ensure_space(row_h + 4)

                # Background
                if is_hdr:
                    c.setFillColorRGB(0.18, 0.36, 0.56)  # professional dark blue
                    c.rect(MARGIN_L, y - row_h, CONTENT_W, row_h, fill=1, stroke=0)
                    c.setFillColorRGB(1, 1, 1)
                elif r_idx % 2 == 1:
                    c.setFillColorRGB(0.95, 0.97, 1.0)
                    c.rect(MARGIN_L, y - row_h, CONTENT_W, row_h, fill=1, stroke=0)
                    c.setFillColorRGB(0.1, 0.1, 0.1)
                else:
                    c.setFillColorRGB(1, 1, 1)
                    c.rect(MARGIN_L, y - row_h, CONTENT_W, row_h, fill=1, stroke=0)
                    c.setFillColorRGB(0.1, 0.1, 0.1)

                # Cell text
                c.setFont(font_n, FONT_SZ)
                for ci, (cl, xp, cw) in enumerate(zip(cell_wrapped, x_pos, col_w)):
                    text_y = y - PAD_V - LINE_H + 2
                    for ln in cl:
                        c.drawString(xp + PAD_H, text_y, ln)
                        text_y -= LINE_H

                # Row bottom border
                c.setStrokeColorRGB(0.70, 0.70, 0.70)
                c.setLineWidth(0.4)
                c.line(MARGIN_L, y - row_h, MARGIN_L + CONTENT_W, y - row_h)

                c.setFillColorRGB(0, 0, 0)
                y -= row_h

            # Outer border (left + right verticals)
            c.setStrokeColorRGB(0.55, 0.55, 0.55)
            c.setLineWidth(0.6)
            y -= 4

        for kind, text in lines:
            if kind == "gap":
                y -= 10
                continue

            if kind == "table":
                _draw_table(text)
                continue

            # Regular text kinds
            if kind == "title":
                font_n, font_s = "Helvetica-Bold", 16
                indent = MARGIN_L
                _ensure_space(font_s + 14)
                c.setFont(font_n, font_s)
                c.drawString(indent, y, _safe(str(text)))
                y -= font_s + 4
                # Decorative underline
                c.setStrokeColorRGB(0.18, 0.36, 0.56)
                c.setLineWidth(1.5)
                c.line(MARGIN_L, y, MARGIN_L + CONTENT_W, y)
                # Increase gap so the next element (headline) starts well below
                # the underline and does not visually overlap it.
                y -= 14
                continue
            elif kind == "section":
                font_n, font_s = "Helvetica-Bold", 11
                indent = MARGIN_L
                y -= 4
            elif kind == "key":
                font_n, font_s = "Helvetica-Bold", 10
                indent = MARGIN_L
            else:  # "item"
                font_n, font_s = "Helvetica", 9
                indent = MARGIN_L + 8

            avail_w = CONTENT_W - (indent - MARGIN_L)
            for ln in _wrap_text(str(text), font_n, font_s, avail_w):
                _ensure_space(font_s + 6)
                c.setFont(font_n, font_s)
                c.drawString(indent, y, ln)
                y -= font_s + 3

        c.save()
        return buf.getvalue()
    except ImportError:
        pass

    # Raw minimal PDF fallback
    def _pdf_str(s):
        return s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    def _pdf_safe(s):
        import re as _re
        s = str(s)
        s = _re.sub(
            r'[\u0400-\u04ff\u0600-\u06ff\u0900-\u0dff\u3040-\u30ff'
            r'\u3400-\u4dbf\u4e00-\u9fff\ua000-\ua4cf\uac00-\ud7af'
            r'\uf900-\ufaff\uff00-\uffef]+',
            '', s,
        )
        s = s.replace('\u2013', '-').replace('\u2014', '-')
        s = s.replace('\u2018', "'").replace('\u2019', "'")
        s = s.replace('\u201c', '"').replace('\u201d', '"')
        s = s.replace('\u2026', '...')
        return s.encode("latin-1", errors="ignore").decode("latin-1")

    stream_lines = ["BT", "/F1 16 Tf", "40 792 Td"]
    for kind, text in lines:
        if kind == "table":
            rows = text
            for r_idx, row in enumerate(rows):
                row_text = " | ".join(str(c) for c in row)
                safe = _pdf_safe(row_text)
                fname = "/FB" if r_idx == 0 else "/F1"
                stream_lines.append(f"{fname} 9 Tf")
                stream_lines.append(f"({_pdf_str(safe)}) Tj")
                stream_lines.append("0 -14 Td")
            continue
        safe = _pdf_safe(str(text))
        if not safe.strip() or kind == "gap":
            stream_lines.append("0 -10 Td")
            continue
        size = 16 if kind == "title" else (12 if kind == "section" else 11)
        bold = kind in ("title", "section", "key")
        fname = "/FB" if bold else "/F1"
        stream_lines.append(f"{fname} {size} Tf")
        stream_lines.append(f"({_pdf_str(safe)}) Tj")
        stream_lines.append("0 -16 Td")
    stream_lines.append("ET")
    stream = "\n".join(stream_lines).encode("latin-1")

    def _obj(n, d, s=None):
        out = f"{n} 0 obj\n{d}\n"
        if s is not None:
            out = out.encode("latin-1") + b"stream\n" + s + b"\nendstream\nendobj\n"
            return out
        return (out + "endobj\n").encode("latin-1")

    o1 = _obj(1, "<< /Type /Catalog /Pages 2 0 R >>")
    o2 = _obj(2, "<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    o3 = _obj(3, "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Contents 4 0 R /Resources << /Font << /F1 5 0 R /FB 6 0 R >> >> >>")
    slen = len(stream)
    o4 = _obj(4, f"<< /Length {slen} >>", stream)
    o5 = _obj(5, "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    o6 = _obj(6, "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>")

    header = b"%PDF-1.4\n"
    body = o1 + o2 + o3 + o4 + o5 + o6
    offsets = []
    pos = len(header)
    for chunk in (o1, o2, o3, o4, o5, o6):
        offsets.append(pos)
        pos += len(chunk)

    xref_offset = len(header) + len(body)
    xref = "xref\n0 7\n0000000000 65535 f \n"
    for off in offsets:
        xref += f"{off:010d} 00000 n \n"
    trailer = f"trailer\n<< /Size 7 /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n"
    return header + body + (xref + trailer).encode("latin-1")


def _build_report_lines(candidate_name: str, criteria_record: dict, assessment_result: dict) -> list:
    """Build the (kind, text) lines for a full assessment report PDF.

    All data sections are rendered as structured tables so that no text
    overflows the page boundary.  Table cells word-wrap automatically.
    """
    role_tag = criteria_record.get("role_tag", "")
    saved_at = criteria_record.get("saved_at", "")
    criteria = criteria_record.get("criteria") or {}

    lines = []

    # ── Title ─────────────────────────────────────────────────────────────────
    lines.append(("title", "Assessment Report"))
    lines.append(("gap", ""))

    # ── Candidate information table ───────────────────────────────────────────
    lines.append(("table", [
        ["FIELD", "VALUE"],
        ["Candidate", candidate_name],
        ["Role", role_tag],
        ["Date", saved_at],
    ]))
    lines.append(("gap", ""))

    if assessment_result:
        # ── Assessment Summary ─────────────────────────────────────────────────
        stars = max(0, min(int(assessment_result.get("stars", 0) or 0), 5))
        star_str = ("*" * stars) + ("." * (5 - stars)) + f" ({stars}/5)"
        overall = assessment_result.get("overall_comment", "") or ""

        lines.append(("section", "ASSESSMENT SUMMARY"))
        summary_rows = [
            ["FIELD", "VALUE"],
            ["Overall Score", str(assessment_result.get("total_score", "-"))],
            ["Stars", star_str],
            ["Assessment Level", str(assessment_result.get("assessment_level", "-"))],
        ]
        if overall:
            summary_rows.append(["Overall Comment", overall])
        lines.append(("table", summary_rows))
        lines.append(("gap", ""))

        # ── Score Breakdown ────────────────────────────────────────────────────
        breakdown = assessment_result.get("criteria") or {}
        if breakdown:
            lines.append(("section", "SCORE BREAKDOWN"))
            bd_rows = [["CATEGORY", "SCORE"]]
            for cat, score in breakdown.items():
                bd_rows.append([str(cat), str(score)])
            lines.append(("table", bd_rows))
            lines.append(("gap", ""))

        # ── Category Appraisals (4-column with wrapped COMMENT) ────────────────
        appraisals = assessment_result.get("category_appraisals") or {}
        if appraisals:
            lines.append(("section", "CATEGORY APPRAISALS"))
            ap_rows = [["CATEGORY", "WEIGHT", "RATING / STATUS", "COMMENT"]]
            for cat, appraisal in appraisals.items():
                if isinstance(appraisal, dict):
                    weight = appraisal.get("weight_percent", "")
                    rating = appraisal.get("rating", "")
                    status = appraisal.get("status", "")
                    comment = appraisal.get("comment", "")
                    rating_status = f"{str(rating)} / {str(status)}" if status else str(rating)
                    ap_rows.append([
                        str(cat),
                        f"{weight}%" if weight not in (None, "") else "-",
                        rating_status,
                        str(comment),
                    ])
                else:
                    ap_rows.append([str(cat), "-", "-", str(appraisal)])
            lines.append(("table", ap_rows))
            lines.append(("gap", ""))

        # ── Skill Comments ─────────────────────────────────────────────────────
        comments_raw = assessment_result.get("comments")
        if comments_raw:
            lines.append(("section", "SKILL COMMENTS"))
            if isinstance(comments_raw, str):
                # Narrative string — render each paragraph as a row in a
                # single-column table so text stays within page boundaries.
                paras = [p.strip() for p in comments_raw.split("\n") if p.strip()]
                if paras:
                    sc_rows = [["COMMENTS"]]
                    for para in paras:
                        sc_rows.append([para])
                    lines.append(("table", sc_rows))
            elif isinstance(comments_raw, (list, tuple)):
                sc_rows = [["SKILL", "STATUS", "COMMENT"]]
                for entry in comments_raw:
                    if isinstance(entry, dict):
                        skill = str(entry.get("skill") or entry.get("category", ""))
                        match = str(entry.get("match") or entry.get("status", ""))
                        note = str(entry.get("comment") or entry.get("note", ""))
                        sc_rows.append([skill, match, note])
                    else:
                        sc_rows.append([str(entry), "", ""])
                lines.append(("table", sc_rows))
            lines.append(("gap", ""))

    # ── Search Criteria (2-column table) ──────────────────────────────────────
    if criteria:
        lines.append(("section", "SEARCH CRITERIA"))
        crit_rows = [["CRITERIA", "VALUE"]]
        for k, v in criteria.items():
            if isinstance(v, list):
                v_str = ", ".join(str(x) for x in v) if v else "-"
            else:
                v_str = str(v) if v is not None else "-"
            crit_rows.append([str(k), v_str])
        lines.append(("table", crit_rows))

    return lines


@app.get("/sourcing/has_report")
def has_report():
    """Check whether a full assessment report can be generated for the given candidate.

    Requires both a criteria JSON file AND a completed assessment result.
    Query params:
        linkedin  – candidate LinkedIn URL
        name      – candidate name
    Returns { "exists": true/false }
    """
    name = (request.args.get("name") or "").strip()
    linkedin = (request.args.get("linkedin") or "").strip()
    if not name and not linkedin:
        return jsonify({"exists": False}), 200
    # Look up name from DB if missing
    if not name and linkedin:
        try:
            conn = _pg_connect()
            try:
                cur = conn.cursor()
                cur.execute(
                    "SELECT name FROM sourcing WHERE linkedinurl=%s AND name IS NOT NULL LIMIT 1",
                    (linkedin,)
                )
                row = cur.fetchone()
                cur.close()
                if row and row[0]:
                    name = row[0].replace("님", "").strip()
            finally:
                conn.close()
        except Exception:
            pass
    if not name:
        return jsonify({"exists": False}), 200
    # criteria JSON must exist
    fpath, _ = _find_criteria_file_for_candidate(name)
    if not fpath:
        return jsonify({"exists": False}), 200
    # assessment result must also exist
    assessment = _find_assessment_for_candidate(linkedin) if linkedin else None
    return jsonify({"exists": assessment is not None}), 200


@app.get("/sourcing/check_reassess")
def check_reassess():
    """Check whether the Reassess File button should be visible for a candidate.

    Logic: query the `process` table for the given LinkedIn URL and inspect the
    `rating` JSON column's `category_appraisals` for the 7 Category Breakdown
    categories (Company, Country, Job Title, Sector, Seniority, Skillset, Tenure).

    Returns JSON:
        {
          "show_reassess": true,   # true = show the button (at least one category rating missing)
          "has_rating":    false,  # true = a completed assessment exists in process.rating
          "missing_fields": ["Seniority", "Tenure"]  # list of categories missing a rating
        }

    The button should be shown when show_reassess is true (i.e. at least one
    Category Breakdown rating is NULL / missing in category_appraisals).
    """
    linkedin = (request.args.get("linkedin") or "").strip().rstrip("/")
    if not linkedin:
        return jsonify({"show_reassess": False, "has_rating": False, "missing_fields": []}), 200

    # Normalise the URL the same way the DB does
    normalized = linkedin.lower().rstrip("/")

    # The 7 required Category Breakdown categories
    _REQUIRED_CATEGORIES = ["Company", "Country", "Job Title", "Sector", "Seniority", "Skillset", "Tenure"]

    try:
        conn = _pg_connect()
        try:
            cur = conn.cursor()
            cur.execute(
                """
                SELECT rating
                FROM process
                WHERE LOWER(TRIM(TRAILING '/' FROM linkedinurl)) = %s
                   OR normalized_linkedin = %s
                ORDER BY id DESC
                LIMIT 1
                """,
                (normalized, normalized),
            )
            row = cur.fetchone()
            cur.close()
        finally:
            conn.close()

        if not row:
            # No process record at all — show the button so the user can upload a CV
            return jsonify({"show_reassess": True, "has_rating": False, "missing_fields": _REQUIRED_CATEGORIES}), 200

        rating_raw = row[0]

        # Parse the rating JSON column
        has_rating = False
        category_appraisals = {}
        if rating_raw is not None:
            r_str = str(rating_raw).strip() if not isinstance(rating_raw, str) else rating_raw.strip()
            if r_str and r_str not in ("{}", "null", "[]", ""):
                has_rating = True
                try:
                    rating_data = json.loads(r_str) if isinstance(r_str, str) else rating_raw
                    if isinstance(rating_data, dict):
                        category_appraisals = rating_data.get("category_appraisals") or {}
                except Exception:
                    pass

        if not has_rating:
            # No assessment yet — show the button for all categories
            return jsonify({"show_reassess": True, "has_rating": False, "missing_fields": _REQUIRED_CATEGORIES}), 200

        # Check each required category for a non-null rating value in category_appraisals
        missing = []
        for cat_name in _REQUIRED_CATEGORIES:
            appraisal = category_appraisals.get(cat_name)
            if not isinstance(appraisal, dict):
                missing.append(cat_name)
                continue
            rating_val = appraisal.get("rating")
            if rating_val is None or rating_val == "":
                missing.append(cat_name)

        show_reassess = len(missing) > 0

        return jsonify({
            "show_reassess": show_reassess,
            "has_rating": has_rating,
            "missing_fields": missing,
        }), 200

    except Exception as e:
        # On any DB error, default to showing the button so the user is never blocked
        return jsonify({"show_reassess": True, "has_rating": False, "missing_fields": [], "error": str(e)}), 200


def _build_report_docx(candidate_name: str, criteria_record: dict, assessment_result: dict,
                       candidate_jobtitle: str = "") -> bytes:
    """Generate a well-structured Word document (.docx) assessment report.

    Uses python-docx tables for every data section so text is always contained
    within column boundaries — no overflow or alignment issues.

    Returns the raw .docx bytes.
    """
    import io as _io
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise RuntimeError("python-docx is not installed; cannot generate DOCX report")

    role_tag = criteria_record.get("role_tag", "")
    saved_at = criteria_record.get("saved_at", "")
    criteria = criteria_record.get("criteria") or {}

    # Format date as YYYY-MM-DD only (strip time component)
    date_str = str(saved_at)
    if "T" in date_str:
        date_str = date_str.split("T")[0]
    elif len(date_str) > 10:
        date_str = date_str[:10]

    # Use candidate's current job title when available, else fall back to role_tag
    display_role = candidate_jobtitle or role_tag

    doc = Document()

    # ── Criteria key → human-readable label (also used as appraisal category lookup) ──
    _CRIT_LABEL = {
        "job_titles": "Job Title",
        "job_title": "Job Title",
        "jobtitle_role_tag": "Job Title",
        "jobtitle": "Job Title",
        "country": "Country",
        "countries": "Country",
        "company": "Company",
        "companies": "Company",
        "sector": "Sector",
        "sectors": "Sector",
        "tenure": "Tenure",
        "min_tenure": "Tenure",
        "skills": "Skillset",
        "skillset": "Skillset",
        "seniority": "Seniority",
    }

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Styles helpers ────────────────────────────────────────────────────────
    DARK_BLUE = RGBColor(0x12, 0x36, 0x5E)   # header bg approximated as font colour
    HDR_BG = "123660"                         # dark navy hex for table header shading
    ALT_BG = "EBF0F8"                         # light blue for alternating rows
    WHITE_BG = "FFFFFF"

    def _shade_cell(cell, hex_color):
        """Apply solid background shading to a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _set_col_widths(table, widths_cm):
        """Set absolute column widths."""
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx < len(widths_cm):
                    cell.width = Cm(widths_cm[idx])

    def _add_table(headers, rows_data, col_widths_cm=None):
        """Add a formatted table with a dark header row and alternating rows."""
        all_rows = [headers] + rows_data
        tbl = doc.add_table(rows=len(all_rows), cols=len(headers))
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        for r_idx, row_vals in enumerate(all_rows):
            is_hdr = r_idx == 0
            row = tbl.rows[r_idx]
            for c_idx, val in enumerate(row_vals):
                cell = row.cells[c_idx]
                cell.text = str(val) if val is not None else ""
                para = cell.paragraphs[0]
                run = para.runs[0] if para.runs else para.add_run(cell.text)
                run.text = str(val) if val is not None else ""
                run.font.size = Pt(9)
                run.font.bold = is_hdr
                if is_hdr:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    _shade_cell(cell, HDR_BG)
                elif r_idx % 2 == 0:
                    _shade_cell(cell, ALT_BG)
                else:
                    _shade_cell(cell, WHITE_BG)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        if col_widths_cm:
            _set_col_widths(tbl, col_widths_cm)
        return tbl

    def _add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text.upper())
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_BLUE
        # Underline the section heading
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(4)
        run2 = p2.add_run("─" * 60)
        run2.font.size = Pt(7)
        run2.font.color.rgb = DARK_BLUE

    # ── Title ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Assessment Report")
    title_run.font.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = DARK_BLUE
    title_p.paragraph_format.space_after = Pt(6)

    # ── Candidate info table ──────────────────────────────────────────────────
    _add_section_heading("Candidate Information")
    _add_table(
        ["Field", "Value"],
        [
            ["Candidate", candidate_name],
            ["Role", display_role],
            ["Date", date_str],
        ],
        col_widths_cm=[4.5, 12.5],
    )
    doc.add_paragraph()

    if assessment_result:
        # ── Assessment Summary ─────────────────────────────────────────────────
        stars = max(0, min(int(assessment_result.get("stars", 0) or 0), 5))
        star_str = ("★" * stars) + ("☆" * (5 - stars)) + f"  ({stars}/5)"
        overall = str(assessment_result.get("overall_comment", "") or "")
        # Ensure overall score shows as e.g. "96%" without duplicating the % symbol
        raw_score = str(assessment_result.get("total_score", "-")).rstrip("%")
        score_display = f"{raw_score}%" if raw_score != "-" else "-"

        _add_section_heading("Assessment Summary")
        summary_data = [
            ["Overall Score", score_display],
            ["Stars", star_str],
            ["Level", str(assessment_result.get("assessment_level", "-"))],
        ]
        if overall:
            summary_data.append(["Overall Comment", overall])
        _add_table(["Field", "Value"], summary_data, col_widths_cm=[4.5, 12.5])
        doc.add_paragraph()

        # ── Category Appraisals (build early so weights are available for criteria table) ──
        appraisals = assessment_result.get("category_appraisals") or {}

        def _get_weight_for_criteria(crit_key):
            """Return weight_percent string from category_appraisals for a criteria key."""
            ap_label = _CRIT_LABEL.get(crit_key.lower(), "")
            for ap_cat, ap_val in appraisals.items():
                if (ap_label and ap_cat.lower() == ap_label.lower()) or ap_cat.lower() == crit_key.lower():
                    if isinstance(ap_val, dict):
                        w = ap_val.get("weight_percent", "")
                        if w not in (None, ""):
                            return f"{w}%"
            return ""

        if criteria:
            _add_section_heading("Search Criteria")
            crit_data = []
            for k, v in criteria.items():
                if isinstance(v, list):
                    v_str = ", ".join(str(x) for x in v) if v else "-"
                else:
                    v_str = str(v) if v is not None else "-"
                display_name = _CRIT_LABEL.get(k.lower(), str(k))
                weight_val = _get_weight_for_criteria(k)
                crit_data.append([display_name, v_str, weight_val])
            _add_table(
                ["Criteria", "Value", "Weight"],
                crit_data,
                col_widths_cm=[4.0, 10.5, 2.5],
            )
            doc.add_paragraph()

        # ── Category Appraisals ────────────────────────────────────────────────
        if appraisals:
            _add_section_heading("Category Appraisals")
            # Build reverse mapping: display_name → criteria breakdown score
            criteria_breakdown = assessment_result.get("criteria") or {}
            _label_to_keys = {}
            for _k, _v in _CRIT_LABEL.items():
                _label_to_keys.setdefault(_v, []).append(_k)

            def _get_score_for_category(display_name):
                """Return actual computed score for a category from the criteria breakdown."""
                dn_lower = display_name.lower()
                try:
                    # Try direct lowercase match
                    if dn_lower in criteria_breakdown:
                        return str(round(float(criteria_breakdown[dn_lower]), 1))
                    # Try mapped internal keys
                    for _key in _label_to_keys.get(display_name, []):
                        if _key in criteria_breakdown:
                            return str(round(float(criteria_breakdown[_key]), 1))
                except (ValueError, TypeError):
                    pass
                return "-"

            ap_data = []
            for cat, appraisal in appraisals.items():
                if isinstance(appraisal, dict):
                    score_val = _get_score_for_category(str(cat))
                    rating = str(appraisal.get("rating", "") or "")
                    status = str(appraisal.get("status", "") or "")
                    comment = str(appraisal.get("comment", "") or "")
                    rating_status = f"{rating} / {status}" if status else rating
                    ap_data.append([str(cat), score_val, rating_status, comment])
                else:
                    ap_data.append([str(cat), "-", "-", str(appraisal)])
            _add_table(
                ["Category", "Score", "Rating / Status", "Comment"],
                ap_data,
                col_widths_cm=[3.5, 1.8, 3.2, 8.5],
            )
            doc.add_paragraph()

        # ── Verified Skillset (below Category Appraisals) ─────────────────────
        vskillset = assessment_result.get("vskillset")
        if vskillset:
            _add_section_heading("Verified Skillset")
            # vskillset may be a list of dicts or a single dict
            if isinstance(vskillset, dict):
                vskillset = [vskillset]
            if isinstance(vskillset, list) and vskillset:
                vs_data = []
                for item in vskillset:
                    if isinstance(item, dict):
                        skill = str(item.get("skill", ""))
                        prob = str(item.get("probability", ""))
                        if prob:
                            prob = prob.rstrip("%") + "%"
                        cat = str(item.get("category", ""))
                        reason = str(item.get("reason", ""))
                        vs_data.append([skill, prob, cat, reason])
                    else:
                        vs_data.append([str(item), "", "", ""])
                _add_table(
                    ["Skill", "Probability", "Category", "Reason"],
                    vs_data,
                    col_widths_cm=[3.0, 2.0, 2.0, 10.0],
                )
                doc.add_paragraph()

        # ── Conclusion (formerly Skill Comments) ───────────────────────────────
        comments_raw = assessment_result.get("comments")
        if comments_raw:
            _add_section_heading("Conclusion")
            if isinstance(comments_raw, str):
                paras = [p.strip() for p in comments_raw.split("\n") if p.strip()]
                if paras:
                    _add_table(["Comments"], [[p] for p in paras], col_widths_cm=[17.0])
            elif isinstance(comments_raw, (list, tuple)):
                sc_data = []
                for entry in comments_raw:
                    if isinstance(entry, dict):
                        skill = str(entry.get("skill") or entry.get("category", ""))
                        match = str(entry.get("match") or entry.get("status", ""))
                        note = str(entry.get("comment") or entry.get("note", ""))
                        sc_data.append([skill, match, note])
                    else:
                        sc_data.append([str(entry), "", ""])
                _add_table(["Skill", "Status", "Comment"], sc_data, col_widths_cm=[4.0, 3.0, 10.0])
            doc.add_paragraph()

    elif criteria:
        # No assessment yet — still show Search Criteria
        _add_section_heading("Search Criteria")
        crit_data = []
        for k, v in criteria.items():
            if isinstance(v, list):
                v_str = ", ".join(str(x) for x in v) if v else "-"
            else:
                v_str = str(v) if v is not None else "-"
            display_name = _CRIT_LABEL.get(k.lower(), str(k))
            crit_data.append([display_name, v_str])
        _add_table(["Criteria", "Value"], crit_data, col_widths_cm=[5.5, 11.5])

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _find_assessment_for_candidate_by_id(process_id: int):
    """Find the latest assessment for a no-LinkedIn candidate by process table id.

    Scans bulk result files for a record whose 'id' or 'process_id' matches.
    Falls back to the DB `rating` column keyed by process primary key.
    """
    if not process_id:
        return None
    # Scan bulk result files
    if os.path.isdir(OUTPUT_DIR):
        try:
            bulk_files = sorted(
                [f for f in os.listdir(OUTPUT_DIR) if f.endswith("_results.json")],
                key=lambda f: os.path.getmtime(os.path.join(OUTPUT_DIR, f)),
                reverse=True,
            )
            for fname in bulk_files:
                fpath = os.path.join(OUTPUT_DIR, fname)
                try:
                    with open(fpath, "r", encoding="utf-8") as fh:
                        records = json.load(fh)
                    if not isinstance(records, list):
                        continue
                    for rec in records:
                        if not isinstance(rec, dict):
                            continue
                        rec_id = rec.get("id") or rec.get("process_id")
                        if rec_id and int(rec_id) == process_id:
                            result = rec.get("result")
                            if result and not result.get("_skipped"):
                                if "vskillset" not in result and rec.get("vskillset"):
                                    result = dict(result)
                                    result["vskillset"] = rec.get("vskillset")
                                return _enrich_assessment_with_db_vskillset(result, process_id=process_id)
                except Exception:
                    continue
        except Exception:
            pass
    # DB fallback — read rating column by process primary key
    try:
        conn = _pg_connect()
        try:
            cur = conn.cursor()
            cur.execute(
                "SELECT rating FROM process WHERE id = %s AND rating IS NOT NULL"
                " ORDER BY rating_updated_at DESC NULLS LAST LIMIT 1",
                (process_id,)
            )
            row = cur.fetchone()
            cur.close()
            if row and row[0]:
                rating = row[0]
                if isinstance(rating, str):
                    try:
                        rating = json.loads(rating)
                    except Exception:
                        rating = None
                if isinstance(rating, dict) and not rating.get("_skipped"):
                    return _enrich_assessment_with_db_vskillset(rating, process_id=process_id)
        finally:
            conn.close()
    except Exception:
        pass
    return None


def _find_assessment_for_candidate_by_name(candidate_name: str):
    """Find the latest assessment for a candidate by name (last-resort fallback).

    Only used when neither LinkedIn URL nor process_id is available.
    """
    if not candidate_name:
        return None
    norm_name = candidate_name.strip().lower()
    if os.path.isdir(OUTPUT_DIR):
        try:
            bulk_files = sorted(
                [f for f in os.listdir(OUTPUT_DIR) if f.endswith("_results.json")],
                key=lambda f: os.path.getmtime(os.path.join(OUTPUT_DIR, f)),
                reverse=True,
            )
            for fname in bulk_files:
                fpath = os.path.join(OUTPUT_DIR, fname)
                try:
                    with open(fpath, "r", encoding="utf-8") as fh:
                        records = json.load(fh)
                    if not isinstance(records, list):
                        continue
                    for rec in records:
                        if not isinstance(rec, dict):
                            continue
                        rec_name = (rec.get("name") or "").strip().lower()
                        if rec_name and rec_name == norm_name:
                            result = rec.get("result")
                            if result and not result.get("_skipped"):
                                if "vskillset" not in result and rec.get("vskillset"):
                                    result = dict(result)
                                    result["vskillset"] = rec.get("vskillset")
                                # Enrich using the process_id embedded in the record (if any)
                                _rec_pid = rec.get("process_id") or rec.get("id")
                                return _enrich_assessment_with_db_vskillset(
                                    result, process_id=_rec_pid if _rec_pid else None
                                )
                except Exception:
                    continue
        except Exception:
            pass
    return None


@app.get("/sourcing/download_report")
def download_report():
    """Generate and download a formal assessment report as a Word document (.docx).

    Combines the criteria JSON and bulk/individual assessment results into one document.
    The generated file is also saved to REPORT_TEMPLATES_DIR for record-keeping.

    Query params:
        linkedin   – candidate LinkedIn URL (optional for no-LinkedIn records)
        name       – candidate name
        process_id – process table primary key (used when linkedin is empty)
    """
    name = (request.args.get("name") or "").strip()
    linkedin = (request.args.get("linkedin") or "").strip()
    process_id_str = (request.args.get("process_id") or "").strip()
    process_id = int(process_id_str) if process_id_str.isdigit() else None
    if not name and not linkedin and not process_id:
        return "name or linkedin required", 400
    # Look up name and linkedin from DB by process_id when LinkedIn URL is absent
    if process_id and not linkedin:
        try:
            conn = _pg_connect()
            try:
                cur = conn.cursor()
                cur.execute(
                    "SELECT name, linkedinurl FROM process WHERE id = %s LIMIT 1",
                    (process_id,)
                )
                row = cur.fetchone()
                cur.close()
                if row:
                    if not name and row[0]:
                        name = row[0].replace("님", "").strip()
                    if not linkedin and row[1]:
                        linkedin = row[1].strip()
            finally:
                conn.close()
        except Exception as exc:
            logger.warning(f"[download_report] process_id lookup failed: {exc}")
    # Look up name from DB if missing (LinkedIn known)
    if not name and linkedin:
        try:
            conn = _pg_connect()
            try:
                cur = conn.cursor()
                cur.execute(
                    "SELECT name FROM sourcing WHERE linkedinurl=%s AND name IS NOT NULL LIMIT 1",
                    (linkedin,)
                )
                row = cur.fetchone()
                cur.close()
                if row and row[0]:
                    name = row[0].replace("님", "").strip()
            finally:
                conn.close()
        except Exception as exc:
            logger.warning(f"[download_report] DB lookup failed: {exc}")
    if not name:
        return "No candidate name found", 404
    # Fetch candidate's current job title from DB for the report
    candidate_jobtitle = ""
    if linkedin:
        try:
            conn = _pg_connect()
            try:
                cur = conn.cursor()
                cur.execute(
                    "SELECT jobtitle FROM sourcing WHERE linkedinurl=%s AND jobtitle IS NOT NULL"
                    " ORDER BY id DESC LIMIT 1",
                    (linkedin,)
                )
                row = cur.fetchone()
                cur.close()
                if row and row[0]:
                    candidate_jobtitle = row[0].strip()
            finally:
                conn.close()
        except Exception as exc:
            logger.warning(f"[download_report] jobtitle lookup failed: {exc}")
    elif process_id:
        try:
            conn = _pg_connect()
            try:
                cur = conn.cursor()
                cur.execute(
                    "SELECT jobtitle FROM process WHERE id = %s AND jobtitle IS NOT NULL LIMIT 1",
                    (process_id,)
                )
                row = cur.fetchone()
                cur.close()
                if row and row[0]:
                    candidate_jobtitle = row[0].strip()
            finally:
                conn.close()
        except Exception as exc:
            logger.warning(f"[download_report] process_id jobtitle lookup failed: {exc}")
    # When process_id is known, look up the candidate's current role_tag+username from
    # the process table and use that to find the correct criteria file — this prevents
    # the name-based scan from returning the wrong criteria file (e.g. cloud) when the
    # candidate's name also appears in a file for a different role.
    fpath = None
    criteria_record = None
    _rt_username = None
    _rt_role_tag = None
    if process_id or linkedin:
        try:
            _rt_conn = _pg_connect()
            try:
                _rt_cur = _rt_conn.cursor()
                if process_id:
                    _rt_cur.execute(
                        "SELECT role_tag, username FROM process WHERE id = %s LIMIT 1",
                        (process_id,)
                    )
                else:
                    _rt_cur.execute(
                        "SELECT role_tag, username FROM process WHERE linkedinurl = %s LIMIT 1",
                        (linkedin,)
                    )
                _rt_row = _rt_cur.fetchone()
                _rt_cur.close()
                if _rt_row:
                    _rt_role_tag = (_rt_row[0] or "").strip()
                    _rt_username = (_rt_row[1] or "").strip()
            finally:
                _rt_conn.close()
        except Exception as _rt_exc:
            logger.warning(f"[download_report] role_tag lookup failed: {_rt_exc}")
    if _rt_role_tag and _rt_username and os.path.isdir(CRITERIA_OUTPUT_DIR):
        # Scan criteria files: prefer exact role_tag match, then substring/prefix match
        try:
            for _cfname in os.listdir(CRITERIA_OUTPUT_DIR):
                if not _cfname.endswith(".json"):
                    continue
                _cfpath = os.path.join(CRITERIA_OUTPUT_DIR, _cfname)
                try:
                    with open(_cfpath, "r", encoding="utf-8") as _cfh:
                        _cfrec = json.load(_cfh)
                    _cf_rt = (_cfrec.get("role_tag") or "").strip()
                    _cf_un = (_cfrec.get("username") or "").strip()
                    if _cf_un == _rt_username and (
                        _cf_rt == _rt_role_tag
                        or any(
                            r.strip().lower() == _cf_rt.lower()
                            for r in _rt_role_tag.split(",")
                        )
                        or any(
                            r.strip().lower() == _rt_role_tag.lower()
                            for r in _cf_rt.split(",")
                        )
                    ):
                        criteria_record = _cfrec
                        fpath = _cfpath
                        break
                except Exception:
                    continue
        except Exception as _scan_exc:
            logger.warning(f"[download_report] criteria scan failed: {_scan_exc}")
    # Fallback: scan criteria files by candidate name (last resort)
    if not criteria_record:
        fpath, criteria_record = _find_criteria_file_for_candidate(name)
    if not criteria_record:
        return "No criteria file found for this candidate", 404
    assessment_result = None
    if linkedin:
        assessment_result = _find_assessment_for_candidate(linkedin)
    if assessment_result is None and process_id:
        assessment_result = _find_assessment_for_candidate_by_id(process_id)
    if assessment_result is None and name:
        assessment_result = _find_assessment_for_candidate_by_name(name)
    try:
        docx_bytes = _build_report_docx(name, criteria_record, assessment_result or {}, candidate_jobtitle=candidate_jobtitle)
    except Exception as exc:
        logger.exception("[download_report] DOCX generation failed")
        return f"Report generation failed: {exc}", 500
    safe_name = re.sub(r'[<>:"/\\|?*]', '_', name)
    safe_role = re.sub(r'[<>:"/\\|?*]', '_', criteria_record.get("role_tag", "report"))
    fname = f"{safe_name} {safe_role}.docx"
    # Persist to REPORT_TEMPLATES_DIR (already created at startup)
    try:
        out_path = os.path.join(REPORT_TEMPLATES_DIR, fname)
        with open(out_path, "wb") as fh:
            fh.write(docx_bytes)
        logger.info(f"[download_report] Saved report to {out_path}")
    except Exception as exc:
        logger.warning(f"[download_report] Could not save to templates dir: {exc}")
    from flask import Response as _Response
    return _Response(
        docx_bytes,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )